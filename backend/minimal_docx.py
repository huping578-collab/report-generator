"""Programmatic DOCX report generation driven by a Markdown skeleton template.

A Markdown template (templates/重庆项目报告模板.md) declares the report
skeleton (headings, static body texts, tables). Dynamic data sections
(overview table, height/bolt statistics, charts, conclusions) are injected at
anchor headings; when an anchor is absent the section is appended at the end.
Requires a Markdown skeleton; explicit anchors <!-- inject:... --> are preferred over keyword fallback.
"""

from __future__ import annotations

import math
import re
from contextlib import contextmanager
from contextvars import ContextVar
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from backend import markdown_skeleton
from backend import report_engine as engine

BODY_FONT = "仿宋_GB2312"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"
BODY_SIZE = Pt(10.5)

_FORMAT_CONFIG = ContextVar("report_format_config", default=None)


def _current_format_config(config=None):
    return config or _FORMAT_CONFIG.get() or markdown_skeleton.DEFAULT_CONFIG


@contextmanager
def format_context(config):
    token = _FORMAT_CONFIG.set(config)
    try:
        yield
    finally:
        _FORMAT_CONFIG.reset(token)


def _alignment(value):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[value]


def _set_rfonts(r_pr, east_asia, latin):
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)


def _set_style_font(style, style_config):
    style.font.name = style_config["latin"]
    style.font.size = Pt(float(style_config["size_pt"]))
    style.font.bold = style_config.get("bold", False)
    _set_rfonts(style._element.get_or_add_rPr(), style_config["east_asia"], style_config["latin"])


def configure_document(doc, config=None):
    config = _current_format_config(config)
    page = config["page"]
    section = doc.sections[0]
    landscape = page["orientation"] == "landscape"
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    width, height = (29.7, 21) if landscape else (21, 29.7)
    section.page_width = Cm(width)
    section.page_height = Cm(height)
    section.left_margin = Cm(page["left_margin_cm"])
    section.right_margin = Cm(page["right_margin_cm"])
    section.top_margin = Cm(page["top_margin_cm"])
    section.bottom_margin = Cm(page["bottom_margin_cm"])

    normal = doc.styles["Normal"]
    _set_style_font(normal, config["body"])
    normal.paragraph_format.alignment = _alignment(config["body"]["alignment"])
    normal.paragraph_format.line_spacing = config["body"]["line_spacing"]
    for level in range(1, 6):
        style = doc.styles[f"Heading {level}"]
        heading_config = config["heading"][str(level)]
        _set_style_font(style, heading_config)
        style.paragraph_format.alignment = _alignment(heading_config["alignment"])
        style.paragraph_format.space_before = Pt(heading_config["space_before_pt"])
        style.paragraph_format.space_after = Pt(heading_config["space_after_pt"])
        if "line_spacing" in heading_config:
            style.paragraph_format.line_spacing = heading_config["line_spacing"]
        style.paragraph_format.keep_with_next = heading_config["keep_with_next"]
        style.paragraph_format.page_break_before = heading_config["page_break_before"]


def _apply_paragraph(paragraph, style_config, *, clear_indent=False):
    fmt = paragraph.paragraph_format
    fmt.alignment = _alignment(style_config["alignment"])
    if "space_before_pt" in style_config:
        fmt.space_before = Pt(style_config["space_before_pt"])
    if "space_after_pt" in style_config:
        fmt.space_after = Pt(style_config["space_after_pt"])
    if "line_spacing" in style_config:
        fmt.line_spacing = style_config["line_spacing"]
    if "first_line_chars" in style_config:
        # first_line_chars is represented as a fixed point-size indent (not w:firstLineChars).
        fmt.first_line_indent = Pt(float(style_config["size_pt"]) * style_config["first_line_chars"])
    if "keep_with_next" in style_config:
        fmt.keep_with_next = style_config["keep_with_next"]
    if "page_break_before" in style_config:
        fmt.page_break_before = style_config["page_break_before"]
    if clear_indent:
        _clear_indent(paragraph)


def _apply_run(run, style_config, *, bold=None):
    if bold is not None:
        run.bold = bold
    elif "bold" in style_config:
        run.bold = style_config["bold"]
    run.italic = False
    run.font.size = Pt(float(style_config["size_pt"]))
    _set_run_fonts(run, style_config["east_asia"], style_config["latin"])


def _set_run_fonts(run, east_asia=BODY_FONT, latin=LATIN_FONT):
    run.font.name = latin
    run.font.color.rgb = RGBColor(0, 0, 0)
    _set_rfonts(run._element.get_or_add_rPr(), east_asia, latin)


def _clear_indent(paragraph):
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    for key in ("left", "right", "firstLine", "hanging", "leftChars", "rightChars", "firstLineChars", "hangingChars"):
        ind.set(qn(f"w:{key}"), "0")


def _warn(paragraph, level):
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level - 1))


_HEADING_NUMBER_FORMATS = {1: "%1.", 2: "%1.%2", 3: "%1.%2.%3"}
_LEADING_NUMBER_RE = re.compile(r"^\d+(?:\.\d+)*\s+")


def _ensure_heading_numbering(doc):
    """一/二/三级标题多级列表（1./1.1/1.1.1），与 Heading 1-3 样式绑定。"""
    # ponytail: 单个共享 numId；下级随上级递增自动重启（Word 默认行为，不显式写 lvlRestart）
    cached = getattr(doc, "_heading_number_id", None)
    if cached is not None:
        return cached
    numbering = doc.part.numbering_part._element
    used_abstract = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    used_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(used_abstract, default=-1) + 1
    num_id = max(used_num, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for ilvl, level in enumerate((1, 2, 3)):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "decimal")
        lvl.append(fmt)
        style = OxmlElement("w:pStyle")
        style.set(qn("w:val"), f"Heading{level}")
        lvl.append(style)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        lvl.append(suff)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), _HEADING_NUMBER_FORMATS[level])
        lvl.append(text)
        align = OxmlElement("w:lvlJc")
        align.set(qn("w:val"), "left")
        lvl.append(align)
        p_pr = OxmlElement("w:pPr")
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "0")
        indent.set(qn("w:hanging"), "0")
        p_pr.append(indent)
        lvl.append(p_pr)
        abstract.append(lvl)
    # numbering.xml 要求 abstractNum 位于 num 之前
    first_num = next((i for i, child in enumerate(numbering) if child.tag == qn("w:num")), len(numbering))
    numbering.insert(first_num, abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    doc._heading_number_id = num_id
    return num_id


def _heading(doc, text, level):
    level = max(1, min(int(level), 5))
    config = _current_format_config()
    style_config = config["heading"][str(level)]
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    _apply_paragraph(paragraph, style_config, clear_indent=True)
    _warn(paragraph, style_config["outline_level"] + 1)
    if level <= 3:
        num_id = _ensure_heading_numbering(doc)
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level - 1))
        num_pr.append(ilvl)
        num_el = OxmlElement("w:numId")
        num_el.set(qn("w:val"), str(num_id))
        num_pr.append(num_el)
        p_pr.append(num_pr)
        text = _LEADING_NUMBER_RE.sub("", str(text), count=1)
    run = paragraph.add_run(text)
    _apply_run(run, style_config)
    return paragraph


def _body(doc, text):
    style_config = _current_format_config()["body"]
    paragraph = doc.add_paragraph()
    _apply_paragraph(paragraph, style_config)
    run = paragraph.add_run(text)
    _apply_run(run, style_config)
    return paragraph


def _rate_text(value):
    return "" if value is None else f"{value:.2f}"


def _caption(doc, prefix, number, title):
    style_config = _current_format_config()["caption"]
    paragraph = doc.add_paragraph()
    _apply_paragraph(paragraph, style_config, clear_indent=True)
    run = paragraph.add_run(f"{prefix}{number} {title}")
    _apply_run(run, style_config)
    return paragraph


def add_toc(doc, toc_config=None):
    """Insert a configured TOC field, or remove TOC update state when disabled."""
    config = _current_format_config()
    toc_config = toc_config or config["toc"]
    settings = doc.settings._element
    if not toc_config["enabled"]:
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is not None:
            settings.remove(update_fields)
        return

    title = doc.add_paragraph(style="TOC Heading")
    title_config = dict(config["heading"]["1"])
    title_config.update({"alignment": "center", "bold": True})
    if toc_config.get("east_asia"):
        title_config["east_asia"] = toc_config["east_asia"]
    _apply_paragraph(title, title_config, clear_indent=True)
    title_run = title.add_run(toc_config["title"])
    _apply_run(title_run, title_config)

    paragraph = doc.add_paragraph(style="TOC Heading")
    _clear_indent(paragraph)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = (
        f" TOC \\o \"{toc_config['min_level']}-{toc_config['max_level']}\""
        " \\h \\z \\u "
    )
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))

    update_fields = settings.find(qn("w:updateFields"))
    if toc_config["update_on_open"]:
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")
    elif update_fields is not None:
        settings.remove(update_fields)


def _add_toc(doc, toc_config=None):
    return add_toc(doc, toc_config)


def _cover_paragraphs(doc, spec: str) -> None:
    """按参考 Word 模板渲染封面：| 分隔 8 段，含主标题、城市、编号、单位与日期。"""
    parts = [p.strip() for p in spec.split("|")]
    if len(parts) != 8:
        raise ValueError(f"封面标记需要 8 段（| 分隔）：{spec}")
    body_font = dict(_current_format_config()["body"])
    bold_center = {**body_font, "alignment": "center", "bold": True, "first_line_chars": 0}
    h1_font = dict(_current_format_config()["heading"]["1"])

    def _para(text, style_config):
        paragraph = doc.add_paragraph()
        _apply_paragraph(paragraph, style_config, clear_indent=True)
        run = paragraph.add_run(text)
        _apply_run(run, style_config)
        return paragraph

    # 主标题两行：22pt 黑体加粗居中
    title_font = {**bold_center, "east_asia": "黑体", "latin": "黑体", "size_pt": 22, "line_spacing": 1.5}
    _para(parts[0], title_font)
    _para(parts[1], title_font)
    # 城市名：18pt 黑体加粗居中
    _para(parts[2], {**bold_center, "east_asia": "黑体", "size_pt": 18, "line_spacing": 1.5})
    # 报告编号、项目名称、委托单位：14pt 宋体加粗（参考模板继承正文字体宋体）
    song = "宋体"
    _para(parts[3], {**bold_center, "east_asia": song, "size_pt": 14, "line_spacing": 1.5})
    left_font = {**body_font, "bold": True, "east_asia": song, "size_pt": 14, "line_spacing": 1.5, "first_line_chars": 0}
    _para(parts[4], left_font)
    _para(parts[5], left_font)
    # 公司与日期：16pt 宋体加粗居中，单倍行距
    company_font = {**bold_center, "east_asia": song, "size_pt": 16}
    _para(parts[6], company_font)
    _para(parts[7], company_font)
    # 封面独立成页
    from docx.enum.text import WD_BREAK
    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


def _notes_paragraphs(doc, spec: str) -> None:
    """注意事项页（参考模板）：标题宋体16居中，条款宋体14，“联系方式”起宋体15。"""
    parts = [p.strip() for p in spec.split("|") if p.strip()]
    if len(parts) < 2:
        raise ValueError(f"注意事项标记至少需要标题和一条内容：{spec}")
    body_font = dict(_current_format_config()["body"])
    song = "宋体"
    contact_size = 15

    def _para(text, style_config):
        paragraph = doc.add_paragraph()
        _apply_paragraph(paragraph, style_config, clear_indent=True)
        run = paragraph.add_run(text)
        _apply_run(run, style_config)
        return paragraph

    _para(parts[0], {**body_font, "east_asia": song, "size_pt": 16, "alignment": "center", "first_line_chars": 0})
    in_contact = False
    for item in parts[1:]:
        if "联系方式" in item:
            in_contact = True
        _para(item, {**body_font, "east_asia": song, "size_pt": contact_size if in_contact else 14, "first_line_chars": 0, "line_spacing": 1.5})
    from docx.enum.text import WD_BREAK
    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


def _table(doc, headers, rows, header_shading=False):
    style_config = _current_format_config()["table"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
        "justify": WD_TABLE_ALIGNMENT.CENTER,
    }[style_config["alignment"]]
    table.style = "Table Grid"

    def fill(cell, value, bold=False):
        cell.text = str(value)
        for paragraph in cell.paragraphs:
            _apply_paragraph(paragraph, style_config, clear_indent=True)
            if header_shading and bold:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:fill"), style_config["header_shading"])
                cell._tc.get_or_add_tcPr().append(shading)
            run = paragraph.runs[0]
            _apply_run(run, style_config, bold=bold and style_config["header_bold"])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        fill(table.rows[0].cells[index], header, True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            fill(cells[index], value)
    if not style_config["allow_row_break"]:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
    return table


def _picture(doc, path, width_cm, height_cm):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))


def _example_table(doc, points):
    style_config = _current_format_config()["table"]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for record in points:
        cells = table.add_row().cells
        merged = cells[0].merge(cells[1])

        def text(paragraph, value, bold=True, size=None):
            _apply_paragraph(paragraph, style_config, clear_indent=True)
            run = paragraph.add_run(value)
            _apply_run(run, style_config, bold=bold)
            if size is not None:
                run.font.size = size

        top = merged.paragraphs[0]
        text(top, f"{engine.format_station(record['electronic_station'])}（{record['height']:.2f}mm） {engine.format_station(record['raw_station'])}")
    return table


def _distribution_images(segments, stats, records, temp_dir):
    return engine.report_images(temp_dir, segments, stats, records)


def generate_guangdong_stub(city, bundle, output_dir, log=lambda _x: None):
    raise NotImplementedError("广东报告需要用户提供第五章模板；当前程序不会伪造该模板。")


def _section_overview(doc, config, segments):
    has_county = bool(segments and any(s.get("county") for s in segments))
    if has_county:
        counties = sorted({str(s.get("county", "")).strip() for s in segments if str(s.get("county", "")).strip()})
        routes = sorted({str(s.get("route", "")).strip() for s in segments if str(s.get("route", "")).strip()})
        county_text = "、".join(counties) if counties else "重庆市"
        route_text = "、".join(routes) if routes else "G210"
        _body(doc, f"本次对{county_text}{route_text}线波形梁护栏{'和'.join(['护栏横梁中心高度', '螺栓缺失'])}进行自动化检测，按采集路段汇总表划分为{len(segments)}个统计分段。{route_text}上行K2264K2325文件按原始桩号统计，其他文件按电子修正桩号统计。")
        _caption(doc, "表", "1.1", f"{county_text}检测路段情况")
        headers = ["序号", "区县", "路线编号", "路线名", "公路等级", "起点桩号", "止点桩号", "里程(km)", "总里程(km)"]
        rows = [[i, s.get("county", ""), s.get("route", "G210"), s.get("route_name", ""), s.get("grade", ""), engine.format_station(s["start"]), engine.format_station(s["end"]), f"{s['mileage']:.3f}", f"{s.get('total_mileage', s['mileage']):.3f}"] for i, s in enumerate(segments, 1)]
    else:
        routes = sorted({str(s.get("route", "")).strip() for s in segments if str(s.get("route", "")).strip()}) if segments else []
        route_text = "、".join(routes) if routes else "G210"
        _body(doc, f"本次对重庆市{route_text}线波形梁护栏{'和'.join(['护栏横梁中心高度', '螺栓缺失'])}进行自动化检测，按采集路段汇总表划分为{len(segments)}个统计分段。{route_text}上行K2264K2325文件按原始桩号统计，其他文件按电子修正桩号统计。")
        _caption(doc, "表", "1.1", f"{route_text}检测路段情况")
        headers = ["序号", "路线编号", "路线名", "公路等级", "起点桩号", "止点桩号", "里程(km)", "总里程(km)"]
        rows = [[i, s.get("route", "G210"), s.get("route_name", ""), s.get("grade", ""), engine.format_station(s["start"]), engine.format_station(s["end"]), f"{s['mileage']:.3f}", f"{s.get('total_mileage', s['mileage']):.3f}"] for i, s in enumerate(segments, 1)]
    _table(doc, headers, rows, True)


def _section_method(doc, height_stats, bolt_stats):
    if height_stats is not None:
        _body(doc, "二波按560、580、620、640mm分档，580≤h≤620mm为合格；三波按657、677、717、737mm分档，677≤h≤717mm为合格。")
    if bolt_stats is not None:
        _body(doc, "螺栓缺失率按缺失数量÷（拼接螺栓数量+连接螺栓数量+缺失数量）×100%计算。")


def _section_height(doc, segments, height_stats, height_records, images, temp_dir, drawing_id=1000):
    # 二级标题：该节位于第 3 章内，用一级会打乱 1-5 章连续编号
    _heading(doc, "波形梁护栏横梁中心高度检测结果", 2)
    from collections import defaultdict
    has_county = bool(segments and any(s.get("county") for s in segments))
    by_county = defaultdict(list)
    for idx, seg in enumerate(segments):
        stat = height_stats[idx] if idx < len(height_stats) else None
        by_county[seg.get("county", "")].append((idx, seg, stat))
    for county_idx, (county, items) in enumerate(by_county.items()):
        _heading(doc, f"{county}整体情况" if county else "整体情况", 2)
        for kind in ("二波", "三波"):
            total = sum(it[2]["types"][kind]["count"] for it in items if it[2] is not None)
            if not total:
                continue
            good = sum(it[2]["types"][kind]["bins"][2] for it in items if it[2] is not None)
            if county:
                _body(doc, f"{county}{kind}形梁护栏有效检测点共{total}个，横梁中心高度合格率为{good * 100 / total:.2f}%。")
            else:
                _body(doc, f"G210线{kind}形梁护栏有效检测点共{total}个，横梁中心高度合格率为{good * 100 / total:.2f}%。")
        labels = ["h＜560", "560≤h＜580", "580≤h≤620", "620＜h≤640", "h＞640"]
        rows = []
        for idx, seg, stat in items:
            if stat is None:
                continue
            data = stat["types"]["二波"] if stat["types"]["二波"]["count"] else None
            if data:
                county_val = seg.get("county", "")
                route_val = seg.get("route", "G210")
                if has_county:
                    rows.append([county_val, route_val, engine.format_station(stat["segment"]["start"]), engine.format_station(stat["segment"]["end"]), *[f"{value:.2f}%" for value in data["pcts"]]])
                else:
                    rows.append([route_val, engine.format_station(stat["segment"]["start"]), engine.format_station(stat["segment"]["end"]), *[f"{value:.2f}%" for value in data["pcts"]]])
        if rows:
            _caption(doc, "表", f"3.{county_idx+1}.1", f"{county}二波形梁护栏横梁中心高度检测结果" if county else "G210线二波形梁护栏横梁中心高度检测结果")
            if has_county:
                _table(doc, ["区县", "路线编号", "起点桩号", "终点桩号", *labels], rows, True)
            else:
                _table(doc, ["路线编号", "起点桩号", "终点桩号", *labels], rows, True)
        labels = ["h＜657", "657≤h＜677", "677≤h≤717", "717＜h≤737", "h＞737"]
        rows = []
        for idx, seg, stat in items:
            if stat is None:
                continue
            data = stat["types"]["三波"] if stat["types"]["三波"]["count"] else None
            if data:
                county_val = seg.get("county", "")
                route_val = seg.get("route", "G210")
                if has_county:
                    rows.append([county_val, route_val, engine.format_station(stat["segment"]["start"]), engine.format_station(stat["segment"]["end"]), *[f"{value:.2f}%" for value in data["pcts"]]])
                else:
                    rows.append([route_val, engine.format_station(stat["segment"]["start"]), engine.format_station(stat["segment"]["end"]), *[f"{value:.2f}%" for value in data["pcts"]]])
        if rows:
            _caption(doc, "表", f"3.{county_idx+1}.2", f"{county}三波形梁护栏横梁中心高度检测结果" if county else "G210线三波形梁护栏横梁中心高度检测结果")
            if has_county:
                _table(doc, ["区县", "路线编号", "起点桩号", "终点桩号", *labels], rows, True)
            else:
                _table(doc, ["路线编号", "起点桩号", "终点桩号", *labels], rows, True)
        for idx, seg, stat in items:
            if stat is None or not any(stat["types"][kind]["count"] for kind in ("二波", "三波")):
                continue
            route_val = seg.get("route", "G210")
            _heading(doc, f"{route_val}线{engine.format_station(seg['start'])}～{engine.format_station(seg['end'])}段", 2)
            for kind in ("二波", "三波"):
                data = stat["types"][kind]
                if not data["count"]:
                    continue
                level = "较高" if data["pass"] >= 80 else ("一般" if data["pass"] >= 60 else "偏低")
                detail = engine.percentage_phrases(kind, data["pcts"])
                _body(doc, f"{route_val}线{engine.format_station(seg['start'])}～{engine.format_station(seg['end'])}段{kind}形梁护栏横梁中心高度有效检测点共{data['count']}个，整体合格率{level}，合格率为{data['pass']:.2f}%。护栏横梁中心高度{detail}。")
                image_set = images.get((idx, kind)) if isinstance(images, dict) else None
                if image_set is None:
                    continue
                drawing_id += 1
                _picture(doc, image_set["line"], 13, 8)
                _caption(doc, "图", f"3.{idx + 2}-{1}", f"{kind}形梁护栏横梁中心高度检测结果")
                _picture(doc, image_set["pie"], 14, 8.5)
                _caption(doc, "图", f"3.{idx + 2}-{2}", f"{kind}形梁护栏横梁中心高度分布情况")
                example_rows = [record for record in (height_records or []) if record["segment"] == idx and record["kind"] == kind]
                example_points = engine.select_height_example_points(example_rows, idx, kind)
                if example_points:
                    _example_table(doc, example_points)
                    _caption(doc, "图", f"3.{idx + 2}-{3}", f"{kind}形梁护栏横梁中心高度自动计算示例")
    return drawing_id


def _section_bolt(doc, segments, bolt_stats, bolt_records, disease_image_index, temp_dir):
    _heading(doc, "波形梁护栏螺栓缺失", 1)
    from collections import defaultdict
    has_county = bool(segments and any(s.get("county") for s in segments))
    by_county = defaultdict(list)
    for idx, seg in enumerate(segments):
        stat = bolt_stats[idx] if idx < len(bolt_stats) else None
        by_county[seg.get("county", "")].append((idx, seg, stat))
    for county_idx, (county, items) in enumerate(by_county.items()):
        _heading(doc, f"{county}整体情况" if county else "整体情况", 2)
        total_splice = sum(it[2]["splice"] for it in items if it[2] is not None)
        total_connection = sum(it[2]["connection"] for it in items if it[2] is not None)
        total_missing = sum(it[2]["missing"] for it in items if it[2] is not None)
        total_rate = engine.bolt_missing_rate(total_splice, total_connection, total_missing)
        if county:
            _body(doc, f"本次采用波形梁护栏螺栓缺失自动化检测方式，对{county}波形梁护栏螺栓缺失情况进行统计，共检出拼接螺栓{total_splice}颗，连接螺栓{total_connection}颗，缺失螺栓{total_missing}颗，整体缺失率为{_rate_text(total_rate)}%。")
        else:
            _body(doc, f"本次采用波形梁护栏螺栓缺失自动化检测方式，对重庆市G210线波形梁护栏螺栓缺失情况进行统计，共检出拼接螺栓{total_splice}颗，连接螺栓{total_connection}颗，缺失螺栓{total_missing}颗，整体缺失率为{_rate_text(total_rate)}%。")
        overall = []
        for idx, seg, stat in items:
            if stat is None:
                continue
            county_val = seg.get("county", "")
            route_val = seg.get("route", "G210")
            if has_county:
                overall.append([county_val, route_val, engine.format_station(stat["segment"]["start"]), engine.format_station(stat["segment"]["end"]), f"{stat['segment']['mileage']:.3f}", stat["splice"], stat["connection"], stat["missing"], f"{_rate_text(stat['rate'])}"])
            else:
                overall.append([route_val, engine.format_station(stat["segment"]["start"]), engine.format_station(stat["segment"]["end"]), f"{stat['segment']['mileage']:.3f}", stat["splice"], stat["connection"], stat["missing"], f"{_rate_text(stat['rate'])}"])
        if has_county:
            overall.append(["合计", "", "", "", "", total_splice, total_connection, total_missing, f"{_rate_text(total_rate)}"])
            _caption(doc, "表", f"4.{county_idx+1}.1", f"{county}波形梁护栏螺栓缺失检测结果" if county else "G210线波形梁护栏螺栓缺失检测结果")
            _table(doc, ["区县", "路线编号", "起点桩号", "止点桩号", "检测里程（km）", "拼接螺栓（颗）", "连接螺栓（颗）", "缺失数量（颗）", "缺失率（%）"], overall, True)
        else:
            overall.append(["合计", "", "", "", total_splice, total_connection, total_missing, f"{_rate_text(total_rate)}"])
            _caption(doc, "表", f"4.{county_idx+1}.1", f"{county}波形梁护栏螺栓缺失检测结果" if county else "G210线波形梁护栏螺栓缺失检测结果")
            _table(doc, ["路线编号", "起点桩号", "止点桩号", "检测里程（km）", "拼接螺栓（颗）", "连接螺栓（颗）", "缺失数量（颗）", "缺失率（%）"], overall, True)
        for idx, seg, stat in items:
            if stat is None:
                continue
            route_val = seg.get("route", "G210")
            section_number = idx + 2
            _heading(doc, f"{route_val}线{engine.format_station(seg['start'])}～{engine.format_station(seg['end'])}段", 2)
            if not stat["points"]:
                _body(doc, "本段无波形护栏")
                continue
            _body(doc, f"{route_val}线{engine.format_station(seg['start'])}～{engine.format_station(seg['end'])}段共检出拼接螺栓{stat['splice']}颗，连接螺栓{stat['connection']}颗，缺失螺栓{stat['missing']}颗，缺失率为{_rate_text(stat['rate'])}%。")
            _caption(doc, "表", f"4.{section_number}-1", f"{route_val}线{engine.format_station(seg['start'])}～{engine.format_station(seg['end'])}段波形梁护栏螺栓缺失检测结果")
            county_val = seg.get("county", "")
            if has_county:
                _table(doc, ["区县", "路线编号", "起点桩号", "止点桩号", "里程（km）", "拼接螺栓（颗）", "连接螺栓（颗）", "缺失数量（颗）"], [[county_val, route_val, engine.format_station(seg["start"]), engine.format_station(seg["end"]), f"{seg['mileage']:.3f}", stat["splice"], stat["connection"], stat["missing"]]], True)
            else:
                _table(doc, ["路线编号", "起点桩号", "止点桩号", "里程（km）", "拼接螺栓（颗）", "连接螺栓（颗）", "缺失数量（颗）"], [[route_val, engine.format_station(seg["start"]), engine.format_station(seg["end"]), f"{seg['mileage']:.3f}", stat["splice"], stat["connection"], stat["missing"]]], True)
            segment_bolt_rows = [record for record in (bolt_records or []) if record["segment"] == idx]
            bolt_examples = engine.select_bolt_example_points(segment_bolt_rows, disease_image_index)
            bolt_examples = [example for example in bolt_examples if example.get("image") is not None]
            if bolt_examples:
                for example in bolt_examples:
                    image_data, image_extension = engine.read_disease_image(example["image"])
                    path = Path(temp_dir) / f"bolt_{section_number}.{image_extension.lstrip('.')}"
                    path.write_bytes(image_data)
                    _picture(doc, path, 16, 8.5)
                    _body(doc, engine.bolt_example_text(example))
                _caption(doc, "图", f"4.{section_number}-1", f"{route_val}线{engine.format_station(seg['start'])}～{engine.format_station(seg['end'])}段波形梁护栏螺栓缺失自动识别示例")


def _section_tci(doc, segments, tci_stats):
    _heading(doc, "沿线设施技术状况评价", 1)
    if tci_stats is None:
        _body(doc, "未提供 TCI 病害清单，沿线设施技术状况未评定。")
        return
    from collections import defaultdict
    has_county = bool(segments and any(s.get("county") for s in segments))
    by_county = defaultdict(list)
    for idx, seg in enumerate(segments):
        st = tci_stats[idx] if idx < len(tci_stats) else None
        by_county[seg.get("county", "")].append((idx, seg, st))
    for county_idx, (county, items) in enumerate(by_county.items()):
        _heading(doc, f"{county}整体情况" if county else "整体情况", 2)
        # 县内汇总：平均 TCI 或按总扣分重新计算？ 此处取算术平均并按等级分布
        vals = [it[2]["tci"] for it in items if it[2] is not None]
        if vals:
            avg = sum(vals)/len(vals)
            # 等级分布
            grades = [it[2]["grade"] for it in items if it[2] is not None]
            cnt = {g: grades.count(g) for g in ["优","良","中","次","差"]}
            _body(doc, f"{county or 'G210线'}沿线设施技术状况共评定{len(vals)}段，平均 TCI{avg:.2f}，等级分布：优{cnt['优']}段、良{cnt['良']}段、中{cnt['中']}段、次{cnt['次']}段、差{cnt['差']}段。")
        rows = []
        for idx, seg, st in items:
            if st is None:
                continue
            if has_county:
                rows.append([seg.get("county",""), seg.get("route","G210"), engine.format_station(seg["start"]), engine.format_station(seg["end"]), f"{seg['mileage']:.3f}", st["light"], st["heavy"], st["sign"], st["marking"], f"{st['tci']:.2f}", st["grade"]])
            else:
                rows.append([seg.get("route","G210"), engine.format_station(seg["start"]), engine.format_station(seg["end"]), f"{seg['mileage']:.3f}", st["light"], st["heavy"], st["sign"], st["marking"], f"{st['tci']:.2f}", st["grade"]])
        if rows:
            _caption(doc, "表", f"2.{county_idx+1}.1", f"{county}沿线设施技术状况评价结果" if county else "沿线设施技术状况评价结果")
            if has_county:
                _table(doc, ["区县","路线编号","起点桩号","止点桩号","里程(km)","防护-轻","防护-重","标志缺损","标线缺损(m)","TCI","等级"], rows, True)
            else:
                _table(doc, ["路线编号","起点桩号","止点桩号","里程(km)","防护-轻","防护-重","标志缺损","标线缺损(m)","TCI","等级"], rows, True)
        # 逐段明细文字
        for idx, seg, st in items:
            if st is None:
                continue
            _body(doc, f"{seg.get('route','G210')}线{engine.format_station(seg['start'])}～{engine.format_station(seg['end'])}段：轻{st['light']}处、重{st['heavy']}处、标志{st['sign']}处、标线{st['marking']}m，TCI{st['tci']:.2f}（{st['grade']}）。")


def _section_conclusion(doc, segments, height_stats, bolt_stats):
    _heading(doc, "结论与建议", 1)
    _heading(doc, "结论", 2)
    has_county = bool(segments and any(s.get("county") for s in segments))
    if has_county:
        from collections import defaultdict
        by_county: dict[str, list[int]] = defaultdict(list)
        for idx, seg in enumerate(segments):
            by_county[seg.get("county", "")].append(idx)
        for county, idxs in by_county.items():
            label = county if county else "G210"
            if height_stats is not None:
                for kind in ("二波", "三波"):
                    total = sum(height_stats[i]["types"][kind]["count"] for i in idxs if i < len(height_stats) and height_stats[i] is not None)
                    if not total:
                        continue
                    good = sum(height_stats[i]["types"][kind]["bins"][2] for i in idxs if i < len(height_stats) and height_stats[i] is not None)
                    _body(doc, f"{label}检测路段{kind}形梁护栏有效检测点{total}个，合格点{good}个，整体合格率{good * 100 / total:.2f}%。")
            if bolt_stats is not None:
                total_splice = sum(bolt_stats[i]["splice"] for i in idxs if i < len(bolt_stats) and bolt_stats[i] is not None)
                total_connection = sum(bolt_stats[i]["connection"] for i in idxs if i < len(bolt_stats) and bolt_stats[i] is not None)
                total_missing = sum(bolt_stats[i]["missing"] for i in idxs if i < len(bolt_stats) and bolt_stats[i] is not None)
                total_rate = engine.bolt_missing_rate(total_splice, total_connection, total_missing)
                _body(doc, f"{label}检测路段共检出拼接螺栓{total_splice}颗、连接螺栓{total_connection}颗，缺失螺栓{total_missing}颗，整体缺失率为{_rate_text(total_rate)}%。")
    else:
        if height_stats is not None:
            for kind in ("二波", "三波"):
                total = sum(item["types"][kind]["count"] for item in height_stats)
                if not total:
                    continue
                good = sum(item["types"][kind]["bins"][2] for item in height_stats)
                _body(doc, f"G210检测路段{kind}形梁护栏有效检测点{total}个，合格点{good}个，整体合格率{good * 100 / total:.2f}%。")
        if bolt_stats is not None:
            total_splice = sum(item["splice"] for item in bolt_stats)
            total_connection = sum(item["connection"] for item in bolt_stats)
            total_missing = sum(item["missing"] for item in bolt_stats)
            total_rate = engine.bolt_missing_rate(total_splice, total_connection, total_missing)
            _body(doc, f"G210检测路段共检出拼接螺栓{total_splice}颗、连接螺栓{total_connection}颗，缺失螺栓{total_missing}颗，整体缺失率为{_rate_text(total_rate)}%。")
    _heading(doc, "建议", 2)
    _body(doc, "建议管养单位优先对横梁中心高度合格率较低的分段开展现场复核，结合路缘石、路面加铺及护栏结构实际情况制定整治计划，并在养护后复测；加强护栏连接件养护巡查，对螺栓缺失位置及时补装同规格螺栓并复核紧固状态。")


def _report_with_skeleton(doc, config, blocks, segments, height_stats, height_records, bolt_stats, bolt_records, tci_stats, tci_records, disease_image_index, images, temp_dir, skeleton_md=None):
    # 显式锚点优先，关键词回退。显式锚点语法（任一满足即注入）：
    #   <!-- inject:overview -->  <!-- inject:height -->  <!-- inject:bolt -->  <!-- inject:conclusion -->
    # 兼容旧模板的关键词匹配：项目概况/整体情况/螺栓/结论
    import re as _re
    anchor_re = _re.compile(r"<!--\s*inject:\s*(overview|tci|height|bolt|conclusion)\s*-->")
    keyword_map = {
        "项目概况": "overview",
        "沿线设施技术状况评价": "tci",
        "沿线设施": "tci",
        "整体情况": "height",
        "螺栓": "bolt",
        "结论": "conclusion",
    }
    def _writer_for(key: str):
        if key == "overview":
            return lambda: _section_overview(doc, config, segments)
        if key == "tci":
            return (lambda: _section_tci(doc, segments, tci_stats)) if tci_stats is not None else (lambda: None)
        if key == "height":
            return (lambda: _section_height(doc, segments, height_stats, height_records, images, temp_dir)) if height_stats is not None else (lambda: None)
        if key == "bolt":
            return (lambda: _section_bolt(doc, segments, bolt_stats, bolt_records, disease_image_index, temp_dir)) if bolt_stats is not None else (lambda: None)
        if key == "conclusion":
            return lambda: _section_conclusion(doc, segments, height_stats, bolt_stats)
        return lambda: None

    # pending 按显式 key 索引，便于锚点直接命中
    pending_by_key = {
        "overview": ("项目概况", True, _writer_for("overview")),
        "tci": ("沿线设施技术状况评价", False, _writer_for("tci")),
        "height": ("整体情况", False, _writer_for("height")),
        "bolt": ("螺栓", False, _writer_for("bolt")),
        "conclusion": ("结论", False, _writer_for("conclusion")),
    }
    pending_keys = set(pending_by_key.keys())

    def _inject(key: str, render_heading: bool, block_text: str = "", level: int = 2):
        if key not in pending_keys:
            return
        keyword, do_render, writer = pending_by_key[key]
        pending_keys.remove(key)
        if do_render and block_text:
            _heading(doc, block_text, level)
        writer()

    skeleton_dir = Path(skeleton_md).parent if skeleton_md else None

    for block in blocks:
        # 1) 显式锚点：任意块（标题/段落）的文本中包含 <!-- inject:xxx -->
        anchor_hit = None
        if block.text:
            m = anchor_re.search(block.text)
            if m:
                anchor_hit = m.group(1)
        if anchor_hit:
            # 显式锚点不渲染原块文本，直接注入
            _inject(anchor_hit, False)
            continue

        if block.kind == "toc":
            if _current_format_config()["toc"]["enabled"]:
                _add_toc(doc)
        elif block.kind == "cover":
            _cover_paragraphs(doc, block.text)
        elif block.kind == "notes":
            _notes_paragraphs(doc, block.text)
        elif block.kind == "heading":
            # 关键词回退
            hit_key = None
            for kw, key in keyword_map.items():
                if kw in block.text and key in pending_keys:
                    hit_key = key
                    break
            if hit_key:
                _keyword, do_render, _ = pending_by_key[hit_key]
                _inject(hit_key, do_render, block.text, block.level)
                continue
            if "G210线K" in block.text and block.level >= 3:
                continue
            _heading(doc, block.text, block.level)
        elif block.kind == "paragraph":
            if block.text:
                _body(doc, block.text)
        elif block.kind == "table":
            if block.rows:
                _table(doc, block.rows[0], block.rows[1:], False)
        elif block.kind == "picture":
            if skeleton_dir is not None:
                try:
                    media_path = (skeleton_dir / block.caption).resolve()
                    if media_path.is_file():
                        # 按 13cm 宽度插入，高度自适应；失败则跳过
                        try:
                            _picture(doc, media_path, 13, 8)
                            continue
                        except Exception:
                            pass
                except Exception:
                    pass
            # 无有效图片时跳过（报告图片由引擎生成）
            continue
    # 未命中锚点追加末尾（不丢数据但位置错，模板编辑时保留关键词或显式锚点可避免）
    for key in list(pending_keys):
        _, _, writer = pending_by_key[key]
        writer()


def make_report(config, segments, height_stats, height_records, bolt_stats, bolt_records, tci_stats, tci_records, disease_image_index, temp_dir, log=lambda _x: None, skeleton_md=None):
    """Build the full report document from computed statistics."""
    images = _distribution_images(segments, height_stats, height_records, temp_dir) if height_stats is not None else {}
    if skeleton_md is None:
        raise FileNotFoundError(f"Markdown 模板不存在：{skeleton_md}，仅支持 .md 模板。")
    template = markdown_skeleton.read_template(skeleton_md)
    doc = Document()
    with format_context(template.config):
        configure_document(doc, template.config)
        _report_with_skeleton(doc, config, template.blocks, segments, height_stats, height_records, bolt_stats, bolt_records, tci_stats, tci_records, disease_image_index, images, temp_dir, skeleton_md=skeleton_md)
    try:
        doc.save(config.out_docx)
    except PermissionError as exc:
        raise PermissionError(f"Word文件被占用：{config.out_docx}") from exc
    return config.out_docx



def run(config, segments, height_stats, height_records, bolt_stats, bolt_records, disease_image_index, log=lambda _x: None, skeleton_md=None, tci_stats=None, tci_records=None):
    import tempfile

    with tempfile.TemporaryDirectory(prefix="g210_report_builtin_") as temp_dir:
        config.output_dir.mkdir(parents=True, exist_ok=True)
        log("未检测到内置 Word 模板，切换到程序化报告生成模式。")
        return make_report(config, segments, height_stats, height_records, bolt_stats, bolt_records, tci_stats, tci_records, disease_image_index, temp_dir, log, skeleton_md=skeleton_md)
