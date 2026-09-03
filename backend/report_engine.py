from __future__ import annotations

from collections import Counter
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from queue import Empty, Queue
from threading import Thread
from zipfile import ZIP_DEFLATED, ZipFile
from concurrent.futures import ThreadPoolExecutor, as_completed
import csv
import math
import os
import posixpath
import random
import re
import shutil
import sys
import tempfile
import traceback
import xml.etree.ElementTree as ET

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import openpyxl
from openpyxl import Workbook
from openpyxl.chart import BarChart, LineChart, PieChart, Reference
from openpyxl.chart.label import DataLabel, DataLabelList
from openpyxl.chart.layout import Layout, ManualLayout
from openpyxl.chart.marker import DataPoint
from openpyxl.chart.shapes import GraphicalProperties
from openpyxl.chart.text import RichText
from openpyxl.drawing.text import CharacterProperties, Font as DrawingFont, Paragraph, ParagraphProperties
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk


X = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
XDR = "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"
PR = "http://schemas.openxmlformats.org/package/2006/relationships"
CT = "http://schemas.openxmlformats.org/package/2006/content-types"
ET.register_namespace("w", W)
ET.register_namespace("r", R)
ET.register_namespace("wp", WP)
ET.register_namespace("a", A)
ET.register_namespace("pic", PIC)
q = lambda ns, tag: f"{{{ns}}}{tag}"
wt = lambda tag: q(W, tag)

PIE_COLORS = ["4472C4", "ED7D31", "A5A5A5", "FFC000", "5B9BD5"]
LINE_COLORS = ["4472C4", "ED7D31", "A5A5A5"]
PROGRAM_NAME = "报告生成工具V0.1"
OUT_XLSX_NAME = "重庆G210护栏统计.xlsx"
OUT_DOCX_NAME = "重庆G210交安设施检测报告.docx"


@dataclass
class GuangdongConfig:
    project_dir: Path
    manual_xlsx: Path
    route_xlsx: Path
    output_dir: Path
    marking_threshold: float
    height_threshold: float
    bolt_threshold: float
    marking_dir: Path | None = None
    guardrail_dir: Path | None = None

    def __post_init__(self):
        values = validate_thresholds(self.marking_threshold, self.height_threshold, self.bolt_threshold)
        self.marking_threshold, self.height_threshold, self.bolt_threshold = values

    @property
    def thresholds(self):
        return {"marking": self.marking_threshold, "height": self.height_threshold, "bolt": self.bolt_threshold}


def validate_thresholds(marking, height, bolt):
    result = []
    for label, value in zip(("标线", "护栏高度", "螺栓缺失"), (marking, height, bolt)):
        if value is None or str(value).strip() == "":
            raise ValueError(f"{label}一致性阈值不能为空")
        try:
            number = float(value)
        except (TypeError, ValueError) as exc:
            raise ValueError(f"{label}一致性阈值必须为数值") from exc
        if not 0 <= number <= 100:
            raise ValueError(f"{label}一致性阈值必须在0～100之间")
        result.append(number)
    return tuple(result)


def application_root():
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent.parent


def resource_template_path(name):
    return application_root() / "templates" / name


def builtin_template_paths():
    """返回内置报告模板配置；新增模板时只需在此处增加名称和路径。"""
    return {
        "重庆模板": resource_template_path("重庆项目报告模板.md"),
    }


BUILTIN_REPORT_TEMPLATES = builtin_template_paths()


@dataclass
class Config:
    project_dir: Path
    summary_xlsx: Path
    detail_dir: Path
    template_docx: Path
    output_dir: Path
    disease_dir: Path | None = None

    @property
    def out_xlsx(self):
        return self.output_dir / OUT_XLSX_NAME

    @property
    def out_docx(self):
        return self.output_dir / OUT_DOCX_NAME


def station_to_m(value):
    s = str(value or "").strip().upper().replace("K", "")
    match = re.search(r"(-?\d+(?:\.\d+)?)\s*\+\s*(-?\d+(?:\.\d+)?)", s)
    if match:
        return float(match.group(1)) * 1000 + float(match.group(2))
    try:
        return float(s) * 1000
    except ValueError:
        return None


def format_station(meters):
    km = math.floor(meters / 1000)
    remainder = int(round(meters - km * 1000))
    if remainder >= 1000:
        km, remainder = km + 1, remainder - 1000
    return f"K{km}+{remainder:03d}"


def format_station_one_decimal(meters):
    """桩号米数保留1位小数，用于螺栓示例的原始桩号。"""
    km = math.floor(meters / 1000)
    remainder = round(meters - km * 1000, 1)
    if remainder >= 1000:
        km, remainder = km + 1, remainder - 1000
    return f"K{km}+{remainder:05.1f}"


def read_segments(summary_xlsx):
    wb = openpyxl.load_workbook(summary_xlsx, read_only=True, data_only=True)
    ws = wb["Sheet1"] if "Sheet1" in wb.sheetnames else wb.worksheets[0]
    # 读取 header 行（第2行）映射列名到索引，兼容旧版无区县时默认 county="" route="G210"
    header = None
    header_row_idx = 2
    for idx in (2, 1):
        if ws.max_row >= idx:
            values = [str(c.value).strip() if c.value is not None else "" for c in ws[idx]]
            if any(v for v in values):
                if any("区县" in v or "公路等级" in v or "路线编号" in v for v in values):
                    header = values
                    header_row_idx = idx
                    break
    # 若 header 含 "区县" 则按新列读，否则按旧固定索引
    if header is not None and any("区县" in h for h in header):
        def _col(keywords):
            for i, h in enumerate(header):
                for kw in keywords:
                    if kw and kw in h:
                        return i
            return None

        county_idx = _col(["区县"])
        route_idx = _col(["路线编号"])
        route_name_idx = _col(["路线名"])
        grade_idx = _col(["公路等级"])
        manager_idx = _col(["管理单位", "管养单位"])
        start_idx = _col(["起点桩号", "起点"])
        end_idx = _col(["止点桩号", "止点"])
        total_idx = _col(["总里程"])
        mileage_idx = None
        for i, h in enumerate(header):
            if "里程" in h and "总里程" not in h:
                mileage_idx = i
                break

        result = []
        for row in ws.iter_rows(min_row=header_row_idx + 1, values_only=True):
            if row is None or all(v is None for v in row):
                continue
            s_val = row[start_idx] if start_idx is not None and start_idx < len(row) else None
            e_val = row[end_idx] if end_idx is not None and end_idx < len(row) else None
            if s_val is None or e_val is None or str(s_val).strip() == "":
                continue
            try:
                start = float(s_val) * 1000
                end = float(e_val) * 1000
            except (TypeError, ValueError):
                continue
            mileage = 0.0
            if mileage_idx is not None and mileage_idx < len(row) and row[mileage_idx] not in (None, ""):
                try:
                    mileage = float(row[mileage_idx])
                except (TypeError, ValueError):
                    mileage = 0.0
            total = mileage
            if total_idx is not None and total_idx < len(row) and row[total_idx] not in (None, ""):
                try:
                    total = float(row[total_idx])
                except (TypeError, ValueError):
                    total = mileage
            county = str(row[county_idx] or "").strip() if county_idx is not None and county_idx < len(row) else ""
            route = str(row[route_idx] or "").strip() if route_idx is not None and route_idx < len(row) else "G210"
            if not route:
                route = "G210"
            route_name = str(row[route_name_idx] or "").strip() if route_name_idx is not None and route_name_idx < len(row) else ""
            grade = str(row[grade_idx] or "").strip() if grade_idx is not None and grade_idx < len(row) else ""
            manager = str(row[manager_idx] or "").strip() if manager_idx is not None and manager_idx < len(row) else ""
            result.append({
                "county": county,
                "route": route,
                "route_name": route_name,
                "grade": grade,
                "manager": manager,
                "start": start,
                "end": end,
                "mileage": mileage,
                "total_mileage": total,
            })
        wb.close()
        if not result:
            raise ValueError("分段汇总表中未读取到起点桩号和终点桩号。")
        return result
    result = []
    for row in ws.iter_rows(min_row=header_row_idx + 1 if header is not None else 3, values_only=True):
        if len(row) < 9 or row[6] is None or row[7] is None:
            continue
        try:
            start = float(row[6]) * 1000
            end = float(row[7]) * 1000
        except (TypeError, ValueError):
            continue
        result.append({
            "county": "",
            "route": "G210",
            "route_name": "",
            "grade": str(row[3] or ""),
            "manager": str(row[4] or ""),
            "start": start,
            "end": end,
            "mileage": float(row[8] or 0),
            "total_mileage": float(row[8] or 0),
        })
    wb.close()
    if not result:
        raise ValueError("分段汇总表中未读取到起点桩号和终点桩号。")
    return result


def cell_value(cell, shared_strings=None):
    if cell.get("t") == "inlineStr":
        return "".join(t.text or "" for t in cell.findall(f".//{q(X, 't')}"))
    value = cell.find(q(X, "v"))
    if value is None or value.text is None:
        return None
    if cell.get("t") == "s" and shared_strings is not None:
        try:
            return shared_strings[int(value.text)]
        except (ValueError, IndexError):
            return None
    try:
        return float(value.text)
    except ValueError:
        return value.text


def iter_height_rows(path):
    # 源文件工作表范围可能错误标记为A1，直接流式读取sheet1.xml。
    with ZipFile(path) as archive, archive.open("xl/worksheets/sheet1.xml") as stream:
        headers = None
        for _, element in ET.iterparse(stream, events=("end",)):
            if element.tag != q(X, "row"):
                continue
            values = {}
            for cell in element.findall(q(X, "c")):
                match = re.match(r"[A-Z]+", cell.get("r", ""))
                if match:
                    values[match.group(0)] = cell_value(cell)
            if headers is None:
                headers = {column: str(value) for column, value in values.items()}
            else:
                yield {headers.get(column, column): value for column, value in values.items()}
            element.clear()


def iter_bolt_rows(path):
    """定位含“拼接螺栓数量”表头的工作表并流式读取。"""
    marker = "拼接螺栓数量".encode("utf-8")
    with ZipFile(path) as archive:
        sheet_name = None
        for name in archive.namelist():
            if name.startswith("xl/worksheets/sheet") and name.endswith(".xml"):
                with archive.open(name) as probe:
                    if marker in probe.read(8192):
                        sheet_name = name
                        break
        if sheet_name is None:
            return
        with archive.open(sheet_name) as stream:
            headers = None
            for _, element in ET.iterparse(stream, events=("end",)):
                if element.tag != q(X, "row"):
                    continue
                values = {}
                for cell in element.findall(q(X, "c")):
                    match = re.match(r"[A-Z]+", cell.get("r", ""))
                    if match:
                        values[match.group(0)] = cell_value(cell)
                if headers is None:
                    headers = {column: str(value) for column, value in values.items()}
                else:
                    yield {headers.get(column, column): value for column, value in values.items()}
                element.clear()


def guardrail_type(value):
    text = str(value or "").strip()
    if "三" in text:
        return "三波"
    if any(word in text for word in ("双", "两", "二")):
        return "二波"
    return None


def bin_index(kind, height):
    limits = (560, 580, 620, 640) if kind == "二波" else (657, 677, 717, 737)
    if height < limits[0]:
        return 0
    if height < limits[1]:
        return 1
    if height <= limits[2]:
        return 2
    if height <= limits[3]:
        return 3
    return 4


def collect_records(segments, detail_dir, log=lambda _: None):
    records, excluded = [], Counter()
    files = sorted(Path(detail_dir).glob("*.xlsx"))
    files = [p for p in files if not p.name.startswith("~$") and "交安设施现场检测" in p.name]
    if not files:
        raise FileNotFoundError("明细文件夹中未找到交安设施现场检测Excel。")
    for index, path in enumerate(files, 1):
        if path.name.startswith("G210-G210-"):
            continue
        log(f"解析明细 {index}/{len(files)}：{path.name}")
        use_raw = path.name.startswith("G210上行K2264K2325-")
        for row in iter_height_rows(path):
            kind = guardrail_type(row.get("护栏类型"))
            height = row.get("梁板中心高度(mm)")
            if kind is None or not isinstance(height, (int, float)) or height <= 0:
                continue
            remark = str(row.get("异常标记") or row.get("备注") or "").strip()
            if remark not in {"", "无备注"}:
                excluded[remark] += 1
                continue
            basis = "原始桩号" if use_raw else "电子修正桩号"
            raw_station = station_to_m(row.get("原始桩号"))
            electronic_station = station_to_m(row.get("电子修正桩号"))
            station = raw_station if use_raw else electronic_station
            if station is None:
                continue
            segment = next(
                (i for i, item in enumerate(segments) if item["start"] <= station <= item["end"]),
                None,
            )
            if segment is None:
                continue
            records.append({
                "file": path.name,
                "direction": str(row.get("方向") or ""),
                "station": station,
                "raw_station": raw_station,
                "electronic_station": electronic_station,
                "basis": basis,
                "kind": kind,
                "height": float(height),
                "segment": segment,
            })
    unique = {}
    for record in records:
        key = (
            record["direction"], round(record["station"], 3),
            record["kind"], round(record["height"], 3),
        )
        unique.setdefault(key, record)
    return list(unique.values()), len(records) - len(unique), excluded


def collect_bolt_records(segments, detail_dir, log=lambda _: None):
    records = []
    files = sorted(Path(detail_dir).glob("*.xlsx"))
    files = [p for p in files if not p.name.startswith("~$") and "交安设施现场检测" in p.name]
    if not files:
        raise FileNotFoundError("明细文件夹中未找到交安设施现场检测Excel。")

    def number(value):
        try:
            return float(value or 0)
        except (TypeError, ValueError):
            return 0.0

    for index, path in enumerate(files, 1):
        log(f"解析螺栓明细 {index}/{len(files)}：{path.name}")
        use_raw = path.name.startswith("G210上行K2264K2325-")
        for row in iter_bolt_rows(path):
            basis = "原始桩号" if use_raw else "电子修正桩号"
            raw_station = station_to_m(row.get("原始桩号"))
            electronic_station = station_to_m(row.get("电子修正桩号"))
            station = raw_station if use_raw else electronic_station
            if station is None:
                continue
            segment = next(
                (i for i, item in enumerate(segments) if item["start"] <= station <= item["end"]),
                None,
            )
            if segment is None:
                continue
            splice = number(row.get("拼接螺栓数量（颗）"))
            splice_missing = number(row.get("拼接螺栓缺失数量（颗）"))
            connection = number(row.get("连接螺栓数量（颗）"))
            connection_missing = number(row.get("连接螺栓缺失数量（颗）"))
            if splice == connection == splice_missing == connection_missing == 0:
                continue
            records.append({
                "file": path.name,
                "direction": str(row.get("方向") or ""),
                "station": station,
                "raw_station": raw_station,
                "electronic_station": electronic_station,
                "basis": basis,
                "segment": segment,
                "splice": splice,
                "splice_missing": splice_missing,
                "connection": connection,
                "connection_missing": connection_missing,
            })
    unique = {}
    for record in records:
        key = (
            record["direction"], round(record["station"], 3),
            record["splice"], record["splice_missing"],
            record["connection"], record["connection_missing"],
        )
        unique.setdefault(key, record)
    return list(unique.values()), len(records) - len(unique)


def normalize_direction(value):
    text = str(value or "")
    if "上" in text:
        return "上行"
    if "下" in text:
        return "下行"
    return text.strip()


def _xlsx_shared_strings(archive):
    if "xl/sharedStrings.xml" not in archive.namelist():
        return []
    root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    return ["".join(node.text or "" for node in item.findall(f".//{q(X, 't')}")) for item in root]


def _xlsx_sheet_rows(archive, sheet_name, shared_strings):
    rows = {}
    headers = None
    with archive.open(sheet_name) as stream:
        for _, element in ET.iterparse(stream, events=("end",)):
            if element.tag != q(X, "row"):
                continue
            row_number = int(element.get("r", "0") or 0)
            values = {}
            for cell in element.findall(q(X, "c")):
                match = re.match(r"[A-Z]+", cell.get("r", ""))
                if match:
                    values[match.group(0)] = cell_value(cell, shared_strings)
            if headers is None and "原始桩号" in {str(value) for value in values.values()}:
                headers = {column: str(value) for column, value in values.items() if value is not None}
            elif headers is not None:
                rows[row_number] = {headers.get(column, column): value for column, value in values.items()}
            element.clear()
    return rows


def collect_disease_image_index(disease_dir, log=lambda _: None):
    """只读取病害清单XML关系，建立原始桩号到嵌入图片的轻量索引。"""
    disease_dir = Path(disease_dir)
    files = sorted(path for path in disease_dir.glob("*.xlsx") if not path.name.startswith("~$") and "病害清单" in path.name)
    if not files:
        raise FileNotFoundError("病害清单文件夹中未找到病害清单Excel。")
    image_index = {}
    indexed = 0
    for file_number, path in enumerate(files, 1):
        log(f"索引病害图片 {file_number}/{len(files)}：{path.name}")
        with ZipFile(path) as archive:
            shared_strings = _xlsx_shared_strings(archive)
            sheet_names = sorted(
                name for name in archive.namelist()
                if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
            )
            for sheet_name in sheet_names:
                sheet_rels_name = posixpath.join(
                    posixpath.dirname(sheet_name), "_rels", posixpath.basename(sheet_name) + ".rels"
                )
                if sheet_rels_name not in archive.namelist():
                    continue
                sheet_relationships = ET.fromstring(archive.read(sheet_rels_name))
                drawing_relation = next(
                    (relation for relation in sheet_relationships if relation.get("Type", "").endswith("/drawing")),
                    None,
                )
                if drawing_relation is None:
                    continue
                drawing_name = posixpath.normpath(posixpath.join(posixpath.dirname(sheet_name), drawing_relation.get("Target")))
                drawing_rels_name = posixpath.join(
                    posixpath.dirname(drawing_name), "_rels", posixpath.basename(drawing_name) + ".rels"
                )
                if drawing_name not in archive.namelist() or drawing_rels_name not in archive.namelist():
                    continue
                rows = _xlsx_sheet_rows(archive, sheet_name, shared_strings)
                drawing_relationships = ET.fromstring(archive.read(drawing_rels_name))
                media_by_rid = {
                    relation.get("Id"): posixpath.normpath(
                        posixpath.join(posixpath.dirname(drawing_name), relation.get("Target"))
                    )
                    for relation in drawing_relationships
                    if relation.get("Type", "").endswith("/image")
                }
                drawing = ET.fromstring(archive.read(drawing_name))
                for anchor in list(drawing):
                    anchor_from = anchor.find(q(XDR, "from"))
                    if anchor_from is None:
                        continue
                    row_node = anchor_from.find(q(XDR, "row"))
                    blip = anchor.find(f".//{q(A, 'blip')}")
                    if row_node is None or blip is None:
                        continue
                    excel_row = int(row_node.text or 0) + 1
                    row = rows.get(excel_row, {})
                    if "螺栓缺失" not in str(row.get("病害类型") or ""):
                        continue
                    raw_station = station_to_m(row.get("原始桩号"))
                    media_name = media_by_rid.get(blip.get(q(R, "embed")))
                    if raw_station is None or not media_name or media_name not in archive.namelist():
                        continue
                    try:
                        quantity = int(round(float(row.get("工程量") or 0)))
                    except (TypeError, ValueError):
                        quantity = 0
                    descriptor = {
                        "workbook": path,
                        "media": media_name,
                        "row": excel_row,
                        "direction": normalize_direction(row.get("方向")),
                        "raw_station": raw_station,
                        "quantity": quantity,
                    }
                    key = (descriptor["direction"], round(raw_station, 1))
                    image_index.setdefault(key, []).append(descriptor)
                    indexed += 1
    for descriptors in image_index.values():
        descriptors.sort(key=lambda item: (str(item["workbook"]), item["row"]))
    log(f"病害图片索引完成：{len(files)}个文件，{indexed}张螺栓缺失图片。")
    return image_index


def match_disease_image(record, image_index):
    if not image_index or record.get("raw_station") is None:
        return None
    key = (normalize_direction(record.get("direction")), round(record["raw_station"], 1))
    candidates = image_index.get(key, [])
    if not candidates:
        return None
    total_missing = int(round(record["splice_missing"] + record["connection_missing"]))
    return next((item for item in candidates if item["quantity"] == total_missing), candidates[0])


def read_disease_image(descriptor):
    with ZipFile(descriptor["workbook"]) as archive:
        data = archive.read(descriptor["media"])
    if data.startswith(b"\x89PNG\r\n\x1a\n"):
        return data, ".png"
    if data.startswith(b"\xff\xd8\xff"):
        return data, ".jpeg"
    extension = Path(descriptor["media"]).suffix.lower() or ".png"
    return data, extension


def bolt_missing_rate(splice, connection, missing):
    """螺栓缺失率：缺失数÷（现有拼接数+现有连接数+缺失数）×100%。"""
    denominator = splice + connection + missing
    return missing * 100 / denominator if denominator else 0


def make_bolt_stats(segments, records):
    stats = []
    for index, segment in enumerate(segments):
        rows = [record for record in records if record["segment"] == index]
        splice = sum(record["splice"] for record in rows)
        connection = sum(record["connection"] for record in rows)
        missing = sum(record["splice_missing"] + record["connection_missing"] for record in rows)
        stats.append({
            "segment": segment,
            "splice": int(round(splice)),
            "connection": int(round(connection)),
            "missing": int(round(missing)),
            "rate": bolt_missing_rate(splice, connection, missing),
            "points": len(rows),
        })
    return stats


def make_stats(segments, records):
    stats = []
    for segment_index, segment in enumerate(segments):
        types = {}
        for kind in ("二波", "三波"):
            heights = [
                record["height"] for record in records
                if record["segment"] == segment_index and record["kind"] == kind
            ]
            bins = [0] * 5
            for height in heights:
                bins[bin_index(kind, height)] += 1
            percentages = [round(value * 100 / len(heights), 2) if heights else 0 for value in bins]
            types[kind] = {
                "count": len(heights), "bins": bins,
                "pcts": percentages, "pass": percentages[2] if heights else 0,
            }
        stats.append({"segment": segment, "types": types})
    return stats


def style_sheet(ws, widths):
    thin = Side(style="thin", color="808080")
    for row in ws.iter_rows():
        for cell in row:
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
    for index, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(index)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def add_interval_sheets(out_path, summary_path, log=lambda _: None):
    segments = read_segments(summary_path)
    wb = openpyxl.load_workbook(out_path)
    for name in list(wb.sheetnames):
        if name.startswith("区间"):
            del wb[name]
    detail = wb["检测明细"]
    headers = {cell.value: cell.column for cell in detail[1]}
    grouped = {(format_station(s["start"]), format_station(s["end"])): [] for s in segments}
    for row in detail.iter_rows(min_row=2, values_only=True):
        key = (row[headers["所属分段起点"] - 1], row[headers["所属分段终点"] - 1])
        if key in grouped:
            normalized_kind = guardrail_type(row[headers["护栏类型"] - 1]) or row[headers["护栏类型"] - 1]
            grouped[key].append((
                row[headers["统计桩号"] - 1], normalized_kind,
                row[headers["梁板中心高度(mm)"] - 1],
            ))
    columns = [
        "序号", "桩号", "护栏类型", "梁板中心高度(mm)",
        "标准值（580mm）", "标准值（620mm）", "标准值（677mm）", "标准值（717mm）",
    ]
    thin = Side(style="thin", color="808080")
    for index, (start, end) in enumerate(grouped, 1):
        ws = wb.create_sheet(f"区间{index:02d}_{start[1:]}-{end[1:]}")
        ws.append(columns)
        rows = sorted(grouped[(start, end)], key=lambda item: (item[0], item[1]))
        for sequence, (station, kind, height) in enumerate(rows, 1):
            ws.append([sequence, station, kind, height, 580, 620, 677, 717])
        for row in ws.iter_rows():
            for cell in row:
                cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4E78")
        for col, width in enumerate([9, 18, 14, 23, 19, 19, 19, 19], 1):
            ws.column_dimensions[get_column_letter(col)].width = width
        ws.freeze_panes, ws.sheet_view.showGridLines = "A2", False
        ws["J1"], ws["K1"], ws["J2"], ws["K2"] = "统计区间", f"{start}～{end}", "有效记录数", len(rows)
        ws["J1"].font = ws["J2"].font = Font(bold=True)
        if rows:
            table = Table(displayName=f"IntervalDetail{index:02d}", ref=f"A1:H{len(rows)+1}")
            table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True)
            ws.add_table(table)
    wb.save(out_path)
    log(f"区间明细表已更新：{len(segments)}个区间，共{sum(map(len, grouped.values()))}条记录。")


def make_excel(
    config, segments,
    height_stats=None, height_records=None, height_duplicates=0, excluded=None,
    bolt_stats=None, bolt_records=None, bolt_duplicates=0,
    log=lambda _: None,
):
    height_records = height_records or []
    bolt_records = bolt_records or []
    excluded = excluded or Counter()
    wb = Workbook()
    wb.remove(wb.active)
    if height_stats is not None:
        for title, kind, labels in (
            ("二波统计", "二波", ["h＜560", "560≤h＜580", "580≤h≤620", "620＜h≤640", "h＞640"]),
            ("三波统计", "三波", ["h＜657", "657≤h＜677", "677≤h≤717", "717＜h≤737", "h＞737"]),
        ):
            ws = wb.create_sheet(title)
            ws.append(["序号", "路线编号", "公路等级", "管理单位", "起点桩号", "终点桩号", "统计口径", "检测点数", *labels, "合格率"])
            sequence = 0
            for segment_index, item in enumerate(height_stats):
                data = item["types"][kind]
                if not data["count"]:
                    continue
                sequence += 1
                segment = item["segment"]
                bases = {r["basis"] for r in height_records if r["segment"] == segment_index and r["kind"] == kind}
                basis = next(iter(bases)) if len(bases) == 1 else "原始桩号+电子修正桩号（按来源文件分别采用）"
                ws.append([
                    sequence, "G210", segment["grade"], segment["manager"],
                    format_station(segment["start"]), format_station(segment["end"]),
                    basis, data["count"], *[value / 100 for value in data["pcts"]], data["pass"] / 100,
                ])
            for row in ws.iter_rows(min_row=2, min_col=9, max_col=14):
                for cell in row:
                    cell.number_format = "0.00%"
            style_sheet(ws, [7, 10, 14, 24, 14, 14, 42, 11, 13, 14, 14, 14, 13, 12])
        ws = wb.create_sheet("检测明细")
        ws.append(["序号", "来源文件", "方向", "统计桩号", "桩号口径", "护栏类型", "梁板中心高度(mm)", "所属分段起点", "所属分段终点"])
        ordered = sorted(height_records, key=lambda item: (item["segment"], item["direction"], item["station"]))
        for sequence, record in enumerate(ordered, 1):
            segment = segments[record["segment"]]
            ws.append([
                sequence, record["file"], record["direction"], format_station(record["station"]),
                record["basis"], record["kind"], record["height"],
                format_station(segment["start"]), format_station(segment["end"]),
            ])
        style_sheet(ws, [8, 62, 10, 16, 14, 12, 20, 16, 16])

    if bolt_stats is not None:
        ws = wb.create_sheet("螺栓缺失统计")
        ws.append(["序号", "路线编号", "起点桩号", "止点桩号", "检测里程（km）", "拼接螺栓（颗）", "连接螺栓（颗）", "缺失数量（颗）", "缺失率（%）"])
        sequence = 0
        for item in bolt_stats:
            sequence += 1
            segment = item["segment"]
            ws.append([
                sequence, "G210", format_station(segment["start"]), format_station(segment["end"]),
                segment["mileage"], item["splice"], item["connection"], item["missing"], item["rate"] / 100,
            ])
        for cell in ws["I"][1:]:
            cell.number_format = "0.00%"
        style_sheet(ws, [8, 11, 16, 16, 16, 18, 18, 18, 14])

        ws = wb.create_sheet("螺栓缺失明细")
        ws.append([
            "序号", "来源文件", "方向", "统计桩号", "桩号口径",
            "拼接螺栓（颗）", "拼接螺栓缺失（颗）", "连接螺栓（颗）", "连接螺栓缺失（颗）",
            "所属分段起点", "所属分段终点",
        ])
        ordered = sorted(bolt_records, key=lambda item: (item["segment"], item["direction"], item["station"]))
        for sequence, record in enumerate(ordered, 1):
            segment = segments[record["segment"]]
            ws.append([
                sequence, record["file"], record["direction"], format_station(record["station"]), record["basis"],
                record["splice"], record["splice_missing"], record["connection"], record["connection_missing"],
                format_station(segment["start"]), format_station(segment["end"]),
            ])
        style_sheet(ws, [8, 62, 10, 16, 14, 18, 22, 18, 22, 16, 16])

    ws = wb.create_sheet("统计说明")
    notes = [
        ("项目", "说明"),
        ("统计分段", "按G210采集路段信息汇总表的起终点桩号分段。"),
        ("桩号口径", "G210上行K2264K2325文件使用原始桩号；其他文件使用电子修正桩号。"),
    ]
    if height_stats is not None:
        notes.extend([
            ("中心高度统计数值", "使用护栏高度表中的梁板中心高度(mm)列。"),
            ("中心高度备注过滤", f"仅保留异常标记为空或无备注的数据；排除{sum(excluded.values())}条：{dict(excluded)}。"),
            ("中心高度去重", f"完全一致的重叠记录去重，剔除{height_duplicates}条。"),
            ("区间明细标准值", "标准值列为580mm、620mm、677mm和717mm；二波折线图使用580mm、620mm，三波折线图使用677mm、717mm。"),
        ])
    if bolt_stats is not None:
        notes.extend([
            ("螺栓统计数值", "使用拼接螺栓数量、拼接螺栓缺失数量、连接螺栓数量和连接螺栓缺失数量列。"),
            ("螺栓缺失总数", "缺失数量为拼接螺栓缺失数量与连接螺栓缺失数量之和。"),
            ("螺栓缺失率", "缺失数量÷（拼接螺栓数量+连接螺栓数量+缺失数量）×100%。"),
            ("螺栓去重", f"完全一致的重叠记录去重，剔除{bolt_duplicates}条。"),
        ])
    for row in notes:
        ws.append(row)
    style_sheet(ws, [18, 110])
    config.output_dir.mkdir(parents=True, exist_ok=True)
    wb.save(config.out_xlsx)
    log(f"统计工作簿已生成：{config.out_xlsx}")
    if height_stats is not None:
        add_interval_sheets(config.out_xlsx, config.summary_xlsx, log)


def run_properties(bold=False, size=21):
    properties = ET.Element(wt("rPr"))
    fonts = ET.SubElement(properties, wt("rFonts"))
    fonts.set(wt("ascii"), "Times New Roman")
    fonts.set(wt("hAnsi"), "Times New Roman")
    fonts.set(wt("eastAsia"), "宋体")
    fonts.set(wt("cs"), "Times New Roman")
    if bold:
        ET.SubElement(properties, wt("b")); ET.SubElement(properties, wt("bCs"))
    if size is not None:
        ET.SubElement(properties, wt("sz")).set(wt("val"), str(size))
        ET.SubElement(properties, wt("szCs")).set(wt("val"), str(size))
    return properties


def paragraph_properties(style=None, center=False, num_id=None, level=None, keep_next=None):
    properties = ET.Element(wt("pPr"))
    if style:
        ET.SubElement(properties, wt("pStyle")).set(wt("val"), str(style))
    if keep_next is not None:
        ET.SubElement(properties, wt("keepNext")).set(wt("val"), "1" if keep_next else "0")
    if num_id is not None:
        number = ET.SubElement(properties, wt("numPr"))
        ET.SubElement(number, wt("ilvl")).set(wt("val"), str(level))
        ET.SubElement(number, wt("numId")).set(wt("val"), str(num_id))
    indent = ET.SubElement(properties, wt("ind"))
    indent.set(wt("left"), "0"); indent.set(wt("right"), "0"); indent.set(wt("firstLine"), "0")
    ET.SubElement(properties, wt("jc")).set(wt("val"), "center" if center else "left")
    return properties


def append_run(p, text, bold=False, size=21):
    run = ET.SubElement(p, wt("r")); run.append(run_properties(bold, size))
    node = ET.SubElement(run, wt("t")); node.text = str(text)
    if str(text).startswith(" ") or str(text).endswith(" "):
        node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return run


def body_paragraph(text="", center=False, bold=False, style=None, keep_next=None):
    p = ET.Element(wt("p")); p.append(paragraph_properties(style, center, keep_next=keep_next))
    append_run(p, text, bold)
    return p


def heading_paragraph(text, level, num_id):
    style = {1: "2", 2: "3", 3: "4"}[level]
    p = ET.Element(wt("p")); p.append(paragraph_properties(style, False, num_id, level - 1))
    append_run(p, text, True, None)
    return p


def append_field(p, instruction, result, bold=False):
    begin = ET.SubElement(p, wt("r")); begin.append(run_properties(bold)); ET.SubElement(begin, wt("fldChar")).set(wt("fldCharType"), "begin")
    code = ET.SubElement(p, wt("r")); code.append(run_properties(bold)); text = ET.SubElement(code, wt("instrText")); text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve"); text.text = f" {instruction} "
    separate = ET.SubElement(p, wt("r")); separate.append(run_properties(bold)); ET.SubElement(separate, wt("fldChar")).set(wt("fldCharType"), "separate")
    append_run(p, result, bold)
    end = ET.SubElement(p, wt("r")); end.append(run_properties(bold)); ET.SubElement(end, wt("fldChar")).set(wt("fldCharType"), "end")


def caption_paragraph(
    title, label, section_number, sequence,
    bookmark_id=None, bookmark_name=None,
):
    """创建包含标题2章节号的图/表题注，编号格式为“章节号-序号”。"""
    p = ET.Element(wt("p")); p.append(paragraph_properties("9", True, keep_next=False))
    if bookmark_id is not None and bookmark_name:
        start = ET.SubElement(p, wt("bookmarkStart"))
        start.set(wt("id"), str(bookmark_id)); start.set(wt("name"), bookmark_name)
    append_run(p, f"{label} ", True)
    # Word“题注编号-包含章节号”的字段结构：章节起始样式为标题2，
    # STYLEREF取得当前标题2编号，SEQ按标题2重新开始计数。
    append_field(p, 'STYLEREF 2 \\s', str(section_number), True)
    append_run(p, "-", True)
    append_field(p, f'SEQ {label} \\* ARABIC \\s 2', str(sequence), True)
    if bookmark_id is not None and bookmark_name:
        end = ET.SubElement(p, wt("bookmarkEnd")); end.set(wt("id"), str(bookmark_id))
    append_run(p, "  " + title, True)
    return p


def append_reference(p, bookmark_name, visible_text):
    append_field(p, f"REF {bookmark_name} \\h", visible_text, False)


def word_table(headers, rows, header_shading=False, keep_together=False):
    table = ET.Element(wt("tbl")); properties = ET.SubElement(table, wt("tblPr"))
    borders = ET.SubElement(properties, wt("tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = ET.SubElement(borders, wt(side)); node.set(wt("val"), "single"); node.set(wt("sz"), "4")
    for row_index, row in enumerate([headers] + rows):
        tr = ET.SubElement(table, wt("tr"))
        if keep_together:
            tr_properties = ET.SubElement(tr, wt("trPr")); ET.SubElement(tr_properties, wt("cantSplit"))
        for value in row:
            tc = ET.SubElement(tr, wt("tc")); tc_properties = ET.SubElement(tc, wt("tcPr"))
            if row_index == 0 and header_shading:
                shading = ET.SubElement(tc_properties, wt("shd"))
                shading.set(wt("val"), "pct15")
                shading.set(wt("color"), "auto")
                shading.set(wt("fill"), "auto")
            p = ET.SubElement(tc, wt("p"))
            # 表格内容允许自然分页，不设置“与下段同页”(w:keepNext)。
            p.append(paragraph_properties(None, True, keep_next=False))
            append_run(p, value, row_index == 0)
    return table


def height_example_table(records):
    """创建2列表格：奇数行留空放图，偶数行合并后显示点位信息。"""
    table = ET.Element(wt("tbl")); properties = ET.SubElement(table, wt("tblPr"))
    width = ET.SubElement(properties, wt("tblW")); width.set(wt("w"), "9000"); width.set(wt("type"), "dxa")
    borders = ET.SubElement(properties, wt("tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = ET.SubElement(borders, wt(side)); node.set(wt("val"), "single"); node.set(wt("sz"), "4"); node.set(wt("color"), "808080")
    grid = ET.SubElement(table, wt("tblGrid"))
    for _ in range(2):
        ET.SubElement(grid, wt("gridCol")).set(wt("w"), "4500")

    for record in records:
        image_row = ET.SubElement(table, wt("tr")); row_properties = ET.SubElement(image_row, wt("trPr"))
        height = ET.SubElement(row_properties, wt("trHeight")); height.set(wt("val"), "2500"); height.set(wt("hRule"), "atLeast")
        for _ in range(2):
            cell = ET.SubElement(image_row, wt("tc")); cell_properties = ET.SubElement(cell, wt("tcPr"))
            cell_width = ET.SubElement(cell_properties, wt("tcW")); cell_width.set(wt("w"), "4500"); cell_width.set(wt("type"), "dxa")
            ET.SubElement(cell_properties, wt("vAlign")).set(wt("val"), "center")
            paragraph = ET.SubElement(cell, wt("p")); paragraph.append(paragraph_properties(None, True))

        info_row = ET.SubElement(table, wt("tr")); cell = ET.SubElement(info_row, wt("tc")); cell_properties = ET.SubElement(cell, wt("tcPr"))
        cell_width = ET.SubElement(cell_properties, wt("tcW")); cell_width.set(wt("w"), "9000"); cell_width.set(wt("type"), "dxa")
        ET.SubElement(cell_properties, wt("gridSpan")).set(wt("val"), "2")
        ET.SubElement(cell_properties, wt("vAlign")).set(wt("val"), "center")
        paragraph = ET.SubElement(cell, wt("p")); paragraph.append(paragraph_properties(None, True))
        electronic = record.get("electronic_station")
        raw = record.get("raw_station")
        electronic_text = format_station(electronic) if electronic is not None else "—"
        raw_text = format_station(raw) if raw is not None else "—"
        append_run(paragraph, f"{electronic_text}（{record['height']:.2f}mm） {raw_text}")
    return table


def replace_text(root, old, new):
    for node in root.findall(f".//{wt('t')}"):
        if node.text:
            node.text = node.text.replace(old, new)


def add_multilevel_numbering(numbering_xml):
    root = ET.fromstring(numbering_xml)
    abstracts = [int(node.get(wt("abstractNumId"))) for node in root.findall(wt("abstractNum"))]
    nums = [int(node.get(wt("numId"))) for node in root.findall(wt("num"))]
    abstract_id = max(abstracts, default=0) + 1
    abstract = ET.Element(wt("abstractNum")); abstract.set(wt("abstractNumId"), str(abstract_id))
    ET.SubElement(abstract, wt("multiLevelType")).set(wt("val"), "multilevel")
    for level, (style, text) in enumerate((("2", "%1"), ("3", "%1.%2"), ("4", "%1.%2.%3"))):
        item = ET.SubElement(abstract, wt("lvl")); item.set(wt("ilvl"), str(level))
        ET.SubElement(item, wt("start")).set(wt("val"), "1")
        ET.SubElement(item, wt("numFmt")).set(wt("val"), "decimal")
        # 不写lvlRestart：Word默认在出现上一级标题时重启下一层编号。
        # 显式写入与当前层级冲突的值会导致Word修复并删除列表定义。
        ET.SubElement(item, wt("pStyle")).set(wt("val"), style)
        # 编号后使用一个空格，不使用Word默认制表符；标题保持无列表缩进。
        ET.SubElement(item, wt("suff")).set(wt("val"), "space")
        ET.SubElement(item, wt("lvlText")).set(wt("val"), text)
        ET.SubElement(item, wt("lvlJc")).set(wt("val"), "left")
        ppr = ET.SubElement(item, wt("pPr")); ind = ET.SubElement(ppr, wt("ind")); ind.set(wt("left"), "0"); ind.set(wt("hanging"), "0")
    # numbering.xml要求所有abstractNum位于num之前。
    first_num_index = next((i for i, child in enumerate(list(root)) if child.tag == wt("num")), len(root))
    root.insert(first_num_index, abstract)
    chapter_nums = {}
    next_id = max(nums, default=0) + 1
    for chapter in (1, 3, 4, 5):
        num = ET.SubElement(root, wt("num")); num.set(wt("numId"), str(next_id))
        ET.SubElement(num, wt("abstractNumId")).set(wt("val"), str(abstract_id))
        override = ET.SubElement(num, wt("lvlOverride")); override.set(wt("ilvl"), "0")
        ET.SubElement(override, wt("startOverride")).set(wt("val"), str(chapter))
        # 每章使用独立编号实例，并显式将小节、子小节从1开始。
        for level in (1, 2):
            override = ET.SubElement(num, wt("lvlOverride")); override.set(wt("ilvl"), str(level))
            ET.SubElement(override, wt("startOverride")).set(wt("val"), "1")
        chapter_nums[chapter] = next_id; next_id += 1
    return ET.tostring(root, encoding="utf-8", xml_declaration=True), chapter_nums


def percentage_phrases(kind, percentages):
    limits = (560, 580, 620, 640) if kind == "二波" else (657, 677, 717, 737)
    labels = [
        f"小于{limits[0]}mm的约占{{:.2f}}%",
        f"介于{limits[0]}～{limits[1]}mm的约占{{:.2f}}%",
        f"介于{limits[1]}～{limits[2]}mm的约占{{:.2f}}%",
        f"介于{limits[2]}～{limits[3]}mm的约占{{:.2f}}%",
        f"大于{limits[3]}mm的约占{{:.2f}}%",
    ]
    return "，".join(labels[index].format(value) for index, value in enumerate(percentages) if value > 0)


def order_example_records(records):
    """示例图片顺序：上行桩号递增，下行桩号递减；混合时上行组在前。"""
    def key(record):
        direction = str(record.get("direction") or "")
        station = record.get("station")
        station = float(station) if station is not None else 0.0
        if "上" in direction:
            return 0, station
        if "下" in direction:
            return 1, -station
        return 2, station
    return sorted(records, key=key)


def select_height_example_points(rows, segment_index=0, kind=""):
    """按检测点数量和高度分位选择自动计算示例点，返回顺序固定且可复现。"""
    ordered = sorted(rows, key=lambda item: (item["height"], item["station"]))
    count = len(ordered)
    if not count:
        return []
    if count > 1000:
        selected = []
        for index in range(4):
            start = index * count // 4
            end = (index + 1) * count // 4
            group = ordered[start:end]
            selected.append(group[len(group) // 2])
        return order_example_records(selected)
    if count > 100:
        quarter_count = max(1, math.ceil(count * 0.25))
        randomizer = random.Random(f"G210|{segment_index}|{kind}|{count}")
        return order_example_records([
            randomizer.choice(ordered[:quarter_count]),
            randomizer.choice(ordered[-quarter_count:]),
        ])
    # 偶数个点时采用靠前的中位点，确保选中的是实际存在的数据点。
    return order_example_records([ordered[(count - 1) // 2]])


def select_bolt_example_points(rows, disease_image_index=None):
    """按缺失阈值和病害图片可匹配性选择拼接、连接螺栓示例。"""
    def attach_image(row):
        descriptor = match_disease_image(row, disease_image_index) if disease_image_index is not None else None
        return row, descriptor

    candidates = [attach_image(row) for row in rows]
    if disease_image_index is not None:
        candidates = [(row, image) for row, image in candidates if image is not None]
    splice_only_rows = [
        (row, image) for row, image in candidates
        if 0 < row["splice_missing"] <= 12 and row["connection_missing"] <= 0
    ]
    connection_rows = [
        (row, image) for row, image in candidates
        if 0 < row["connection_missing"] <= 2 and row["splice_missing"] <= 12
    ]
    splice_only_rows.sort(key=lambda item: (-item[0]["splice_missing"], item[0]["station"]))
    connection_rows.sort(key=lambda item: (
        -item[0]["connection_missing"],
        -(item[0]["splice_missing"] + item[0]["connection_missing"]),
        item[0]["station"],
    ))

    selected_rows = []
    if splice_only_rows:
        selected_rows.append(splice_only_rows[0])
    if connection_rows:
        selected_rows.append(connection_rows[0])
    elif selected_rows:
        # 没有任何连接螺栓缺失点时，只保留一个仅拼接螺栓缺失点。
        selected_rows = selected_rows[:1]

    selected = [
        {
            "record": row,
            "missing": int(round(row["splice_missing"] + row["connection_missing"])),
            "image": image,
        }
        for row, image in selected_rows
    ]
    ordered_records = order_example_records([example["record"] for example in selected])
    by_identity = {id(example["record"]): example for example in selected}
    return [by_identity[id(record)] for record in ordered_records]


def bolt_example_text(example):
    raw = example["record"].get("raw_station")
    raw_text = format_station_one_decimal(raw) if raw is not None else "—"
    return f"螺栓缺失{example['missing']}颗 {raw_text}"


def report_images(temp_dir, segments, stats, records):
    images = {}
    plt.rcParams["font.sans-serif"] = ["SimSun", "Microsoft YaHei", "Arial Unicode MS"]
    plt.rcParams["axes.unicode_minus"] = False
    for segment_index, item in enumerate(stats):
        segment = item["segment"]
        for kind in ("二波", "三波"):
            rows = sorted(
                [record for record in records if record["segment"] == segment_index and record["kind"] == kind],
                key=lambda record: record["station"],
            )
            if not rows:
                continue
            key = (segment_index, kind)
            line_path = Path(temp_dir) / f"line_{segment_index}_{kind}.png"
            pie_path = Path(temp_dir) / f"pie_{segment_index}_{kind}.png"
            x = list(range(len(rows))); heights = [row["height"] for row in rows]
            figure, axis = plt.subplots(figsize=(13 / 2.54, 8 / 2.54), dpi=180)
            axis.plot(x, heights, color="#4472C4", linewidth=1, label="梁板中心高度(mm)")
            standards = (580, 620) if kind == "二波" else (677, 717)
            for standard, color in zip(standards, ("#ED7D31", "#A5A5A5")):
                axis.plot(
                    x, [standard] * len(x), color=color, linewidth=2,
                    label=f"标准值（{standard}mm）",
                )
            axis.set_ylim(300, 850); axis.grid(True, alpha=0.25)
            ticks = sorted(set(int(i * (len(x) - 1) / min(9, max(1, len(x) - 1))) for i in range(min(10, len(x)))))
            axis.set_xticks(ticks); axis.set_xticklabels([format_station(rows[i]["station"]) for i in ticks], rotation=30, ha="right", fontsize=7)
            axis.legend(loc="upper center", bbox_to_anchor=(0.5, -0.22), ncol=3, frameon=False)
            figure.subplots_adjust(left=0.10, right=0.98, top=0.96, bottom=0.30)
            figure.savefig(line_path, transparent=False); plt.close(figure)

            values = item["types"][kind]["pcts"]
            labels = (["h＜560", "560≤h＜580", "580≤h≤620", "620＜h≤640", "h＞640"] if kind == "二波" else
                      ["h＜657", "657≤h＜677", "677≤h≤717", "717＜h≤737", "h＞737"])
            nonzero = [(label, value, f"#{PIE_COLORS[i]}") for i, (label, value) in enumerate(zip(labels, values)) if value > 0]
            figure, axis = plt.subplots(figsize=(14 / 2.54, 8.5 / 2.54), dpi=180)
            wedges, _, _ = axis.pie(
                [value for _, value, _ in nonzero], colors=[color for _, _, color in nonzero],
                autopct=lambda pct: f"{pct:.2f}%" if pct > 0 else "", pctdistance=1.15,
                textprops={"fontsize": 8},
            )
            axis.legend(wedges, [label for label, _, _ in nonzero], loc="center left", bbox_to_anchor=(1.0, 0.5), frameon=False)
            figure.subplots_adjust(left=0.02, right=0.76, top=0.88, bottom=0.05)
            figure.savefig(pie_path, transparent=False); plt.close(figure)
            images[key] = {"line": line_path, "pie": pie_path}
    return images


def picture_paragraph(rel_id, drawing_id, width_cm=13, height_cm=8):
    cx, cy = int(width_cm * 360000), int(height_cm * 360000)
    p = ET.Element(wt("p")); p.append(paragraph_properties(None, True))
    run = ET.SubElement(p, wt("r")); run.append(run_properties())
    drawing = ET.SubElement(run, wt("drawing")); inline = ET.SubElement(drawing, q(WP, "inline"))
    extent = ET.SubElement(inline, q(WP, "extent")); extent.set("cx", str(cx)); extent.set("cy", str(cy))
    effect = ET.SubElement(inline, q(WP, "effectExtent"));
    for side in ("l", "t", "r", "b"): effect.set(side, "0")
    docpr = ET.SubElement(inline, q(WP, "docPr")); docpr.set("id", str(drawing_id)); docpr.set("name", f"G210Chart{drawing_id}")
    frame = ET.SubElement(inline, q(WP, "cNvGraphicFramePr")); ET.SubElement(frame, q(A, "graphicFrameLocks")).set("noChangeAspect", "1")
    graphic = ET.SubElement(inline, q(A, "graphic")); data = ET.SubElement(graphic, q(A, "graphicData")); data.set("uri", "http://schemas.openxmlformats.org/drawingml/2006/picture")
    pic = ET.SubElement(data, q(PIC, "pic")); nv = ET.SubElement(pic, q(PIC, "nvPicPr"));
    cpr = ET.SubElement(nv, q(PIC, "cNvPr")); cpr.set("id", "0"); cpr.set("name", f"G210Image{drawing_id}")
    ET.SubElement(nv, q(PIC, "cNvPicPr")); fill = ET.SubElement(pic, q(PIC, "blipFill"))
    blip = ET.SubElement(fill, q(A, "blip")); blip.set(q(R, "embed"), rel_id)
    stretch = ET.SubElement(fill, q(A, "stretch")); ET.SubElement(stretch, q(A, "fillRect"))
    shape = ET.SubElement(pic, q(PIC, "spPr")); transform = ET.SubElement(shape, q(A, "xfrm"))
    offset = ET.SubElement(transform, q(A, "off")); offset.set("x", "0"); offset.set("y", "0")
    size = ET.SubElement(transform, q(A, "ext")); size.set("cx", str(cx)); size.set("cy", str(cy))
    geometry = ET.SubElement(shape, q(A, "prstGeom")); geometry.set("prst", "rect"); ET.SubElement(geometry, q(A, "avLst"))
    return p


def make_docx(
    config, segments,
    height_stats=None, height_records=None,
    bolt_stats=None, bolt_records=None,
    disease_image_index=None,
    log=lambda _: None,
    require_template=True,
):
    if config.template_docx.suffix.lower() != ".md":
        raise ValueError(f"仅支持 Markdown 模板：{config.template_docx}，请使用 .md 模板。")
    if not config.template_docx.is_file():
        raise FileNotFoundError(f"Word模板不存在：{config.template_docx}")
    from backend import minimal_docx
    log("使用 Markdown 报告模板。")
    return minimal_docx.run(
        config, segments,
        height_stats=height_stats, height_records=height_records,
        bolt_stats=bolt_stats, bolt_records=bolt_records,
        disease_image_index=disease_image_index, log=log,
        skeleton_md=config.template_docx,
    )

def add_distribution_charts(wb, source_name, target_name):
    if target_name in wb.sheetnames:
        del wb[target_name]
    source, target = wb[source_name], wb.create_sheet(target_name)
    target.sheet_view.showGridLines = False
    display_name = source_name.replace("双波", "二波").replace("两波", "二波")
    target["A1"] = f"{display_name}各区段波形梁护栏横梁中心高度分布情况图"
    target["A1"].font = Font(size=14, bold=True)
    count = 0
    for row in range(2, source.max_row + 1):
        start, end = source.cell(row, 5).value, source.cell(row, 6).value
        if not start or not end:
            continue
        chart = PieChart()
        chart.add_data(Reference(source, min_col=9, max_col=13, min_row=row, max_row=row), from_rows=True)
        chart.set_categories(Reference(source, min_col=9, max_col=13, min_row=1, max_row=1))
        chart.title, chart.legend.position, chart.roundedCorners = f"{start}-{end}", "r", False
        chart.title.tx.rich.p[0].r[0].rPr = CharacterProperties(
            latin=DrawingFont(typeface="Times New Roman"),
            ea=DrawingFont(typeface="宋体"),
            sz=1050,
            b=True,
        )
        chart.height, chart.width, chart.varyColors = 8.5, 14, True
        chart.layout = Layout(manualLayout=ManualLayout(x=-2 / 14, xMode="factor"))
        chart.dataLabels = DataLabelList()
        chart.dataLabels.showVal, chart.dataLabels.showPercent = True, False
        chart.dataLabels.showCatName = chart.dataLabels.showSerName = chart.dataLabels.showLegendKey = False
        chart.dataLabels.dLblPos, chart.dataLabels.numFmt = "outEnd", "0.00%"
        chart.dataLabels.showLeaderLines = True
        # 对占比为0的扇区单独设置删除标签，避免显示0.00%。
        zero_labels = []
        for index in range(5):
            if not (source.cell(row, 9 + index).value or 0):
                label = DataLabel(idx=index)
                label.delete = True
                zero_labels.append(label)
        chart.dataLabels.dLbl = zero_labels
        chart.series[0].data_points = [DataPoint(idx=i, spPr=GraphicalProperties(solidFill=color)) for i, color in enumerate(PIE_COLORS)]
        grid_row, grid_col = divmod(count, 2)
        target.add_chart(chart, f"{'A' if grid_col == 0 else 'J'}{3 + grid_row * 17}")
        count += 1
    return count


def configure_line_chart(chart):
    chart.title, chart.style, chart.roundedCorners = None, 2, False
    chart.height, chart.width = 8, 13
    chart.y_axis.scaling.min, chart.y_axis.scaling.max, chart.y_axis.majorUnit = 300, 850, 50
    chart.x_axis.title = chart.y_axis.title = None
    chart.x_axis.delete = chart.y_axis.delete = False
    # 水平轴与图例间距约0.1厘米。
    chart.layout = Layout(manualLayout=ManualLayout(
        x=0.10, y=0.05, w=0.80, h=0.8125,
        xMode="edge", yMode="edge", wMode="factor", hMode="factor",
    ))
    chart.legend.position = "b"
    chart.legend.layout = Layout(manualLayout=ManualLayout(
        x=0.10, y=0.875, w=0.80, h=0.10,
        xMode="edge", yMode="edge", wMode="factor", hMode="factor",
    ))
    for index, series in enumerate(chart.series):
        series.graphicalProperties.line.solidFill = LINE_COLORS[index]
        series.graphicalProperties.line.width = 12700 if index == 0 else 25400
        series.graphicalProperties.line.noFill = False


def add_interval_line_charts(wb):
    helper_name = "_折线图数据"
    if helper_name in wb.sheetnames:
        del wb[helper_name]
    helper = wb.create_sheet(helper_name); helper.sheet_state = "hidden"
    count = helper_group = 0
    for sheet in [ws for ws in wb.worksheets if ws.title.startswith("区间")]:
        sheet._charts = []
        data_rows = [row for row in range(2, sheet.max_row + 1) if sheet.cell(row, 1).value is not None]
        if not data_rows:
            sheet["H4"] = "本区间无有效护栏高度数据，未绘制折线图。"
            continue
        sheet["H4"] = None
        groups = {"二波": [], "三波": []}
        for row in data_rows:
            kind = str(sheet.cell(row, 3).value or "")
            if "三" in kind:
                groups["三波"].append(row)
            elif any(word in kind for word in ("双", "两", "二")):
                groups["二波"].append(row)
        sheet_chart_index = 0
        for kind in ("二波", "三波"):
            rows = groups[kind]
            if not rows:
                continue
            start_col = 1 + helper_group * 5; helper_group += 1
            standard_columns = (5, 6) if kind == "二波" else (7, 8)
            standard_values = (580, 620) if kind == "二波" else (677, 717)
            helper_headers = [
                "桩号", "梁板中心高度(mm)",
                f"标准值（{standard_values[0]}mm）", f"标准值（{standard_values[1]}mm）",
            ]
            for offset, header in enumerate(helper_headers):
                helper.cell(1, start_col + offset, header)
            for target_row, source_row in enumerate(rows, 2):
                helper.cell(target_row, start_col, sheet.cell(source_row, 2).value)
                helper.cell(target_row, start_col + 1, sheet.cell(source_row, 4).value)
                helper.cell(target_row, start_col + 2, sheet.cell(source_row, standard_columns[0]).value)
                helper.cell(target_row, start_col + 3, sheet.cell(source_row, standard_columns[1]).value)
            chart = LineChart(); chart.visible_cells_only = False
            chart.add_data(Reference(helper, min_col=start_col + 1, max_col=start_col + 3, min_row=1, max_row=len(rows) + 1), titles_from_data=True)
            chart.set_categories(Reference(helper, min_col=start_col, min_row=2, max_row=len(rows) + 1))
            configure_line_chart(chart)
            sheet.add_chart(chart, f"K{4 + sheet_chart_index * 17}")
            sheet_chart_index += 1; count += 1
    return count


def add_charts(workbook_path, log=lambda _: None):
    wb = openpyxl.load_workbook(workbook_path)
    if "双波分布图" in wb.sheetnames:
        del wb["双波分布图"]
    if "二波统计" not in wb.sheetnames and "双波统计" in wb.sheetnames:
        wb["双波统计"].title = "二波统计"
    two_wave_source = "二波统计"
    double_pies = add_distribution_charts(wb, two_wave_source, "二波分布图")
    triple_pies = add_distribution_charts(wb, "三波统计", "三波分布图")
    line_charts = add_interval_line_charts(wb)
    for sheet in wb.worksheets:
        for chart in sheet._charts:
            chart.roundedCorners = False
    wb.save(workbook_path)
    log(f"图表已生成：二波饼图{double_pies}个，三波饼图{triple_pies}个，分类型折线图{line_charts}个。")


def generate_statistics_and_report(
    config, log=lambda _: None, generate_charts_first=False,
    process_height=True, process_bolts=False, process_alongline=False,
    require_template=True,
):
    segments = read_segments(config.summary_xlsx)
    height_records = []; height_stats = None; height_duplicates = 0; excluded = Counter()
    bolt_records = []; bolt_stats = None; bolt_duplicates = 0
    disease_image_index = None
    if process_alongline:
        log("沿线设施选项已勾选；该细分项暂未开发，本次不会生成沿线设施统计或报告内容。")
    if process_height:
        height_records, height_duplicates, excluded = collect_records(segments, config.detail_dir, log)
        height_stats = make_stats(segments, height_records)
    if process_bolts:
        bolt_records, bolt_duplicates = collect_bolt_records(segments, config.detail_dir, log)
        bolt_stats = make_bolt_stats(segments, bolt_records)
        if config.disease_dir is None or not Path(config.disease_dir).is_dir():
            raise FileNotFoundError("处理螺栓缺失时，请选择有效的病害清单文件夹。")
        disease_image_index = collect_disease_image_index(config.disease_dir, log)
    make_excel(
        config, segments,
        height_stats=height_stats, height_records=height_records,
        height_duplicates=height_duplicates, excluded=excluded,
        bolt_stats=bolt_stats, bolt_records=bolt_records, bolt_duplicates=bolt_duplicates,
        log=log,
    )
    if generate_charts_first and process_height:
        add_charts(config.out_xlsx, log)
    make_docx(
        config, segments,
        height_stats=height_stats, height_records=height_records,
        bolt_stats=bolt_stats, bolt_records=bolt_records,
        disease_image_index=disease_image_index, log=log,
        require_template=require_template,
    )
    if process_height:
        log(f"中心高度有效记录{len(height_records)}条；备注排除{sum(excluded.values())}条；重复排除{height_duplicates}条。")
    if process_bolts:
        log(f"螺栓有效记录{len(bolt_records)}条；重复排除{bolt_duplicates}条。")
    return {"height": height_stats, "bolts": bolt_stats}


def discover_paths(folder):
    folder = Path(folder)
    summary = next(iter(folder.rglob("G210采集路段信息汇总.xlsx")), None)
    detail_candidates = []
    disease_candidates = []
    for directory in [folder, *[p for p in folder.rglob("*") if p.is_dir()]]:
        detail_count = sum(1 for path in directory.glob("*.xlsx") if "交安设施现场检测" in path.name and "明细" in path.name)
        disease_count = sum(1 for path in directory.glob("*.xlsx") if "交安设施现场检测" in path.name and "病害清单" in path.name)
        if detail_count:
            detail_candidates.append((detail_count, directory))
        if disease_count:
            disease_candidates.append((disease_count, directory))
    detail = max(detail_candidates, default=(0, None))[1]
    disease = max(disease_candidates, default=(0, None))[1]
    return summary, detail, disease


class GuardrailApp(tk.Tk):
    MODES = (
        "完整生成（统计、报告、区间明细、图表）",
        "生成统计与报告（含区间明细）",
        "仅更新区间明细",
        "仅生成图表",
    )

    def __init__(self):
        super().__init__()
        self.title("G210交安设施统计与报告工具")
        self.geometry("980x760")
        self.minsize(860, 620)
        self.queue = Queue()
        self.running = False
        self.vars = {name: tk.StringVar() for name in ("project", "summary", "detail", "disease", "output")}
        self.template_name = tk.StringVar(value=next(iter(BUILTIN_REPORT_TEMPLATES)))
        self.item_vars = {
            "沿线设施": tk.BooleanVar(value=False),
            "中心高度": tk.BooleanVar(value=True),
            "螺栓缺失": tk.BooleanVar(value=True),
        }
        self.mode = tk.StringVar(value=self.MODES[0])
        self._build_ui()
        self.after(100, self._poll_queue)

    def _build_ui(self):
        style = ttk.Style(self)
        style.configure("Title.TLabel", font=("Microsoft YaHei UI", 16, "bold"))
        container = ttk.Frame(self, padding=18); container.pack(fill="both", expand=True)
        ttk.Label(container, text="G210交安设施统计与报告工具", style="Title.TLabel").pack(anchor="w", pady=(0, 14))
        form = ttk.LabelFrame(container, text="文件与文件夹", padding=12); form.pack(fill="x")
        rows = [
            ("项目文件夹", "project", self.choose_project, "选择文件夹"),
            ("分段汇总表", "summary", lambda: self.choose_file("summary", [("Excel", "*.xlsx")]), "选择文件"),
            ("检测明细文件夹", "detail", lambda: self.choose_folder("detail"), "选择文件夹"),
            ("病害清单文件夹", "disease", lambda: self.choose_folder("disease"), "选择文件夹"),
            ("输出文件夹", "output", lambda: self.choose_folder("output"), "选择文件夹"),
        ]
        for index, (label, key, command, button_text) in enumerate(rows[:4]):
            ttk.Label(form, text=label, width=15).grid(row=index, column=0, sticky="w", pady=5)
            ttk.Entry(form, textvariable=self.vars[key]).grid(row=index, column=1, sticky="ew", padx=8, pady=5)
            ttk.Button(form, text=button_text, command=command, width=12).grid(row=index, column=2, pady=5)
        template_row = 4
        ttk.Label(form, text="Word报告模板", width=15).grid(row=template_row, column=0, sticky="w", pady=5)
        ttk.Combobox(
            form,
            textvariable=self.template_name,
            values=tuple(BUILTIN_REPORT_TEMPLATES),
            state="readonly",
        ).grid(row=template_row, column=1, sticky="ew", padx=8, pady=5)
        ttk.Label(form, text="内置配置", foreground="#777777", width=12).grid(row=template_row, column=2, pady=5)
        output_label, output_key, output_command, output_button = rows[-1]
        output_row = 5
        ttk.Label(form, text=output_label, width=15).grid(row=output_row, column=0, sticky="w", pady=5)
        ttk.Entry(form, textvariable=self.vars[output_key]).grid(row=output_row, column=1, sticky="ew", padx=8, pady=5)
        ttk.Button(form, text=output_button, command=output_command, width=12).grid(row=output_row, column=2, pady=5)
        form.columnconfigure(1, weight=1)

        items = ttk.LabelFrame(container, text="处理细分项", padding=10); items.pack(fill="x", pady=(14, 0))
        for name in ("沿线设施", "中心高度", "螺栓缺失"):
            ttk.Checkbutton(items, text=name, variable=self.item_vars[name]).pack(side="left", padx=(0, 22))
        ttk.Label(items, text="沿线设施暂保留选择项，尚未开发", foreground="#777777").pack(side="left")

        action = ttk.LabelFrame(container, text="运行功能", padding=12); action.pack(fill="x", pady=14)
        ttk.Combobox(action, textvariable=self.mode, values=self.MODES, state="readonly", width=45).pack(side="left", fill="x", expand=True)
        self.run_button = ttk.Button(action, text="开始运行", command=self.start_run, width=14); self.run_button.pack(side="left", padx=(12, 0))
        self.open_button = ttk.Button(action, text="打开输出文件夹", command=self.open_output, width=16); self.open_button.pack(side="left", padx=(8, 0))

        self.progress = ttk.Progressbar(container, mode="indeterminate"); self.progress.pack(fill="x", pady=(0, 10))
        log_frame = ttk.LabelFrame(container, text="运行日志", padding=8); log_frame.pack(fill="both", expand=True)
        self.log_box = scrolledtext.ScrolledText(log_frame, wrap="word", font=("Consolas", 10), state="disabled")
        self.log_box.pack(fill="both", expand=True)
        ttk.Label(container, text="提示：若提示文件被占用，请先关闭已打开的Excel或Word文件。", foreground="#666666").pack(anchor="w", pady=(8, 0))

    def log(self, text):
        self.queue.put(("log", text))

    def choose_project(self):
        folder = filedialog.askdirectory(title="选择包含统计资料的项目文件夹")
        if not folder:
            return
        self.vars["project"].set(folder)
        summary, detail, disease = discover_paths(folder)
        if summary: self.vars["summary"].set(str(summary))
        if detail: self.vars["detail"].set(str(detail))
        if disease: self.vars["disease"].set(str(disease))
        self.vars["output"].set(str(summary.parent if summary else Path(folder)))
        self._append_log("已导入项目文件夹并自动识别相关文件。")

    def choose_file(self, key, types):
        path = filedialog.askopenfilename(filetypes=types)
        if path: self.vars[key].set(path)

    def choose_folder(self, key):
        path = filedialog.askdirectory()
        if path: self.vars[key].set(path)

    def build_config(self):
        project = Path(self.vars["project"].get() or ".").resolve()
        output = Path(self.vars["output"].get() or project).resolve()
        template_name = self.template_name.get()
        if template_name not in BUILTIN_REPORT_TEMPLATES:
            raise ValueError(f"未知的内置Word报告模板：{template_name}")
        return Config(
            project_dir=project,
            summary_xlsx=Path(self.vars["summary"].get()).resolve(),
            detail_dir=Path(self.vars["detail"].get()).resolve(),
            template_docx=BUILTIN_REPORT_TEMPLATES[template_name].resolve(),
            output_dir=output,
            disease_dir=Path(self.vars["disease"].get()).resolve() if self.vars["disease"].get() else None,
        )

    def validate(self, config):
        mode = self.mode.get()
        selected = {name for name, variable in self.item_vars.items() if variable.get()}
        if not selected:
            raise ValueError("请至少勾选一个处理细分项。")
        if not selected.intersection({"中心高度", "螺栓缺失"}):
            raise ValueError("沿线设施功能暂未开发，请同时勾选中心高度或螺栓缺失。")
        if mode in self.MODES[:2]:
            if not config.summary_xlsx.is_file(): raise FileNotFoundError("请选择有效的分段汇总表。")
            if not config.detail_dir.is_dir(): raise FileNotFoundError("请选择有效的检测明细文件夹。")
            if "螺栓缺失" in selected and (config.disease_dir is None or not config.disease_dir.is_dir()):
                raise FileNotFoundError("处理螺栓缺失时，请选择有效的病害清单文件夹。")
            if not config.template_docx.is_file():
                raise FileNotFoundError(f"内置Word报告模板不存在：{config.template_docx}")
        elif mode == self.MODES[2]:
            if "中心高度" not in selected: raise ValueError("更新区间明细仅适用于中心高度，请勾选中心高度。")
            if not config.out_xlsx.is_file(): raise FileNotFoundError(f"未找到统计工作簿：{config.out_xlsx}")
            if not config.summary_xlsx.is_file(): raise FileNotFoundError("请选择有效的分段汇总表。")
        elif mode == self.MODES[3]:
            if "中心高度" not in selected: raise ValueError("生成图表仅适用于中心高度，请勾选中心高度。")
            if not config.out_xlsx.is_file(): raise FileNotFoundError(f"未找到统计工作簿：{config.out_xlsx}")

    def start_run(self):
        if self.running:
            return
        try:
            config = self.build_config(); self.validate(config)
        except Exception as error:
            messagebox.showerror("参数错误", str(error)); return
        self.running = True; self.run_button.configure(state="disabled"); self.progress.start(10)
        self._append_log("=" * 60); self._append_log(f"开始：{self.mode.get()}")
        selected = {name for name, variable in self.item_vars.items() if variable.get()}
        Thread(target=self._worker, args=(config, self.mode.get(), selected), daemon=True).start()

    def _worker(self, config, mode, selected):
        try:
            options = {
                "process_height": "中心高度" in selected,
                "process_bolts": "螺栓缺失" in selected,
                "process_alongline": "沿线设施" in selected,
            }
            if mode == self.MODES[0]:
                generate_statistics_and_report(config, self.log, generate_charts_first=True, **options)
            elif mode == self.MODES[1]:
                generate_statistics_and_report(config, self.log, **options)
            elif mode == self.MODES[2]:
                add_interval_sheets(config.out_xlsx, config.summary_xlsx, self.log)
            else:
                add_charts(config.out_xlsx, self.log)
            self.queue.put(("done", "处理完成。"))
        except PermissionError:
            self.queue.put(("error", "文件被占用，无法保存。请关闭Excel或Word后重试。"))
        except Exception as error:
            self.queue.put(("error", f"{error}\n\n{traceback.format_exc()}"))

    def _poll_queue(self):
        try:
            while True:
                kind, text = self.queue.get_nowait()
                if kind == "log": self._append_log(text)
                elif kind == "done":
                    self._append_log(text); self._finish(); messagebox.showinfo("完成", text)
                elif kind == "error":
                    self._append_log("错误：" + text); self._finish(); messagebox.showerror("运行失败", text)
        except Empty:
            pass
        self.after(100, self._poll_queue)

    def _finish(self):
        self.running = False; self.progress.stop(); self.run_button.configure(state="normal")

    def _append_log(self, text):
        self.log_box.configure(state="normal")
        self.log_box.insert("end", str(text) + "\n"); self.log_box.see("end")
        self.log_box.configure(state="disabled")

    def open_output(self):
        folder = Path(self.vars["output"].get() or self.vars["project"].get() or ".")
        if folder.is_dir(): os.startfile(folder)
        else: messagebox.showwarning("提示", "输出文件夹不存在。")


# ==================== 广东项目模板层 ====================

def _norm_header(value):
    return re.sub(r"[\s（）()\[\]【】_:：\-]+", "", str(value or "")).lower()


def _first(row, *aliases, default=None):
    normalized = {_norm_header(k): v for k, v in row.items()}
    for alias in aliases:
        key = _norm_header(alias)
        if key in normalized and normalized[key] not in (None, ""):
            return normalized[key]
    return default


def _float(value):
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _route(value):
    return re.sub(r"\s+", "", str(value or "")).upper()


def _safe_city_component(value):
    text = str(value or "").strip()
    reserved = {"CON","PRN","AUX","NUL",*[f"COM{i}" for i in range(1,10)],*[f"LPT{i}" for i in range(1,10)]}
    if (not text or text in {".",".."} or Path(text).is_absolute() or
            re.search(r'[<>:"/\\|?*\x00-\x1f]', text) or text.rstrip(" .") != text or
            text.upper().split(".")[0] in reserved):
        raise ValueError(f"地市名称包含非法路径字符：{value}")
    return text


def _safe_excel_value(value):
    if isinstance(value, str) and value.lstrip().startswith(("=", "+", "-", "@")):
        return "'" + value
    return value


def _guardrail_type(value):
    text = str(value or "")
    if any(word in text for word in ("二波", "两波", "双波")):
        return "二波"
    if "三波" in text:
        return "三波"
    return text.strip()


def _station_first(value):
    text = str(value or "")
    match = re.search(r"K?\s*(\d+(?:\.\d+)?)\s*\+\s*(\d+(?:\.\d+)?)", text, re.I)
    if match:
        return float(match.group(1)) * 1000 + float(match.group(2))
    match = re.search(r"K\s*(\d+(?:\.\d+)?)", text, re.I)
    if match:
        return float(match.group(1)) * 1000
    return station_to_m(text.split("-")[0].split("~")[0])


class RouteCategoryIndex:
    VALID = {"高速公路", "普通国省道"}

    def __init__(self, mapping):
        self.mapping = {}
        self.display = {}
        for (city, route), value in mapping.items():
            norm = self._norm_city(city)
            self.mapping[(norm, _route(route))] = str(value).strip()
            self._remember_display(norm, str(city).strip())

    def _remember_display(self, norm, raw):
        """记录地市原始写法；同地市出现“韶关/韶关市”两种写法时优先保留带“市”的完整形式。"""
        current = self.display.get(norm)
        if current is None:
            self.display[norm] = raw
        elif raw.endswith("市") and not current.endswith("市") and raw[:-1] == current:
            self.display[norm] = raw

    @staticmethod
    def _norm_city(value):
        """地市匹配前归一化：去空白；末尾“市”仅在三字及以上时去掉（韶关市↔韶关）。"""
        text = re.sub(r"\s+", "", str(value or ""))
        return text[:-1] if text.endswith("市") and len(text) > 2 else text

    @classmethod
    def _sheet_category(cls, title):
        """按工作表名推断道路类别（附件3_4 格式无“道路类别”列）。"""
        normalized = _norm_header(title)
        if "国省道" in normalized:
            return "普通国省道"
        if "高速" in normalized:
            return "高速公路"
        return None

    @classmethod
    def from_file(cls, path):
        """支持两种路线分类表格式：
        1) 旧格式：任意工作表，首行表头含 地市/地区、路线/路线编号、道路类别/公路类别；
        2) 附件3_4 格式：工作表名区分类别（附件3-高速公路明细、附件4-普通国省道明细），
           首行为大标题、第2行表头，路线列为“路线”或“路线编码”，地市可带或不带“市”字。
        无地市+路线表头的工作表（如“里程汇总”）自动跳过。
        """
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        mapping = {}
        seen = {}
        for ws in wb.worksheets:
            sheet_category = cls._sheet_category(ws.title)
            rows = ws.iter_rows(values_only=True)
            headers = None
            for values in rows:
                cells = {_norm_header(v) for v in values if v not in (None, "")}
                if ("地市" in cells or "地区" in cells) and any(name in cells for name in ("路线", "路线编号", "路线编码")):
                    headers = list(values)
                    break
            if headers is None:
                continue
            for values in rows:
                if not any(v not in (None, "") for v in values):
                    continue
                row = dict(zip(headers, values))
                city = _first(row, "地市", "地区")
                route = _first(row, "路线", "路线编号", "路线编码")
                if not city or not route:
                    continue
                category = _first(row, "道路类别", "公路类别")
                category = str(category).strip() if category not in (None, "") else sheet_category
                if not category:
                    continue
                if category not in cls.VALID:
                    raise ValueError(f"路线分类值无效：{city}/{route}/{category}")
                raw_city = str(city).strip()
                key = (cls._norm_city(raw_city), _route(route))
                if key in seen and seen[key] != category:
                    raise ValueError(f"路线分类冲突：{raw_city}/{key[1]}")
                seen[key] = category
                mapping[(raw_city, _route(route))] = category
        wb.close()
        if not mapping:
            raise ValueError("路线分类表未读取到有效记录")
        return cls(mapping)

    def category(self, city, route):
        key = (self._norm_city(city), _route(route))
        if key not in self.mapping:
            raise KeyError(f"路线分类缺失：{str(city).strip()}/{key[1]}")
        return self.mapping[key]

    def resolve_city(self, route):
        cities = sorted({self.display.get(norm, norm) for norm, number in self.mapping if number == _route(route)})
        if not cities:
            raise KeyError(f"路线分类缺失：{_route(route)}")
        if len(cities) > 1:
            raise ValueError(f"路线{_route(route)}跨地市，无法唯一补全城市：{','.join(cities)}")
        return cities[0]

    def rows(self, city=None):
        target = self._norm_city(city) if city is not None else None
        return [{"city": self.display.get(norm, norm), "route": number, "category": value}
                for (norm, number), value in sorted(self.mapping.items()) if target is None or norm == target]

    def category_or_none(self, city, route):
        """查路线类别，查不到返回 None（不抛异常）。用于对人工对比记录等补充类别信息。"""
        key = (self._norm_city(city), _route(route))
        return self.mapping.get(key)

    def row_category(self, row):
        """按记录中的 地市/地区 + 路线 字段查类别，查不到返回 None。"""
        city = _first(row, "地市", "地区")
        route = _first(row, "路线", "路线编号", "路线编码")
        if city is None or route is None:
            return None
        return self.category_or_none(city, route)


class GuangdongInputScanner:
    """递归扫描项目资料，先按表头判型，再补全已判型标线 CSV 的文件名元数据。"""
    EXCLUDE_DIRS = ("原始数据",)  # 原始采集数据体量巨大且非统计表，默认跳过，避免误选根目录时卡死
    MAX_FILE_BYTES = 50 * 1024 * 1024  # 单文件上限，超大的直接跳过并记入问题清单，作为二次防护

    def __init__(self, root, route_index=None, default_city=None, log=lambda _x: None):
        self.root = Path(root)
        self.route_index = route_index
        self.default_city = default_city or self._city_from_folder()
        self.issues = []
        self.log = log

    def _city_from_folder(self):
        """若项目文件夹名含某地市（如“佛山市标线统计数据”），则其内容默认归属该市；
        用于跨地市贯通路线（G15/G105等）无法单凭路线号唯一补全城市时的回退。"""
        name = self.root.name
        if self.route_index:
            for (norm, _route_num) in self.route_index.mapping:
                if norm and norm in name:
                    return self.route_index.display.get(norm, norm)
        return None

    @staticmethod
    def _kind(headers):
        h = {_norm_header(x) for x in headers if x is not None}
        has = lambda prefix: any(value.startswith(_norm_header(prefix)) for value in h)
        if (has("逆反亮度系数") or has("逆反射亮度系数")) and (has("计算区间") or has("桩号")):
            return "marking"
        if has("拼接螺栓数量") and has("连接螺栓缺失数量"):
            return "bolt"
        if (has("梁板中心高度") or has("护栏中心高度")) and has("护栏类型"):
            return "height"
        return None

    def _xlsx_tables(self, path):
        dimension_a1 = False
        try:
            with ZipFile(path) as archive:
                xml = archive.read("xl/worksheets/sheet1.xml")
            dimension_a1 = b'<dimension ref="A1"' in xml
        except Exception:
            xml = b""
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        for index, ws in enumerate(wb.worksheets):
            if index == 0 and dimension_a1:
                continue
            values = list(ws.iter_rows(values_only=True))
            if values:
                yield ws.title, values[0], values[1:]
        wb.close()
        # 真实护栏文件 dimension 可能错误标为 A1；复用基线 XML 流读取回退。
        if dimension_a1:
            try:
                rows = list(iter_height_rows(path))
                if rows:
                    headers = list(rows[0])
                    yield "sheet1.xml回退", headers, [[row.get(h) for h in headers] for row in rows]
            except Exception as exc:
                self.issues.append(f"{path.name}: XML读取回退失败：{exc}")

    @staticmethod
    def _csv_table(path):
        last = None
        for encoding in ("utf-8-sig", "gbk", "utf-8"):
            try:
                with path.open("r", encoding=encoding, newline="") as stream:
                    rows = list(__import__("csv").reader(stream))
                return rows[0], rows[1:]
            except UnicodeDecodeError as exc:
                last = exc
        raise last

    def _metadata_from_marking_name(self, path):
        match = re.search(r"([GSHXY]\d+)\s*(上行|下行)", path.stem, re.I)
        if not match:
            return {}
        route, direction = match.groups()
        position_match = re.search(r"标线\d+", path.stem)
        position = position_match.group(0) if position_match else ""
        city = None
        km = re.search(r"K(\d+(?:\.\d+)?)K(\d+(?:\.\d+)?)", path.stem, re.I)
        segment = f"{_route(route)}{direction}K{km.group(1)}～K{km.group(2)}" if km else ""
        return {"city": city, "route": _route(route), "direction": direction,
                "marking_position": position, "file_segment": segment}

    def _convert(self, kind, row, source, sheet):
        city = str(_first(row, "地市", "地区", default="") or "").strip()
        route = _route(_first(row, "路线编号", "路线"))
        direction = str(_first(row, "方向", default="") or "").strip()
        metadata = self._metadata_from_marking_name(source)
        if not city: city = str(metadata.get("city") or "")
        if not route: route = metadata.get("route", "")
        if not direction: direction = metadata.get("direction", "")
        if not city and route and self.route_index:
            try:
                city = self.route_index.resolve_city(route)
            except ValueError:
                # 跨地市贯通路线（G15/G105等）：回退到项目文件夹隐含的地市
                if self.default_city:
                    city = self.default_city
                else:
                    raise
        station_raw = _first(row, "标注修正桩号", "电子修正桩号", "原始桩号", "桩号", "计算区间", "桩号范围")
        explicit_segment = _first(row, "检测区段", "统计区段", "区段", "桩号范围")
        base = {"city": city, "route": route, "direction": direction, "station_m": _station_first(station_raw),
                "segment": str(explicit_segment or metadata.get("file_segment") or station_raw or ""), "source": str(source), "sheet": sheet}
        if base["station_m"] is None:
            raise ValueError(f"{Path(source).name}/{sheet}：桩号无法解析：{station_raw}")
        if kind == "height":
            value = _float(_first(row, "梁板中心高度(mm)", "护栏中心高度(mm)", "梁板中心高度", "护栏中心高度", "梁板中心高度毫米", "护栏中心高度mm"))
            remark = str(_first(row, "异常标记", "备注标记", "备注", default="") or "").strip()
            if remark not in ("", "无备注"):
                # 桥梁/隧道路段标记行：不参与高度统计，但记录供报告生成区段说明
                if "桥梁" in remark or "隧道" in remark:
                    base.update(guardrail_note=remark)
                    return base
                return None
            if value is None or value <= 0: return None
            base.update(guardrail_type=_guardrail_type(_first(row, "护栏类型")), height=value)
        elif kind == "bolt":
            remark = str(_first(row, "备注标记", "异常标记", "备注", default="") or "").strip()
            vals = [_float(_first(row, name, name+"颗")) for name in ("拼接螺栓数量","拼接螺栓缺失数量","连接螺栓数量","连接螺栓缺失数量")]
            if any(v is None for v in vals):
                if remark and ("桥梁" in remark or "隧道" in remark):
                    base.update(guardrail_note=remark); return base
                return None
            base.update(splice=vals[0], splice_missing=vals[1], connection=vals[2], connection_missing=vals[3])
            if remark and ("桥梁" in remark or "隧道" in remark):
                base["guardrail_note"] = remark
        else:
            value = _float(_first(row, "逆反亮度系数", "逆反射亮度系数"))
            if value is None: return None
            for key in ("city","route","direction"):
                if not base[key]: base[key] = metadata.get(key, "")
            base.update(marking_position=str(_first(row,"标线位置","标线名称",default=metadata.get("marking_position", ""))),
                        value=value, target=_float(_first(row,"逆反亮度系数目标值",default=80)) or 80)
        if not base["city"] or not base["route"]:
            raise ValueError(f"{Path(source).name}/{sheet}：缺少地市或路线元数据")
        return base

    def _process_file(self, path):
        """单个文件处理（并行任务单元）"""
        records = {"height": [], "bolt": [], "marking": [], "notes": [], "issues": []}
        try:
            tables = self._xlsx_tables(path) if path.suffix.lower() == ".xlsx" else [(str(path.stem), *self._csv_table(path))]
            for sheet, headers, rows in tables:
                kind = self._kind(headers)
                if not kind: continue
                for values in rows:
                    row = dict(zip(headers, values))
                    try:
                        record = self._convert(kind, row, path, sheet)
                        if record:
                            if record.get("guardrail_note"):
                                records["notes"].append(record)
                            else:
                                records[kind].append(record)
                    except Exception as exc:
                        records["issues"].append(str(exc))
        except Exception as exc:
            records["issues"].append(f"{path.name}：读取失败：{exc}")
        return records

    def scan(self):
        """并行扫描项目资料，优化大数据量性能。"""
        if not self.root.is_dir(): raise FileNotFoundError(f"项目资料文件夹不存在：{self.root}")
        files = [p for p in sorted(self.root.rglob("*")) if p.is_file() and not p.name.startswith("~$") and p.suffix.lower() in (".xlsx", ".csv")]
        # 跳过显式排除的文件夹（如“原始数据”），避免误选根目录时读取海量原始采集文件而卡死
        kept, skipped_dir = [], []
        for p in files:
            parts = p.relative_to(self.root).parts
            (skipped_dir if any(part in self.EXCLUDE_DIRS for part in parts) else kept).append(p)
        files = kept
        # 单文件体积上限，超大的直接跳过并记入问题清单，作为二次防护
        huge = [(p, p.stat().st_size) for p in files if p.stat().st_size > self.MAX_FILE_BYTES]
        files = [p for p in files if p.stat().st_size <= self.MAX_FILE_BYTES]
        for p, size in huge:
            self.issues.append(f"{p.name}：文件过大（约{size/1048576:.0f}MB）已跳过，疑似非统计表原始数据")
        if skipped_dir or huge:
            self.log(f"待识别文件 {len(files)} 个（已跳过 {len(skipped_dir)} 个排除文件夹内文件、{len(huge)} 个超大文件）")
        else:
            self.log(f"待识别文件 {len(files)} 个")
        result = {"height": [], "bolt": [], "marking": [], "notes": [], "issues": self.issues}
        # ponytail: 并行处理，max_workers=4 避免内存溢出，必要时可调高
        with ThreadPoolExecutor(max_workers=4) as executor:
            futures = {executor.submit(self._process_file, f): f for f in files}
            done = 0
            for future in as_completed(futures):
                partial = future.result()
                done += 1
                for kind in ("height", "bolt", "marking", "notes"):
                    result[kind].extend(partial[kind])
                result["issues"].extend(partial["issues"])
                if done % 10 == 0 or done == len(files):
                    self.log(f"已识别文件 {done}/{len(files)}")
        return result



def detect_guangdong_data_folders(project_dir):
    """扫描项目文件夹，自动识别标线数据和护栏数据所在的子文件夹。
    返回 (marking_dir, guardrail_dir)，可能为 None。
    """
    project_dir = Path(project_dir)
    marking_candidates = {}
    guardrail_candidates = {}

    dirs_to_check = [project_dir] + sorted(
        d for d in project_dir.iterdir()
        if d.is_dir() and d.name not in GuangdongInputScanner.EXCLUDE_DIRS and not d.name.startswith(".")
    )

    for subdir in dirs_to_check:
        marking_count = 0
        guardrail_count = 0
        files = sorted(
            f for f in subdir.iterdir()
            if f.is_file() and not f.name.startswith("~$") and f.suffix.lower() in (".xlsx", ".csv")
        )
        for f in files[:20]:  # 每个目录最多检查20个文件
            try:
                if f.stat().st_size > 50 * 1024 * 1024:
                    continue
            except OSError:
                continue
            try:
                if f.suffix.lower() == ".xlsx":
                    wb = openpyxl.load_workbook(f, read_only=True, data_only=True)
                    for ws in wb.worksheets[:2]:
                        headers = [str(v) for v in next(ws.iter_rows(min_row=1, max_row=1, values_only=True), []) if v is not None]
                        kind = GuangdongInputScanner._kind(headers)
                        if kind == "marking":
                            marking_count += 1
                        elif kind in ("height", "bolt"):
                            guardrail_count += 1
                    wb.close()
                elif f.suffix.lower() == ".csv":
                    hdrs = None
                    for encoding in ("utf-8-sig", "gbk", "utf-8"):
                        try:
                            with open(f, "r", encoding=encoding) as fh:
                                hdrs = next(csv.reader(fh), [])
                            break
                        except UnicodeDecodeError:
                            continue
                    if hdrs is None:
                        continue
                    kind = GuangdongInputScanner._kind(hdrs)
                    if kind == "marking":
                        marking_count += 1
                    elif kind in ("height", "bolt"):
                        guardrail_count += 1
            except Exception:
                continue
        if marking_count:
            marking_candidates[str(subdir)] = marking_count
        if guardrail_count:
            guardrail_candidates[str(subdir)] = guardrail_count

    marking_dir = max(marking_candidates, key=marking_candidates.get, default=None)
    guardrail_dir = max(guardrail_candidates, key=guardrail_candidates.get, default=None)
    return marking_dir, guardrail_dir


class GuangdongStatistics:
    HEIGHT_LIMITS = {"二波": (580, 620), "三波": (677, 717)}

    @classmethod
    def height_summary(cls, rows):
        result = {}
        for kind in ("二波", "三波"):
            selected = [r for r in rows if _guardrail_type(r.get("guardrail_type")) == kind and _float(r.get("height")) is not None]
            low, high = cls.HEIGHT_LIMITS[kind]
            values = [float(r["height"]) for r in selected]
            qualified = sum(low <= v <= high for v in values)
            over = sum(v < low - 100 or v > high + 100 for v in values)
            result[kind] = {"valid_count": len(values), "average": sum(values)/len(values) if values else None,
                            "qualified_count": qualified, "qualified_rate": qualified/len(values) if values else None,
                            "over_10cm_count": over}
        return result

    @staticmethod
    def bolt_summary(rows):
        totals = {k: sum(float(r.get(k, 0) or 0) for r in rows) for k in ("splice","splice_missing","connection","connection_missing")}
        missing = totals["splice_missing"] + totals["connection_missing"]
        denominator = totals["splice"] + totals["connection"] + missing
        totals.update(missing_total=missing, missing_rate=missing/denominator if denominator else None)
        return totals

    @staticmethod
    def marking_summary(rows):
        values=[float(r["value"]) for r in rows if _float(r.get("value")) is not None]
        qualified=sum(float(r["value"]) >= float(r.get("target",80)) for r in rows if _float(r.get("value")) is not None)
        return {"valid_count":len(values),"average":sum(values)/len(values) if values else None,
                "qualified_count":qualified,"qualified_rate":qualified/len(values) if values else None}

    @staticmethod
    def _group(rows, fields):
        grouped = {}
        for row in rows:
            key = tuple(row.get(field, "") for field in fields)
            grouped.setdefault(key, []).append(row)
        return grouped

    @classmethod
    def marking_segment_summary(cls, rows):
        fields = ("category", "route", "direction", "segment", "marking_position")
        result = []
        for key, selected in cls._group(rows, fields).items():
            summary = cls.marking_summary(selected)
            result.append(dict(zip(fields, key), **summary))
        return sorted(result, key=lambda row: tuple(str(row.get(field, "")) for field in fields))

    @classmethod
    def marking_segment_pair_summary(cls, rows):
        fields = ("category", "route", "direction", "segment")
        result = []
        for key, selected in cls._group(rows, fields).items():
            item = dict(zip(fields, key))
            side_names = marking_side_names(r.get("marking_position") for r in selected)
            for pos in side_names:
                summary = cls.marking_summary([r for r in selected if str(r.get("marking_position", "")) == pos])
                item[f"{pos}_valid_count"] = summary["valid_count"]
                item[f"{pos}_average"] = summary["average"]
                item[f"{pos}_qualified_count"] = summary["qualified_count"]
                item[f"{pos}_qualified_rate"] = summary["qualified_rate"]
            item["_side_names"] = side_names
            result.append(item)
        return sorted(result, key=lambda row: tuple(str(row.get(field, "")) for field in fields))

    @classmethod
    def height_segment_summary(cls, rows):
        fields = ("category", "route", "direction", "segment", "guardrail_type")
        result = []
        for key, selected in cls._group(rows, fields).items():
            kind = _guardrail_type(key[-1])
            if kind not in cls.HEIGHT_LIMITS:
                continue
            summary = cls.height_summary(selected).get(kind, {})
            result.append(dict(zip(fields, key), **summary))
        return sorted(result, key=lambda row: tuple(str(row.get(field, "")) for field in fields))

    @classmethod
    def bolt_segment_summary(cls, rows):
        fields = ("category", "route", "direction", "segment")
        result = []
        for key, selected in cls._group(rows, fields).items():
            result.append(dict(zip(fields, key), **cls.bolt_summary(selected)))
        return sorted(result, key=lambda row: tuple(str(row.get(field, "")) for field in fields))

    @staticmethod
    def continuous_marking_weak(rows, minimum_length_m=3000):
        groups = {}
        for row in rows:
            key=(row.get("route"),row.get("direction"),row.get("marking_position"))
            groups.setdefault(key,[]).append(row)
        weak=[]
        for key, selected in groups.items():
            selected=sorted(selected,key=lambda r:r.get("station_m") if r.get("station_m") is not None else float("inf"))
            start=previous=None
            def flush():
                if start is not None and previous is not None and abs(previous-start) >= minimum_length_m:
                    weak.append({"route":key[0],"direction":key[1],"marking_position":key[2],"start_m":start,"end_m":previous})
            for row in selected:
                station=row.get("station_m"); value=_float(row.get("value")); target=_float(row.get("target")) or 80
                bad=station is not None and value is not None and value < target
                if not bad:
                    flush(); start=previous=None; continue
                if start is None:
                    start=previous=station
                elif abs(station-previous) <= 50:
                    previous=station
                else:
                    flush(); start=previous=station
            flush()
        return weak


class ManualAutoComparator:
    def __init__(self, thresholds): self.thresholds = thresholds

    @staticmethod
    def read_file(path):
        wb=openpyxl.load_workbook(path,read_only=True,data_only=True)
        records=[]; issues=[]
        for ws in wb.worksheets:
            title=ws.title
            rows=list(ws.iter_rows(values_only=True))
            if not rows: continue
            first="|".join(_norm_header(value) for value in rows[0] if value not in (None,""))
            second="|".join(_norm_header(value) for value in rows[1] if value not in (None,"")) if len(rows)>1 else ""
            combined=first+"|"+second
            if ("人工复核护栏中心平均高度" in combined or "人工护栏中心高度" in combined) and "自动化护栏中心高度" in combined:
                indicator="height"; start=1
            elif "人工逆反射亮度系数平均值" in combined and "自动化逆反射亮度系数平均值" in combined:
                indicator="marking"; start=1
            elif ("拼接螺栓缺失数量" in combined and "连接螺栓缺失数量" in combined and
                  ("人工复核螺栓缺失数量" in combined or "自动化螺栓缺失数量" in combined)):
                indicator="bolt"; start=2
            else:
                issues.append(f"{Path(path).name}/{title}：未按业务表头识别到人工复核指标")
                continue
            headers=list(rows[0])
            for row_index, values in enumerate(rows[start:], start=start+1):
                if not any(v not in (None,"") for v in values): continue
                row=dict(zip(headers,values)); city=_first(row,"地市","地区"); route=_first(row,"路线","路线编号"); direction=_first(row,"方向"); segment=_first(row,"桩号范围","桩号","计算区间")
                # 原样保留展示字段，供报告生成三张分项对比表
                display={"gtype":str(_first(row,"护栏类型",default="") or "").strip(),
                         "position":str(_first(row,"护栏位置","标线位置",default="") or "").strip(),
                         "remark":str(_first(row,"备注",default="") or "").strip()}
                if indicator=="bolt":
                    manual=(_float(values[7]) or 0)+(_float(values[8]) or 0) if len(values)>10 else None
                    automatic=(_float(values[9]) or 0)+(_float(values[10]) or 0) if len(values)>10 else None
                    display.update(msplice=_float(values[7]),mconn=_float(values[8]),asplice=_float(values[9]),aconn=_float(values[10]))
                elif indicator=="height":
                    manual=_float(_first(row,"人工复核护栏中心平均高度mm","人工护栏中心高度","人工值")); automatic=_float(_first(row,"自动化护栏中心高度mm","自动化护栏中心高度","自动化值"))
                else:
                    manual=_float(_first(row,"人工逆反射亮度系数平均值","人工值")); automatic=_float(_first(row,"自动化逆反射亮度系数平均值","自动化值"))
                if not city or not route or manual is None or automatic is None:
                    issues.append(f"{Path(path).name}/{title}/第{row_index}行：人工对比记录字段缺失或非数值"); continue
                records.append({"indicator":indicator,"city":str(city).strip(),"route":_route(route),"direction":str(direction or ""),"segment":str(segment or ""),"manual":manual,"automatic":automatic,**display})
        wb.close(); return records,issues

    @staticmethod
    def summarize(detail):
        """按指标汇总人工/自动化对比明细（可传入按道路类别过滤后的明细）。
        标线、螺栓缺失使用相对偏差(%)；护栏中心高度使用绝对偏差(mm)。"""
        summary={}
        for indicator in ("marking","height","bolt"):
            selected=[x for x in detail if x.get("indicator")==indicator]
            if indicator=="height":
                computable=[x for x in selected if x.get("absolute_difference") is not None]
                metric="absolute_difference"
            else:
                computable=[x for x in selected if x.get("relative_deviation") is not None]
                metric="relative_deviation"
            deviations=[x[metric] for x in computable]; within=sum(bool(x.get("within_threshold")) for x in computable)
            summary[indicator]={"paired_count":len(selected),"computable_count":len(computable),"min":min(deviations) if deviations else None,
                "max":max(deviations) if deviations else None,"average":sum(deviations)/len(deviations) if deviations else None,
                "within_count":within,"consistency_rate":within/len(computable) if computable else None,
                "zero_manual_count":sum(x.get("relative_deviation") is None for x in selected)}
        return summary

    def compare(self, records):
        detail=[]
        for record in records:
            item=dict(record); manual=_float(item.get("manual")); automatic=_float(item.get("automatic"))
            if manual is None or automatic is None: continue
            item["absolute_difference"]=abs(automatic-manual)
            item["relative_deviation"]=None if manual == 0 else abs(automatic-manual)/abs(manual)*100
            indicator=item["indicator"]; threshold=self.thresholds[indicator]
            if indicator=="height":
                # 护栏中心高度：使用绝对偏差(mm)判断，不使用百分比
                item["within_threshold"]=item["absolute_difference"] <= threshold
            else:
                # 标线/螺栓缺失数量：使用相对偏差(%)判断
                item["within_threshold"]=None if manual == 0 else item["relative_deviation"] <= threshold
            detail.append(item)
        return detail, self.summarize(detail)




class GuangdongChapterWriter:
    @staticmethod
    def _fmt(value, digits=2):
        if value is None or value == "": return "—"
        if isinstance(value, (int, float)): return f"{value:,.{digits}f}"
        return str(value)

    @staticmethod
    def _pct(value):
        return "—" if value is None else f"{value:.2%}"

    @staticmethod
    def _segment_text(value):
        text=str(value or "—")
        return text if text.endswith("段") or text == "—" else text + "段"

    @staticmethod
    def _add_table(doc, headers, rows, merge=None):
        """表格：所有单元格文字无缩进、水平居中、垂直居中；表头加粗仿宋_GB2312。
        merge={起始列: 跨列数}：两级表头——第一行主表头（横向合并），其下列子表头；单列表头纵向合并两行。"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def _style(cell, bold=False):
            for p in cell.paragraphs:
                p.alignment=1  # 水平居中
                pPr=p._p.get_or_add_pPr()
                ind=pPr.find(qn('w:ind'))
                if ind is None:
                    ind=OxmlElement('w:ind'); pPr.append(ind)
                ind.set(qn('w:firstLineChars'),'0'); ind.set(qn('w:firstLine'),'0')
                tcPr=cell._tc.get_or_add_tcPr()
                vAlign=tcPr.find(qn('w:vAlign'))
                if vAlign is None:
                    vAlign=OxmlElement('w:vAlign'); tcPr.append(vAlign)
                vAlign.set(qn('w:val'),'center')  # 垂直居中
                for run in p.runs:
                    run.bold=bold; run.font.name="仿宋_GB2312"

        if merge:
            # 两级表头：merge={列号: (主表头名, 跨列数)}，原生 XML 构建（gridSpan/vMerge），物理列 = 逻辑列 - Σ(跨列数-1)
            from docx.table import _Cell
            merge = {c: (v[0], int(v[1])) for c, v in merge.items()}
            phys_of={}; p=0; i=0
            while i < len(headers):
                if i in merge:
                    n=merge[i][1]
                    for j in range(i, i+n): phys_of[j]=p
                    p+=n-1; i+=n
                else:
                    phys_of[i]=p; p+=1; i+=1
            physical=p
            table=doc.add_table(rows=2,cols=physical); table.style="Table Grid"
            # 清空默认两行，手工构建 tr
            tbl=table._tbl
            for tr in list(tbl.findall(qn('w:tr'))): tbl.remove(tr)
            def _tc(text, span=None, vmerge=None, bold=True):
                tc=OxmlElement('w:tc'); tcPr=OxmlElement('w:tcPr')
                if span: sp=OxmlElement('w:gridSpan'); sp.set(qn('w:val'),str(span)); tcPr.append(sp)
                if vmerge is not None:
                    vm=OxmlElement('w:vMerge'); vm.set(qn('w:val'),vmerge) if vmerge!="continue" else None; tcPr.append(vm)
                tc.append(tcPr); par=OxmlElement('w:p'); r=OxmlElement('w:r')
                rPr=OxmlElement('w:rPr')
                if bold: b=OxmlElement('w:b'); rPr.append(b)
                fonts=OxmlElement('w:rFonts'); fonts.set(qn('w:eastAsia'),"仿宋_GB2312"); fonts.set(qn('w:ascii'),"Times New Roman"); fonts.set(qn('w:hAnsi'),"Times New Roman"); rPr.append(fonts)
                r.append(rPr)
                t=OxmlElement('w:t'); t.text=str(text); r.append(t); par.append(r); tc.append(par)
                return tc
            top=OxmlElement('w:tr'); sub=OxmlElement('w:tr')
            i=0
            while i < len(headers):
                if i in merge and merge[i][0]:
                    parent,n=merge[i]
                    top.append(_tc(parent, span=n-1))
                    for j in range(i+1, i+n):
                        sub.append(_tc(headers[j]))
                    i+=n
                elif i in merge:
                    top.append(_tc(headers[i], vmerge="restart"))
                    sub.append(_tc("", vmerge="continue"))
                    i+=1
                else:
                    top.append(_tc(headers[i], vmerge="restart"))
                    sub.append(_tc("", vmerge="continue"))
                    i+=1
            tbl.append(top); tbl.append(sub)
        else:
            table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"
            for index, header in enumerate(headers):
                cell=table.rows[0].cells[index]; cell.text=str(header); _style(cell, True)
        if not rows:
            blank_cols = physical if merge else len(headers)
            rows = [["—"] + [""] * (blank_cols - 1)]
        for row in rows:
            cells=table.add_row().cells
            for index,value in enumerate(row):
                cells[index].text=str(value)
                _style(cells[index])
        return table

    @classmethod
    def _body(cls, doc, text):
        """正文段落：仿宋_GB2312五号，首行缩进2字符（firstLineChars=200）。"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        paragraph=doc.add_paragraph(text)
        pPr=paragraph._p.get_or_add_pPr()
        ind=pPr.find(qn('w:ind'))
        if ind is None:
            ind=OxmlElement('w:ind'); pPr.append(ind)
        ind.set(qn('w:firstLineChars'),'200')
        ind.set(qn('w:firstLine'),'420')  # 备用磅值：2字符×10.5pt≈21pt=420twips
        return paragraph

    @classmethod
    def _indent_existing_body(cls, doc):
        """给已有正文段落（模板遗留正文）补首行缩进2字符；跳过标题与空段。"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
            style=paragraph.style.name if paragraph.style is not None else ""
            if style.startswith("Heading"):
                continue
            pPr=paragraph._p.get_or_add_pPr()
            ind=pPr.find(qn('w:ind'))
            if ind is None:
                ind=OxmlElement('w:ind'); pPr.append(ind)
            ind.set(qn('w:firstLineChars'),'200')
            ind.set(qn('w:firstLine'),'420')

    @classmethod
    def _remove_template_placeholder(cls, doc):
        """删除模板自带的占位章节结构：从首个 '（一）' 起，到首个 '（三）'前，
        移除所有直系 body 段落/表格元素（含 Heading2/3/4 占位标题及其下属正文），
        但保留末尾 sectPr（节属性）。若模板不含 '（三）'，则删到 sectPr 之前为止。"""
        from docx.oxml.ns import qn
        body=doc.element.body
        children=list(body)
        start_idx=None; end_idx=len(children)
        for i,el in enumerate(children):
            if el.tag.endswith('}sectPr'):
                end_idx=i  # 永远不删除节属性
                break
            if el.tag.endswith('}p'):
                texts="".join(node.text or "" for node in el.iter(qn('w:t')))
                if start_idx is None and texts.startswith("（一）"):
                    start_idx=i
                elif start_idx is not None and texts.startswith("（三）"):
                    end_idx=i; break
        if start_idx is None:
            return
        for el in children[start_idx:end_idx]:
            body.remove(el)

    @staticmethod
    def _outline_level(paragraph):
        """返回 Word 实际大纲等级（1-based）；优先读取 w:outlineLvl。"""
        from docx.oxml.ns import qn
        import re
        pPr = paragraph._p.pPr
        if pPr is not None:
            outline = pPr.find(qn("w:outlineLvl"))
            if outline is not None:
                try:
                    return int(outline.get(qn("w:val"))) + 1
                except (TypeError, ValueError):
                    pass
        style = paragraph.style.name if paragraph.style is not None else ""
        match = re.search(r"Heading\s*(\d+)", style, re.I)
        return int(match.group(1)) if match else None

    @staticmethod
    def _clear_paragraph_indent(paragraph):
        """标题/题注不保留模板首行、左右或悬挂缩进。"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
        fmt = paragraph.paragraph_format
        fmt.left_indent = Pt(0)
        fmt.right_indent = Pt(0)
        fmt.first_line_indent = Pt(0)
        pPr = paragraph._p.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        for key in ("left", "right", "firstLine", "hanging", "leftChars", "rightChars", "firstLineChars", "hangingChars"):
            ind.set(qn(f"w:{key}"), "0")

    @classmethod
    def _apply_heading_format(cls, paragraph, level, main_title=False):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
        level = max(1, int(level or 1))
        pPr = paragraph._p.get_or_add_pPr()
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            pPr.append(outline)
        outline.set(qn("w:val"), str(level - 1))
        paragraph.alignment = 0
        cls._clear_paragraph_indent(paragraph)
        east_asia = "黑体" if level in (1, 2, 3) else "仿宋_GB2312"
        size = 16 if main_title else {1: 16, 2: 15, 3: 14, 4: 12, 5: 10.5}.get(level, 10.5)
        for run in paragraph.runs:
            run.bold = False
            run.italic = False
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            cls._set_run_fonts(run, east_asia)

    @staticmethod
    def _outline_level(paragraph):
        from docx.oxml.ns import qn
        import re
        pPr = paragraph._p.pPr
        if pPr is not None:
            node = pPr.find(qn("w:outlineLvl"))
            if node is not None:
                try:
                    return int(node.get(qn("w:val"))) + 1
                except (TypeError, ValueError):
                    pass
        style = paragraph.style.name if paragraph.style is not None else ""
        match = re.search(r"Heading\s*(\d+)", style, re.I)
        return int(match.group(1)) if match else None

    @staticmethod
    def _clear_paragraph_indent(paragraph):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
        fmt = paragraph.paragraph_format
        fmt.left_indent = Pt(0)
        fmt.right_indent = Pt(0)
        fmt.first_line_indent = Pt(0)
        pPr = paragraph._p.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        for key in ("left", "right", "firstLine", "hanging", "leftChars", "rightChars", "firstLineChars", "hangingChars"):
            ind.set(qn(f"w:{key}"), "0")

    @classmethod
    def _apply_heading_format(cls, paragraph, level, main_title=False):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
        level = max(1, int(level or 1))
        pPr = paragraph._p.get_or_add_pPr()
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            pPr.append(outline)
        outline.set(qn("w:val"), str(level - 1))
        paragraph.alignment = 0
        cls._clear_paragraph_indent(paragraph)
        east_asia = "黑体" if level in (1, 2, 3) else "仿宋_GB2312"
        size = 16 if main_title else {1: 16, 2: 15, 3: 14, 4: 12, 5: 10.5}.get(level, 10.5)
        for run in paragraph.runs:
            run.bold = False
            run.italic = False
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            cls._set_run_fonts(run, east_asia)

    @classmethod
    def _heading(cls, doc, text, level):
        paragraph = doc.add_paragraph()
        if 1 <= level <= 5 and f"Heading {level}" in doc.styles:
            paragraph.style = doc.styles[f"Heading {level}"]
        paragraph.add_run(text)
        cls._apply_heading_format(paragraph, level)
        return paragraph

    @classmethod
    def _comparison_table(cls, doc, detail, thresholds, table_adder=None):
        """人工复核对比：先给判断标准说明，再按 护栏中心高度/螺栓缺失/标线逆反射 生成三张分项对比表。"""
        cls._body(doc, f"人工复核分别采用以下一致性判断标准：标线逆反射亮度系数按相对偏差判断，阈值{thresholds['marking']}%；波形梁护栏中心高度按绝对偏差判断，阈值{thresholds['height']} mm；波形梁护栏螺栓缺失数量按相对偏差判断，阈值{thresholds['bolt']}%。相对偏差为自动化检测值与人工复核值之差的绝对值占人工复核值的百分比，绝对偏差为二者之差的绝对值；偏差在阈值范围内的记录判定为一致，一致性占比为一致记录数占可计算记录数的比例。")
        # 精简版偏差分析：只保留平均偏差+一致性占比，不输出 min～max 偏差范围
        summary = ManualAutoComparator.summarize(detail)
        names = {"marking": "标线逆反射亮度系数", "height": "波形梁护栏中心高度", "bolt": "波形梁护栏螺栓缺失"}
        sentences = []
        for key, name in names.items():
            s = summary.get(key, {})
            if s.get("computable_count"):
                unit = "%" if key in ("marking", "bolt") else " mm"
                sentences.append(f"{name}:自动化检测与人工复核平均偏差{cls._fmt(s.get('average'))}{unit}，一致性占比{cls._pct(s.get('consistency_rate'))}")
        if sentences:
            cls._body(doc, "偏差范围与一致性占比分析：" + "；".join(sentences) + "。")
        fmt = lambda v: "" if v is None else (f"{v:g}" if isinstance(v, float) else str(v))
        def emit(headers, rows, title, merge=None):
            if table_adder is None:
                cls._add_table(doc, headers, rows, merge=merge)
            else:
                table_adder(headers, rows, title, merge)

        height=[r for r in detail if r.get("indicator")=="height"]
        bolt=[r for r in detail if r.get("indicator")=="bolt"]
        marking=[r for r in detail if r.get("indicator")=="marking"]

        emit(["路线","护栏类型","护栏位置","方向","桩号范围","人工复核护栏中心平均高度(mm)","自动化护栏中心高度(mm)","备注"],
             [[r.get("route"),r.get("gtype"),r.get("position"),r.get("direction"),r.get("segment"),fmt(r.get("manual")),fmt(r.get("automatic")),r.get("remark")] for r in height],
             "人工复核护栏中心高度对比表")

        # 螺栓两级表头：主表头横跨2个子列（拼接/连接），无空单元格
        if bolt:
            emit(["路线","护栏类型","护栏位置","方向","桩号范围","人工复核螺栓缺失数量","拼接螺栓缺失数量","连接螺栓缺失数量","自动化螺栓缺失数量","拼接螺栓缺失数量","连接螺栓缺失数量","备注"],
                 [[r.get("route"),r.get("gtype"),r.get("position"),r.get("direction"),r.get("segment"),fmt(r.get("msplice")),fmt(r.get("mconn")),fmt(r.get("asplice")),fmt(r.get("aconn")),r.get("remark")] for r in bolt],
                 "人工复核螺栓缺失对比表", merge={5:("人工复核螺栓缺失数量",3),8:("自动化螺栓缺失数量",3)})

        emit(["路线","标线位置","方向","桩号范围","人工逆反射亮度系数平均值","自动化逆反射亮度系数平均值"],
             [[r.get("route"),r.get("position"),r.get("direction"),r.get("segment"),fmt(r.get("manual")),fmt(r.get("automatic"))] for r in marking],
             "人工复核标线逆反射对比表")

    @staticmethod
    def _set_run_fonts(run, east_asia="仿宋_GB2312", latin="Times New Roman"):
        from docx.oxml.ns import qn
        rPr=run._element.get_or_add_rPr()
        rFonts=rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts=rPr.makeelement(qn('w:rFonts'),{}); rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), east_asia)
        rFonts.set(qn('w:ascii'), latin)
        rFonts.set(qn('w:hAnsi'), latin)

    @classmethod
    def _configure_heading_styles(cls, doc):
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import OxmlElement
        if "Caption" not in doc.styles:
            doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
        for level in range(1, 6):
            name = f"Heading {level}"
            if name not in doc.styles:
                style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            else:
                style = doc.styles[name]
            east_asia = "黑体" if level in (1, 2, 3) else "仿宋_GB2312"
            style.font.name = east_asia
            style.font.size = Pt({1: 16, 2: 15, 3: 14, 4: 12, 5: 10.5}[level])
            style.font.bold = False
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.left_indent = Pt(0)
            style.paragraph_format.right_indent = Pt(0)
            style.paragraph_format.first_line_indent = Pt(0)
            style.paragraph_format.alignment = 0
            rPr = style._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), east_asia)
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")

    @classmethod
    def _format_existing_headings(cls, doc):
        main_prefix = "五、交通安全设施技术状况检测评价"
        for paragraph in doc.paragraphs:
            if paragraph.text.strip().startswith(main_prefix) and "结果" in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace("结果", "情况")
            style = paragraph.style.name if paragraph.style is not None else ""
            text = paragraph.text.strip()
            level = cls._outline_level(paragraph)
            is_main = text.startswith(main_prefix)
            is_heading = is_main or level is not None or style.startswith("Heading")
            if not is_heading:
                continue
            level = level or (1 if is_main else 1)
            if level <= 5 and f"Heading {level}" in doc.styles:
                paragraph.style = doc.styles[f"Heading {level}"]
            cls._apply_heading_format(paragraph, level, main_title=is_main)

    @classmethod
    def _format_chart_caption(cls, paragraph):
        from docx.shared import Pt, RGBColor
        paragraph.alignment = 1
        cls._clear_paragraph_indent(paragraph)
        for run in paragraph.runs:
            run.bold = False
            run.italic = False
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
            cls._set_run_fonts(run, "黑体")

    @classmethod
    def _format_all_run_fonts(cls, doc):
        main_prefix = "五、交通安全设施技术状况检测评价"
        for paragraph in doc.paragraphs:
            style = paragraph.style.name if paragraph.style is not None else ""
            text = paragraph.text.strip()
            level = cls._outline_level(paragraph)
            is_main = text.startswith(main_prefix)
            is_heading = is_main or level is not None or style.startswith("Heading")
            if style == "Caption":
                cls._format_chart_caption(paragraph)
            elif is_heading:
                cls._apply_heading_format(paragraph, level or 1, main_title=is_main)
            else:
                for run in paragraph.runs:
                    cls._set_run_fonts(run, "仿宋_GB2312")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            cls._set_run_fonts(run, "仿宋_GB2312")

    @staticmethod
    def _alpha_label(index):
        value = int(index)
        result = ""
        while value:
            value, remainder = divmod(value - 1, 26)
            result = chr(ord("a") + remainder) + result
        return result + "."

    @staticmethod
    def _chapter_number(doc):
        import re
        chinese = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}
        for paragraph in doc.paragraphs:
            match = re.match(r"^([一二三四五六七八九十]+|\d+)、", paragraph.text.strip())
            if not match:
                continue
            token = match.group(1)
            if token.isdigit():
                return int(token)
            if token == "十":
                return 10
            if len(token) == 2 and token[0] == "十":
                return 10 + chinese[token[1]]
            if len(token) == 2 and token[1] == "十":
                return chinese[token[0]] * 10
            if len(token) == 3 and token[1] == "十":
                return chinese[token[0]] * 10 + chinese[token[2]]
            return chinese.get(token, 1)
        return 1

    @classmethod
    def _marking_segment_sentence(cls, row):
        names = row.get("_side_names") or {}
        ordered = sorted(names)
        if not ordered:
            ordered = ["标线2", "标线3"]; names = marking_side_names(ordered)
        parts = []
        for pos in ordered:
            name = names[pos]
            count, avg, rate = row.get(f"{pos}_valid_count", 0) or 0, cls._fmt(row.get(f"{pos}_average")), cls._pct(row.get(f"{pos}_qualified_rate"))
            parts.append(f"{name}共{count:,}个有效计算单元，平均逆反射亮度系数{avg}，合格率{rate}")
        return "该区段" + "；".join(parts) + "。"

    @staticmethod
    def _height_segment_sentence(rows):
        parts = []
        total = 0
        for row in rows:
            count = int(row.get("valid_count") or 0)
            total += count
            parts.append(f"{_guardrail_type(row.get('guardrail_type'))}护栏{count:,}个有效点，平均高度{GuangdongChapterWriter._fmt(row.get('average'))} mm，合格率{GuangdongChapterWriter._pct(row.get('qualified_rate'))}")
        return f"该区段护栏中心高度共检测{total:,}个有效点，其中" + "；".join(parts) + "。"

    @staticmethod
    def _bolt_segment_sentence(row):
        return (f"该区段识别现有拼接螺栓{int(row.get('splice', 0) or 0):,}颗、连接螺栓{int(row.get('connection', 0) or 0):,}颗，检出缺失螺栓{int(row.get('missing_total', 0) or 0):,}颗，螺栓缺失率为{GuangdongChapterWriter._pct(row.get('missing_rate'))}。")

    @staticmethod
    def _guardrail_note_counts(notes):
        """按区段聚合桥梁/隧道标记：{(路线,方向,区段): {"桥梁":n,"隧道":n}}。"""
        result = {}
        for row in notes or []:
            remark = str(row.get("guardrail_note") or "")
            key = (row.get("route"), row.get("direction"), row.get("segment"))
            counts = result.setdefault(key, {"桥梁": 0, "隧道": 0})
            if "桥梁" in remark: counts["桥梁"] += 1
            if "隧道" in remark: counts["隧道"] += 1
        return result

    @staticmethod
    def _guardrail_note_sentence(counts):
        if not counts or not (counts["桥梁"] or counts["隧道"]): return None
        if counts["桥梁"] and counts["隧道"]: return "该区段为桥梁和隧道路段，区段内无护栏"
        if counts["桥梁"]: return "该区段为桥梁路段，区段内无护栏"
        return "该区段为隧道路段，区段内无护栏"

    @staticmethod
    def _guardrail_no_valid_point_sentence(counts):
        if not counts or not (counts["桥梁"] or counts["隧道"]): return None
        if counts["桥梁"] and counts["隧道"]: return "当前区段为桥梁和隧道路段，无有效检测点位。"
        if counts["桥梁"]: return "当前区段为桥梁路段，无有效检测点位。"
        return "当前区段为隧道路段，无有效检测点位。"

    @classmethod
    def write(cls,city,bundle,output_dir,template,thresholds):
        from docx import Document
        from docx.shared import Pt
        city=_safe_city_component(city)
        template=Path(template)
        if not template.is_file(): raise FileNotFoundError(f"Word模板不存在：{template}")
        if template.suffix.lower() != ".md":
            raise ValueError(f"仅支持 Markdown 模板：{template}，请使用 .md 模板。")
        from backend import markdown_skeleton
        doc = Document()
        blocks = markdown_skeleton.read_blocks(template)
        # 渲染完整骨架：标题/正文/表格/图片，支持 {{地市}} 占位与显式锚点（见 minimal_docx 锚点文档）。
        # 若模板包含 <!-- inject:* --> 标记，动态章节将在对应位置注入；否则追加末尾（兼容旧模板）。
        def _skeleton_picture(caption: str):
            try:
                media_path = (template.parent / caption).resolve()
                if media_path.is_file():
                    try:
                        doc.add_picture(str(media_path), width=Pt(380))
                    except Exception:
                        pass
            except Exception:
                pass

        markdown_skeleton.render_skeleton(
            doc, blocks,
            heading=lambda text, level: cls._heading(doc, text, min(level, 2)),
            body=lambda text: cls._body(doc, text),
            table=lambda headers, rows, shading: cls._add_table(doc, headers, rows),
            picture=_skeleton_picture,
            replace={"{{地市}}": city},
        )
        chapter_no = 1
        for style_name in ("Normal","Body Text"):
            if style_name in doc.styles:
                doc.styles[style_name].font.name="仿宋_GB2312"; doc.styles[style_name].font.size=Pt(10.5)
        caption_counts={"figure":0,"table":0}

        def _caption(kind,title):
            caption_counts[kind]+=1
            prefix="图" if kind=="figure" else "表"
            cap=doc.add_paragraph(f"{prefix}{chapter_no}-{caption_counts[kind]} {title}")
            cap.style="Caption"
            cls._format_chart_caption(cap)

        def _table(headers,rows,title,merge=None):
            _caption("table",title)
            return cls._add_table(doc,headers,rows,merge=merge)

        def _figure(path,title):
            doc.add_picture(path,width=Pt(440))
            _caption("figure",title)

        for category_number,(category,chapter_title) in enumerate((("高速公路","（一）高速公路交安设施技术状况"),("普通国省道","（二）普通国省道交安设施技术状况")),1):
            all_mark=[r for r in bundle.get("marking",[]) if r.get("category")==category]
            all_height=[r for r in bundle.get("height",[]) if r.get("category")==category]
            all_bolt=[r for r in bundle.get("bolt",[]) if r.get("category")==category]
            all_detail=[r for r in bundle.get("comparison_detail",[]) if r.get("category")==category]
            chart_prefix=f"{city}_{category}_"
            base=Path(output_dir)/city/"charts"
            def _images(suffix):
                return sorted(str(p) for p in base.glob(f"{chart_prefix}*{suffix}*.png")) if base.exists() else []

            cls._heading(doc,chapter_title,2)
            cls._heading(doc,"1.沿线设施技术状况TCI",3)

            cls._heading(doc,"2.标线、护栏自动化检测",3)
            cls._heading(doc,"（1）标线逆反射亮度系数",4)
            segment_sort_key=lambda row: tuple(str(row.get(field) or "") for field in ("route","direction","segment"))
            marking_summary=GuangdongStatistics.marking_summary(all_mark)
            cls._body(doc,f"{category}标线检测共获得{marking_summary['valid_count']:,}个有效计算单元，平均逆反射亮度系数为{cls._fmt(marking_summary['average'])}，合格单元{marking_summary['qualified_count']:,}个，合格率为{cls._pct(marking_summary['qualified_rate'])}。" if all_mark else f"{category}未读取到有效标线逆反射数据。")
            pair_summary=sorted(GuangdongStatistics.marking_segment_pair_summary(all_mark),key=segment_sort_key)
            def _pair_cells(row, pos):
                return [row.get(f"{pos}_valid_count"), cls._fmt(row.get(f"{pos}_average")), cls._pct(row.get(f"{pos}_qualified_rate"))]
            pair_rows=[[row.get("route"),row.get("direction"),cls._segment_text(row.get("segment")),*[cell for pos in sorted(row.get("_side_names") or {}) for cell in _pair_cells(row,pos)]] for row in pair_summary]
            side_labels=sorted({tuple(sorted((row.get("_side_names") or {}).items())) for row in pair_summary})
            labels=side_labels[0] if len(side_labels)==1 else tuple(sorted(marking_side_names(("标线2","标线3")).items()))
            pair_header=["路线","方向","检测区段"]+[x for _,name in labels for x in (f"{name}单元数",f"{name}均值",f"{name}合格率")]
            _table(pair_header,pair_rows,"标线逆反射区段汇总表")
            marking_images=_images("marking")
            for idx,row in enumerate(pair_summary,1):
                seg_text=cls._segment_text(row.get("segment"))
                cls._heading(doc,f"{cls._alpha_label(idx)} {seg_text}",5)
                cls._body(doc,cls._marking_segment_sentence(row))
                names=row.get("_side_names") or {}
                ordered=sorted(names) or ["标线2","标线3"]
                _table(["检测区段",*[x for pos in ordered for x in (f"{names.get(pos,pos)}单元数",f"{names.get(pos,pos)}均值",f"{names.get(pos,pos)}合格率")]],[[seg_text,*[x for pos in ordered for x in _pair_cells(row,pos)]]],f"{seg_text}标线逆反射统计表")
                seg=str(row.get("segment") or "")
                for pos in ordered:
                    for img in [p for p in marking_images if seg.replace("/","_").replace("\\","_").replace(" ","_") in Path(p).name and pos in Path(p).name]:
                        try: _figure(img,f"{seg_text}{names.get(pos,pos)}逆反射亮度系数检测结果")
                        except Exception: pass

            cls._heading(doc,"（2）护栏中心高度",4)
            total_height=len(all_height)
            if total_height:
                hs=GuangdongStatistics.height_summary(all_height)
                parts=[]
                for kind in ("二波","三波"):
                    item=hs[kind]
                    if item["valid_count"]: parts.append(f"{kind}护栏{item['valid_count']:,}个有效点，平均高度{cls._fmt(item['average'])} mm，合格率{cls._pct(item['qualified_rate'])}")
                cls._body(doc,f"{category}波形梁护栏中心高度有效检测点共{total_height:,}个，其中"+"；".join(parts)+"。")
            else: cls._body(doc,f"{category}未读取到有效护栏中心高度数据。")
            height_summary_rows=sorted(GuangdongStatistics.height_segment_summary(all_height),key=lambda row:(*segment_sort_key(row),str(row.get("guardrail_type") or "")))
            note_counts=cls._guardrail_note_counts([r for r in bundle.get("notes",[]) if r.get("category")==category])
            def _note_counts_for(route_name,direction_name,segment_name):
                segment_name=str(segment_name)
                for key in ((route_name,direction_name,segment_name),(route_name,direction_name,segment_name.removesuffix("段")),(route_name,direction_name,segment_name+"段")):
                    if key in note_counts: return note_counts[key]
                return None
            def _note_for(route_name,direction_name,segment_name):
                return cls._guardrail_note_sentence(_note_counts_for(route_name,direction_name,segment_name))
            height_rows=[[row.get("route"),row.get("direction"),cls._segment_text(row.get("segment")),_guardrail_type(row.get("guardrail_type")),row.get("valid_count"),cls._fmt(row.get("average")),row.get("qualified_count"),cls._pct(row.get("qualified_rate")),row.get("over_10cm_count")] for row in height_summary_rows]
            _table(["路线","方向","检测区段","护栏类型","有效点数","平均高度（mm）","合格点数","合格率","偏差超10cm点数"],height_rows,"护栏中心高度区段汇总表")
            height_images=_images("height")
            segment_key=lambda key:(key[0],key[1],str(key[2]).removesuffix("段"))
            height_groups={}
            for row in height_summary_rows:
                height_groups.setdefault(segment_key((row.get("route"),row.get("direction"),row.get("segment"))),[]).append(row)
            for note_key,counts in sorted(note_counts.items()):
                height_groups.setdefault(segment_key(note_key),[])
            for idx,(key,rows) in enumerate(sorted(height_groups.items(),key=lambda item:tuple(str(x) for x in item[0])),1):
                route_name,direction_name,segment_name=key
                seg_text=cls._segment_text(segment_name)
                note_sentence=_note_for(route_name,direction_name,segment_name)
                cls._heading(doc,f"{cls._alpha_label(idx)} {seg_text}",5)
                total=sum(int(row.get("valid_count") or 0) for row in rows)
                if total > 0:
                    cls._body(doc,cls._height_segment_sentence(rows))
                elif note_sentence:
                    cls._body(doc,cls._guardrail_no_valid_point_sentence(_note_counts_for(*key)))
                else:
                    cls._body(doc,cls._height_segment_sentence(rows))
                if rows:
                    _table(["检测区段","护栏类型","有效点数","平均高度（mm）","合格点数","合格率","偏差超10cm点数"],[[seg_text,_guardrail_type(row.get("guardrail_type")),row.get("valid_count"),cls._fmt(row.get("average")),row.get("qualified_count"),cls._pct(row.get("qualified_rate")),row.get("over_10cm_count")] for row in rows],f"{seg_text}护栏中心高度统计表")
                safe_seg=str(segment_name).replace("/","_").replace("\\","_").replace(" ","_")
                if not note_sentence or rows:
                    for kind in ("二波","三波"):
                        for img in [p for p in height_images if safe_seg in Path(p).name and kind in Path(p).name and "_line." in p]:
                            try: _figure(img,f"{seg_text}{kind}护栏中心高度检测结果")
                            except Exception: pass
                        for img in [p for p in height_images if safe_seg in Path(p).name and kind in Path(p).name and "_pie." in p]:
                            try: _figure(img,f"{seg_text}{kind}护栏中心高度结果分布")
                            except Exception: pass

            cls._heading(doc,"（3）螺栓安装情况",4)
            bs=GuangdongStatistics.bolt_summary(all_bolt)
            cls._body(doc,f"{category}共识别现有拼接螺栓{int(bs['splice']):,}颗、连接螺栓{int(bs['connection']):,}颗，检出缺失螺栓{int(bs['missing_total']):,}颗，螺栓缺失率为{cls._pct(bs['missing_rate'])}。" if all_bolt else f"{category}未读取到有效波形梁护栏螺栓数据。")
            bolt_summary_rows=sorted(GuangdongStatistics.bolt_segment_summary(all_bolt),key=segment_sort_key)
            bolt_existing={(r.get("route"),r.get("direction"),str(r.get("segment") or "").removesuffix("段")) for r in bolt_summary_rows}
            bolt_note_keys={(k[0],k[1],str(k[2]).removesuffix("段")) for k in note_counts if (k[0],k[1],str(k[2]).removesuffix("段")) not in bolt_existing}
            bolt_rows=[[row.get("route"),row.get("direction"),cls._segment_text(row.get("segment")),int(row.get("splice",0) or 0),int(row.get("connection",0) or 0),int(row.get("missing_total",0) or 0),cls._pct(row.get("missing_rate"))] for row in bolt_summary_rows]
            _table(["路线","方向","检测区段","拼接螺栓（颗）","连接螺栓（颗）","缺失数量（颗）","缺失率"],bolt_rows,"护栏螺栓缺失区段汇总表")
            bolt_images=_images("bolt")
            bolt_groups={(row.get("route"),row.get("direction"),str(row.get("segment") or "").removesuffix("段")):row for row in bolt_summary_rows}
            for note_key in bolt_note_keys: bolt_groups[note_key]=None
            for idx,(bolt_key,row) in enumerate(sorted(bolt_groups.items(),key=lambda item:tuple(str(x or "") for x in item[0])),1):
                if row is None:
                    seg_text=cls._segment_text(bolt_key[2])
                    cls._heading(doc,f"{cls._alpha_label(idx)} {seg_text}",5)
                    cls._body(doc,cls._guardrail_no_valid_point_sentence(_note_counts_for(*bolt_key)))
                    continue
                seg_text=cls._segment_text(row.get("segment"))
                cls._heading(doc,f"{cls._alpha_label(idx)} {seg_text}",5)
                cls._body(doc,cls._bolt_segment_sentence(row))
                _table(["路线","方向","检测区段","拼接螺栓（颗）","连接螺栓（颗）","缺失数量（颗）","缺失率"],[[row.get("route"),row.get("direction"),seg_text,int(row.get("splice",0) or 0),int(row.get("connection",0) or 0),int(row.get("missing_total",0) or 0),cls._pct(row.get("missing_rate"))]],f"{seg_text}护栏螺栓缺失统计表")
                for img in [p for p in bolt_images if str(row.get("segment") or "").replace("/","_").replace("\\","_").replace(" ","_") in Path(p).name and "_bar." in p]:
                    try: _figure(img,f"{seg_text}护栏螺栓缺失检测结果")
                    except Exception: pass
            cls._heading(doc,"（4）人工复核对比情况",4)
            cls._comparison_table(doc,all_detail,thresholds,table_adder=_table)

        cls._heading(doc,"（三）典型路段及成因分析",2)
        weak=bundle.get("weak_segments",[])
        if not weak: cls._body(doc,"未识别到满足连续3 km标线不合格、护栏中心高度偏差超过10 cm或螺栓缺失率超过5%的典型薄弱路段。")
        else: _table(["类型","路线","方向","检测区段","原因"],[[r.get("type"),r.get("route"),r.get("direction"),r.get("segment"),r.get("reason")] for r in weak],"典型薄弱路段及成因表")
        cls._heading(doc,"（四）交安工作建议",2)
        cls._heading(doc,"1.交安薄弱路段处治建议",3)
        cls._body(doc,"（1）对标线逆反射性能连续3 km不合格路段开展现场复核，对确认存在磨损、污染或逆反射性能不足的标线进行清理和重新施划，并在完工后复测。")
        cls._body(doc,"（2）对波形梁护栏中心高度偏差超过10 cm的区段核查路面加铺、路缘石、沉陷和立柱埋深等影响因素，结合护栏类型实施抬升、调整或更换。")
        cls._body(doc,"（3）对螺栓缺失率超过5%的区段优先补齐缺失螺栓，同步检查梁板搭接、连接件、防阻块、立柱和端头的牢固性。")
        cls._heading(doc,"2.交安养护管理建议",3)
        cls._body(doc,"建立日常巡查、定期检测和专项排查相结合的预防性养护机制，统一路线、方向、桩号和区段编码，完善检测、复核、处治、复测和销号闭环。")
        cls._indent_existing_body(doc)
        cls._format_all_run_fonts(doc)
        folder=Path(output_dir)/city; folder.mkdir(parents=True,exist_ok=True); path=folder/f"{city}在役公路技术状况检测评价报告第五部分.docx"
        try: doc.save(path)
        except PermissionError as exc: raise PermissionError(f"Word文件被占用：{path}") from exc
        return path



# ==================== 广东项目图表生成 ====================

def _gd_height_charts(rows, route, direction, chart_dir, prefix=""):
    """护栏中心高度：每个区段、每种护栏类型分别绘制折线图与结果分布饼图。"""
    images = {}
    safe_route = route.replace("/", "_").replace("\\", "_")
    safe_dir = direction.replace("/", "_")
    for segment in sorted({str(r.get("segment", "")) for r in rows}):
        seg_rows = [r for r in rows if str(r.get("segment", "")) == segment]
        safe_seg = segment.replace("/", "_").replace("\\", "_").replace(" ", "_") or "segment"
        for kind in ("二波", "三波"):
            kind_rows = sorted(
                [r for r in seg_rows if _guardrail_type(r.get("guardrail_type")) == kind
                 and _float(r.get("height")) is not None and r.get("station_m") is not None],
                key=lambda r: r["station_m"],
            )
            if not kind_rows:
                continue
            heights = [float(r["height"]) for r in kind_rows]
            x = list(range(len(kind_rows)))
            seg_text = str(segment or "") + ("段" if segment and not str(segment).endswith("段") else "")

            line_path = chart_dir / f"{prefix}height_{safe_route}_{safe_dir}_{safe_seg}_{kind}_line.png"
            fig, ax = plt.subplots(figsize=(13 / 2.54, 8 / 2.54), dpi=180)
            ax.plot(x, heights, color="#4472C4", linewidth=1, label="梁板中心高度(mm)")
            standards = (580, 620) if kind == "二波" else (677, 717)
            for std, color in zip(standards, ("#ED7D31", "#A5A5A5")):
                ax.plot(x, [std] * len(x), color=color, linewidth=2, label=f"标准值（{std}mm）")
            ax.set_ylim(300, 850)
            ax.grid(True, alpha=0.25)
            tick_count = min(10, len(x))
            ticks = sorted(set(int(i * (len(x) - 1) / (tick_count - 1)) for i in range(tick_count))) if tick_count > 1 else [0]
            ax.set_xticks(ticks)
            ax.set_xticklabels([format_station(kind_rows[i]["station_m"]) for i in ticks], rotation=30, ha="right", fontsize=7)
            ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.22), ncol=3, frameon=False)
            fig.subplots_adjust(left=0.10, right=0.98, top=0.92, bottom=0.30)
            fig.savefig(line_path, transparent=False, bbox_inches="tight", pad_inches=0.15)
            plt.close(fig)

            pie_path = chart_dir / f"{prefix}height_{safe_route}_{safe_dir}_{safe_seg}_{kind}_pie.png"
            limits = (560, 580, 620, 640) if kind == "二波" else (657, 677, 717, 737)
            labels = [
                f"h＜{limits[0]}", f"{limits[0]}≤h＜{limits[1]}",
                f"{limits[1]}≤h≤{limits[2]}", f"{limits[2]}＜h≤{limits[3]}", f"h＞{limits[3]}",
            ]
            bins = [0] * 5
            for h in heights:
                if h < limits[0]: bins[0] += 1
                elif h < limits[1]: bins[1] += 1
                elif h <= limits[2]: bins[2] += 1
                elif h <= limits[3]: bins[3] += 1
                else: bins[4] += 1
            total = len(heights)
            pcts = [b * 100 / total if total else 0 for b in bins]
            nonzero = [(lbl, v, f"#{PIE_COLORS[i]}") for i, (lbl, v) in enumerate(zip(labels, pcts)) if v > 0]
            if nonzero:
                fig, ax = plt.subplots(figsize=(14 / 2.54, 8.5 / 2.54), dpi=180)
                wedges, _, _ = ax.pie(
                    [v for _, v, _ in nonzero], colors=[c for _, _, c in nonzero],
                    autopct=lambda pct: f"{pct:.2f}%" if pct > 0 else "", pctdistance=1.15,
                    textprops={"fontsize": 8},
                )
                ax.legend(wedges, [lbl for lbl, _, _ in nonzero], loc="center left", bbox_to_anchor=(1.0, 0.5), frameon=False)
                fig.subplots_adjust(left=0.02, right=0.76, top=0.88, bottom=0.05)
                fig.savefig(pie_path, transparent=False, bbox_inches="tight", pad_inches=0.15)
                plt.close(fig)
                images[f"height_{safe_route}_{safe_dir}_{safe_seg}_{kind}"] = {"line": line_path, "pie": pie_path}
            else:
                images[f"height_{safe_route}_{safe_dir}_{safe_seg}_{kind}"] = {"line": line_path}
    return images

def _gd_bolt_charts(rows, route, direction, chart_dir, prefix=""):
    """护栏螺栓缺失：每个区段绘制一组柱状图（拼接/连接/缺失数量）。"""
    images = {}
    safe_route = route.replace("/", "_").replace("\\", "_")
    safe_dir = direction.replace("/", "_")
    segments = {}
    for r in rows:
        segments.setdefault(str(r.get("segment", "")), []).append(r)
    for seg in sorted(segments.keys()):
        seg_rows = segments[seg]
        s = sum(float(r.get("splice", 0) or 0) for r in seg_rows)
        c = sum(float(r.get("connection", 0) or 0) for r in seg_rows)
        m = sum(float(r.get("splice_missing", 0) or 0) + float(r.get("connection_missing", 0) or 0) for r in seg_rows)
        safe_seg = seg.replace("/", "_").replace("\\", "_").replace(" ", "_") or "segment"
        seg_text = str(seg or "") + ("段" if seg and not str(seg).endswith("段") else "")
        bar_path = chart_dir / f"{prefix}bolt_{safe_route}_{safe_dir}_{safe_seg}_bar.png"
        fig, ax = plt.subplots(figsize=(13 / 2.54, 8 / 2.54), dpi=180)
        category_vals = [int(round(s)), int(round(c)), int(round(m))]
        bar_labels = ["拼接螺栓", "连接螺栓", "缺失数量"]
        bar_colors = ["#4472C4", "#ED7D31", "#FF0000"]
        ax.bar(range(3), category_vals, 0.5, color=bar_colors)
        ax.set_xticks(range(3))
        ax.set_xticklabels(bar_labels, fontsize=9)
        ax.set_ylabel("螺栓数量（颗）")
        for i, v in enumerate(category_vals):
            ax.text(i, v, str(v), ha="center", va="bottom", fontsize=9)
        ax.grid(True, axis="y", alpha=0.25)
        fig.subplots_adjust(left=0.10, right=0.98, top=0.92, bottom=0.12)
        fig.savefig(bar_path, transparent=False, bbox_inches="tight", pad_inches=0.15)
        plt.close(fig)
        images[f"bolt_{safe_route}_{safe_dir}_{safe_seg}"] = {"bar": bar_path}
    return images

def _gd_marking_charts(rows, route, direction, chart_dir, prefix=""):
    """广东标线逆反射：每个区段的标线2/3分别折线图，并同时叠加80、50两条合格线。"""
    images = {}
    groups = {}
    for r in rows:
        key = (str(r.get("segment", "")), str(r.get("marking_position", "")))
        groups.setdefault(key, []).append(r)

    for (segment, pos), pos_rows in groups.items():
        pos_rows = sorted(
            [r for r in pos_rows if _float(r.get("value")) is not None and r.get("station_m") is not None],
            key=lambda r: r["station_m"],
        )
        if len(pos_rows) < 2:
            continue
        values = [float(r["value"]) for r in pos_rows]
        x = list(range(len(pos_rows)))

        safe_route = route.replace("/", "_").replace("\\", "_")
        safe_dir = direction.replace("/", "_")
        safe_seg = segment.replace("/", "_").replace("\\", "_").replace(" ", "_") or "segment"
        safe_pos = pos.replace("/", "_").replace(" ", "_") or "all"
        line_path = chart_dir / f"{prefix}marking_{safe_route}_{safe_dir}_{safe_seg}_{safe_pos}_line.png"
        fig, ax = plt.subplots(figsize=(13 / 2.54, 8 / 2.54), dpi=180)
        ax.plot(x, values, color="#4472C4", linewidth=1, label="逆反射亮度系数")
        for target, label, color in ((80, "白色标线合格线（80）", "#A5A5A5"), (50, "黄色标线合格线（50）", "#ED7D31")):
            ax.plot(x, [target] * len(x), color=color, linewidth=2, linestyle="--", label=label)
        ax.grid(True, alpha=0.25)
        tick_count = min(10, len(x))
        ticks = sorted(set(int(i * (len(x) - 1) / (tick_count - 1)) for i in range(tick_count))) if tick_count > 1 else [0]
        ax.set_xticks(ticks)
        ax.set_xticklabels([format_station(pos_rows[i]["station_m"]) for i in ticks], rotation=30, ha="right", fontsize=7)
        ax.set_ylabel("逆反射亮度系数")
        ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.22), ncol=3, frameon=False)
        fig.subplots_adjust(left=0.10, right=0.98, top=0.92, bottom=0.30)
        fig.savefig(line_path, transparent=False, bbox_inches="tight", pad_inches=0.15)
        plt.close(fig)

        images[f"marking_{safe_route}_{safe_dir}_{safe_seg}_{safe_pos}"] = {"line": line_path}
    return images

def guangdong_report_images(bundle, output_dir, log=lambda _: None):
    """为广东项目城市 bundle 生成护栏高度、螺栓缺失、标线逆反射图表。"""
    city = _safe_city_component(bundle["city"])
    chart_dir = Path(output_dir) / city / "charts"
    chart_dir.mkdir(parents=True, exist_ok=True)

    plt.rcParams["font.sans-serif"] = ["SimSun", "Microsoft YaHei", "Arial Unicode MS"]
    plt.rcParams["axes.unicode_minus"] = False

    all_images = {}
    for kind in ("height", "bolt", "marking"):
        rows = bundle.get(kind, [])
        if not rows:
            continue
        groups = {}
        for row in rows:
            key = (str(row.get("route", "")), str(row.get("direction", "")))
            groups.setdefault(key, []).append(row)
        for (route, direction), group_rows in groups.items():
            category = group_rows[0].get("category") or ""
            prefix = f"{city}_{category}_" if category else ""
            try:
                if kind == "height":
                    all_images.update(_gd_height_charts(group_rows, route, direction, chart_dir, prefix))
                elif kind == "bolt":
                    all_images.update(_gd_bolt_charts(group_rows, route, direction, chart_dir, prefix))
                elif kind == "marking":
                    all_images.update(_gd_marking_charts(group_rows, route, direction, chart_dir, prefix))
            except Exception as exc:
                log(f"图表生成警告（{kind}/{route}/{direction}）：{exc}")

    if all_images:
        log(f"{city}图表已生成：{len(all_images)}组，保存至{chart_dir}")
    return all_images


def marking_side_names(positions):
    """区段内两条标线按序号命名：序号小的显示为左侧标线，大的为右侧标线。

    兼容「标线2/标线3」「标线1/标线2」等编号形态；仅返回实际出现的最多两条。
    """
    ordered = sorted({str(p) for p in positions if str(p)})
    return {pos: name for pos, name in zip(ordered[:2], ("左侧标线", "右侧标线"))}


def _side_display(pos):
    """单个标线位置编号的显示名（图表标题用）。"""
    return marking_side_names([pos]).get(str(pos), str(pos))


def write_guangdong_chart_workbook(bundle, output_dir, log=lambda _: None):
    """生成专门的图表统计工作簿：护栏中心高度、护栏螺栓缺失、标线逆反射各占一个工作表。

    图表为可编辑的原生Excel图表；配色、线条粗细、图例位置与标题文字均与Word文档插图一致。
    """
    city = _safe_city_component(bundle["city"])
    folder = Path(output_dir) / city; folder.mkdir(parents=True, exist_ok=True)
    path = folder / f"{city}交安设施统计图表.xlsx"
    thin = Side(style="thin", color="B7B7B7")

    def font_axes(chart):
        for axis in (chart.x_axis, chart.y_axis):
            try:
                axis.txPr = RichText(p=[Paragraph(
                    pPr=ParagraphProperties(defRPr=CharacterProperties(
                        latin=DrawingFont(typeface="Times New Roman"), ea=DrawingFont(typeface="宋体"), sz=700)),
                    endParaRPr=CharacterProperties(sz=700))])
            except Exception:
                pass

    def line_style(series, color, width_emu, dash=None):
        series.graphicalProperties.line.solidFill = color
        series.graphicalProperties.line.width = width_emu
        if dash:
            series.graphicalProperties.line.dashStyle = dash
        series.smooth = False

    def header_cells(ws, row, headers, start_col=1):
        for offset, text in enumerate(headers):
            cell = ws.cell(row, start_col + offset, text)
            cell.font = Font(name="宋体", size=10, bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def seg_text(segment):
        return str(segment) + ("" if str(segment).endswith("段") or not str(segment) else "段")

    wb = Workbook(); wb.remove(wb.active)
    counts = {"高度折线": 0, "高度分布": 0, "螺栓柱状": 0, "标线折线": 0}

    # ---- 工作表1：护栏中心高度（逐区段、逐类型折线图+分布饼图，与Word插图同款） ----
    ws = wb.create_sheet("护栏中心高度")
    groups = {}
    for r in bundle.get("height", []):
        kind = _guardrail_type(r.get("guardrail_type"))
        value = _float(r.get("height"))
        if kind not in ("二波", "三波") or value is None or r.get("station_m") is None:
            continue
        groups.setdefault((str(r.get("segment") or ""), kind), []).append((r["station_m"], value))
    line_slot = pie_slot = 2
    for (segment, kind), points in sorted(groups.items()):
        points.sort(key=lambda item: item[0])
        if not points:
            continue
        standards = (580, 620) if kind == "二波" else (677, 717)
        limits = (560, 580, 620, 640) if kind == "二波" else (657, 677, 717, 737)
        header_row = line_slot
        header_cells(ws, header_row, [
            "桩号", "梁板中心高度(mm)", f"标准值（{standards[0]}mm）", f"标准值（{standards[1]}mm）",
            "高度区间", "点数",
        ])
        bin_labels = [f"h＜{limits[0]}", f"{limits[0]}≤h＜{limits[1]}", f"{limits[1]}≤h≤{limits[2]}",
                      f"{limits[2]}＜h≤{limits[3]}", f"h＞{limits[3]}"]
        bins = [0] * 5
        for index, (station, value) in enumerate(points, 1):
            row = header_row + index
            ws.cell(row, 1, format_station(station))
            ws.cell(row, 2, value)
            ws.cell(row, 3, standards[0])
            ws.cell(row, 4, standards[1])
            if value < limits[0]: bins[0] += 1
            elif value < limits[1]: bins[1] += 1
            elif value <= limits[2]: bins[2] += 1
            elif value <= limits[3]: bins[3] += 1
            else: bins[4] += 1
        for index, (label, count) in enumerate(zip(bin_labels, bins), 1):
            ws.cell(header_row + index, 5, label)
            ws.cell(header_row + index, 6, count)
        last = header_row + len(points)

        chart = LineChart(); chart.visible_cells_only = False
        chart.add_data(Reference(ws, min_col=2, max_col=4, min_row=header_row, max_row=last), titles_from_data=True)
        chart.set_categories(Reference(ws, min_col=1, min_row=header_row + 1, max_row=last))
        chart.height, chart.width = 8, 13
        chart.y_axis.scaling.min, chart.y_axis.scaling.max, chart.y_axis.majorUnit = 300, 850, 50
        chart.x_axis.delete = chart.y_axis.delete = False
        chart.legend.position = "b"
        line_style(chart.series[0], "4472C4", 12700)
        line_style(chart.series[1], "ED7D31", 25400)
        line_style(chart.series[2], "A5A5A5", 25400)
        ws.add_chart(chart, f"H{line_slot}")
        counts["高度折线"] += 1; line_slot += 17

        pie = PieChart()
        pie.add_data(Reference(ws, min_col=6, min_row=header_row, max_row=header_row + 5), titles_from_data=True)
        pie.set_categories(Reference(ws, min_col=5, min_row=header_row + 1, max_row=header_row + 5))
        pie.height, pie.width = 8.5, 14
        pie.legend.position = "r"
        pie.dataLabels = DataLabelList()
        pie.dataLabels.showVal, pie.dataLabels.showPercent = False, True
        pie.dataLabels.numFmt = "0.00%"
        pie.dataLabels.showCatName = pie.dataLabels.showSerName = pie.dataLabels.showLegendKey = False
        pie.series[0].data_points = [DataPoint(idx=i, spPr=GraphicalProperties(solidFill=color)) for i, color in enumerate(PIE_COLORS)]
        zero_labels = []
        for index, count in enumerate(bins):
            if not count:
                label = DataLabel(idx=index); label.delete = True; zero_labels.append(label)
        if zero_labels:
            pie.dataLabels.dLbl = zero_labels
        ws.add_chart(pie, f"S{pie_slot}")
        counts["高度分布"] += 1; pie_slot = max(pie_slot + 17, header_row + len(points) + 3)

    # ---- 工作表2：护栏螺栓缺失（区段汇总表+逐区段柱状图，柱色与Word一致） ----
    ws = wb.create_sheet("护栏螺栓缺失")
    headers = ["路线", "方向", "检测区段", "拼接螺栓（颗）", "连接螺栓（颗）", "缺失数量（颗）"]
    header_cells(ws, 1, headers)
    summary_rows = {}
    for r in bundle.get("bolt", []):
        key = (str(r.get("route") or ""), str(r.get("direction") or ""), str(r.get("segment") or ""))
        item = summary_rows.setdefault(key, {"splice": 0, "connection": 0, "missing": 0})
        item["splice"] += float(r.get("splice", 0) or 0) + 0
        item["connection"] += float(r.get("connection", 0) or 0)
        item["missing"] += float(r.get("splice_missing", 0) or 0) + float(r.get("connection_missing", 0) or 0)
    for index, (key, item) in enumerate(sorted(summary_rows.items()), 2):
        ws.cell(index, 1, key[0]); ws.cell(index, 2, key[1]); ws.cell(index, 3, key[2])
        ws.cell(index, 4, int(round(item["splice"]))); ws.cell(index, 5, int(round(item["connection"]))); ws.cell(index, 6, int(round(item["missing"])))
    for column, width in zip("ABCDEF", (10, 10, 26, 15, 15, 15)):
        ws.column_dimensions[column].width = width
    bar_slot = 2
    for index in range(2, len(summary_rows) + 2):
        chart = BarChart(); chart.type = "col"
        chart.add_data(Reference(ws, min_col=4, max_col=6, min_row=index, max_row=index), from_rows=True)
        chart.set_categories(Reference(ws, min_col=4, max_col=6, min_row=1))
        chart.height, chart.width = 8, 13
        chart.legend = None
        chart.x_axis.delete = chart.y_axis.delete = False
        chart.y_axis.title = "螺栓数量（颗）"
        chart.gapWidth = 80
        chart.series[0].data_points = [DataPoint(idx=i, spPr=GraphicalProperties(solidFill=color))
                                       for i, color in enumerate(("4472C4", "ED7D31", "FF0000"))]
        chart.dataLabels = DataLabelList()
        chart.dataLabels.showVal = True
        chart.dataLabels.dLblPos = "outEnd"
        ws.add_chart(chart, f"H{bar_slot}")
        counts["螺栓柱状"] += 1; bar_slot += 17

    # ---- 工作表3：标线逆反射（逐20m明细折线图，叠加80/50虚线合格线） ----
    ws = wb.create_sheet("标线逆反射")
    marking_groups = {}
    for r in bundle.get("marking", []):
        value = _float(r.get("value"))
        if value is None or r.get("station_m") is None:
            continue
        key = (str(r.get("segment") or ""), str(r.get("marking_position") or ""))
        marking_groups.setdefault(key, []).append((r["station_m"], value))
    slot = 2
    for (segment, pos), points in sorted(marking_groups.items()):
        points.sort(key=lambda item: item[0])
        if len(points) < 2:
            continue
        header_row = slot
        header_cells(ws, header_row, ["桩号", "逆反射亮度系数", "白色标线合格线（80）", "黄色标线合格线（50）"])
        for index, (station, value) in enumerate(points, 1):
            row = header_row + index
            ws.cell(row, 1, format_station(station))
            ws.cell(row, 2, value)
            ws.cell(row, 3, 80)
            ws.cell(row, 4, 50)
        last = header_row + len(points)
        chart = LineChart(); chart.visible_cells_only = False
        chart.add_data(Reference(ws, min_col=2, max_col=4, min_row=header_row, max_row=last), titles_from_data=True)
        chart.set_categories(Reference(ws, min_col=1, min_row=header_row + 1, max_row=last))
        chart.height, chart.width = 8, 13
        chart.x_axis.delete = chart.y_axis.delete = False
        chart.y_axis.title = "逆反射亮度系数"
        chart.legend.position = "b"
        line_style(chart.series[0], "4472C4", 12700)
        line_style(chart.series[1], "A5A5A5", 25400, dash="dash")
        line_style(chart.series[2], "ED7D31", 25400, dash="dash")
        ws.add_chart(chart, f"F{slot}")
        counts["标线折线"] += 1; slot = max(slot + 17, last + 2)

    for column, width in (("A", 14), ("B", 18), ("C", 18), ("D", 18), ("E", 16), ("F", 10)):
        ws.column_dimensions[column].width = width
        wb["护栏中心高度"].column_dimensions[column].width = width
    try:
        wb.save(path)
    except PermissionError as exc:
        raise PermissionError(f"Excel文件被占用：{path}") from exc
    log(f"图表统计工作簿已生成：{path}（" + "，".join(f"{name}{count}个" for name, count in counts.items()) + "）")
    return path


class GuangdongBatchRunner:
    @staticmethod
    def add_weak_segments(bundle):
        weak=[]
        for row in GuangdongStatistics.continuous_marking_weak(bundle.get("marking",[])):
            weak.append({"type":"标线连续3km不合格",**row,"segment":f"{format_station(row['start_m'])}~{format_station(row['end_m'])}","reason":"同一路线、方向和标线位置连续不合格长度达到3 km"})
        for row in GuangdongStatistics.height_segment_summary(bundle.get("height",[])):
            if (row.get("over_10cm_count") or 0) > 0:
                weak.append({"type":"护栏高度偏差超10cm","route":row.get("route"),"direction":row.get("direction"),"segment":row.get("segment"),"reason":f"{_guardrail_type(row.get('guardrail_type'))}护栏偏差超10 cm点数{row.get('over_10cm_count')}个"})
        for row in GuangdongStatistics.bolt_segment_summary(bundle.get("bolt",[])):
            if (row.get("missing_rate") or 0) > 0.05:
                weak.append({"type":"螺栓缺失率超5%","route":row.get("route"),"direction":row.get("direction"),"segment":row.get("segment"),"reason":f"螺栓缺失率{row.get('missing_rate'):.2%}"})
        bundle["weak_segments"]=weak
        return bundle

    @staticmethod
    def run_bundles(bundles,output_dir,template,thresholds,log=lambda _x:None):
        result={"success":[],"failed":{},"warnings":[]}
        for city,bundle in bundles.items():
            try:
                if bundle.get("_error"):
                    raise ValueError(bundle["_error"])
                log(f"生成{city}数据图表")
                guangdong_report_images(bundle, output_dir, log)
                log(f"生成{city}第五章Word")
                GuangdongChapterWriter.write(city,bundle,output_dir,template,thresholds)
                log(f"生成{city}图表统计工作簿")
                write_guangdong_chart_workbook(bundle,output_dir,log)
                charts_dir=Path(output_dir)/_safe_city_component(city)/"charts"
                if charts_dir.is_dir():
                    shutil.rmtree(charts_dir,ignore_errors=True)
                    log(f"已清理临时图片文件夹：{charts_dir}")
                result["success"].append(city)
            except Exception as exc:
                result["failed"][city]=str(exc); log(f"{city}失败：{exc}")
        return result

    @staticmethod
    def build_bundles(scanned,route_index,manual_records=None,thresholds=None):
        cities=sorted({r["city"] for kind in ("height","bolt","marking","notes") for r in scanned.get(kind,[])})
        bundles={}
        comparator=ManualAutoComparator(thresholds or {"marking":0,"height":0,"bolt":0})
        for city in cities:
            bundle={"city":city,"route_rows":route_index.rows(city),"issues":list(scanned.get("issues",[])),"weak_segments":[]}
            try:
                for kind in ("height","bolt","marking","notes"):
                    bundle[kind]=[dict(r) for r in scanned.get(kind,[]) if r["city"]==city]
                    for row in bundle[kind]: row["category"]=route_index.category(city,row["route"])
                detail,summary=comparator.compare([r for r in (manual_records or []) if r.get("city")==city]); bundle["comparison_detail"]=detail; bundle["comparison_summary"]=summary
                for row in detail: row["category"]=route_index.category_or_none(row.get("city"),row.get("route"))
                GuangdongBatchRunner.add_weak_segments(bundle)
            except Exception as exc:
                bundle["_error"]=str(exc)
                for kind in ("height","bolt","marking","notes"):
                    bundle.setdefault(kind,[])
                bundle.setdefault("comparison_detail",[]); bundle.setdefault("comparison_summary",{})
            bundles[city]=bundle
        return bundles


PROJECT_TEMPLATES = ("重庆项目模板", "广东项目模板")


def run_guangdong_project(config, log=lambda _x: None):
    for label,path,kind in (("项目资料",config.project_dir,"dir"),("人工自动化对比表",config.manual_xlsx,"file"),("路线分类表",config.route_xlsx,"file")):
        exists=path.is_dir() if kind=="dir" else path.is_file()
        if not exists: raise FileNotFoundError(f"{label}不存在：{path}")
    if config.marking_dir and not config.marking_dir.is_dir():
        raise FileNotFoundError(f"标线数据文件夹不存在：{config.marking_dir}")
    if config.guardrail_dir and not config.guardrail_dir.is_dir():
        raise FileNotFoundError(f"护栏数据文件夹不存在：{config.guardrail_dir}")
    route_index=RouteCategoryIndex.from_file(config.route_xlsx); log("路线分类表读取完成")

    scanned = {"height": [], "bolt": [], "marking": [], "notes": [], "issues": []}
    if config.marking_dir or config.guardrail_dir:
        marking_dir = config.marking_dir
        guardrail_dir = config.guardrail_dir
        if marking_dir and guardrail_dir and Path(marking_dir) == Path(guardrail_dir):
            log(f"扫描数据文件夹：{marking_dir}")
            partial = GuangdongInputScanner(marking_dir, route_index, log=log).scan()
            scanned["marking"].extend(partial["marking"])
            scanned["height"].extend(partial["height"])
            scanned["bolt"].extend(partial["bolt"])
            scanned["notes"].extend(partial["notes"])
            scanned["issues"].extend(partial["issues"])
        else:
            if marking_dir:
                log(f"扫描标线数据文件夹：{marking_dir}")
                partial = GuangdongInputScanner(marking_dir, route_index, log=log).scan()
                scanned["marking"].extend(partial["marking"])
                scanned["issues"].extend(partial["issues"])
            if guardrail_dir:
                log(f"扫描护栏数据文件夹：{guardrail_dir}")
                partial = GuangdongInputScanner(guardrail_dir, route_index, log=log).scan()
                scanned["height"].extend(partial["height"])
                scanned["bolt"].extend(partial["bolt"])
                scanned["notes"].extend(partial["notes"])
                scanned["issues"].extend(partial["issues"])
    else:
        scanned = GuangdongInputScanner(config.project_dir, route_index, log=log).scan()

    log(f"数据识别完成：标线{len(scanned['marking'])}、高度{len(scanned['height'])}、螺栓{len(scanned['bolt'])}")
    if not any(scanned[k] for k in ("marking","height","bolt")): raise ValueError("未识别到任何有效数据")
    manual,manual_issues=ManualAutoComparator.read_file(config.manual_xlsx); scanned["issues"].extend(manual_issues)
    bundles=GuangdongBatchRunner.build_bundles(scanned,route_index,manual,config.thresholds)
    template=resource_template_path("广东项目第五章模板.md")
    return GuangdongBatchRunner.run_bundles(bundles,config.output_dir,template,config.thresholds,log)


class ReportGeneratorApp(tk.Tk):
    """双模板统一 GUI；重庆逻辑调用原基线，广东逻辑调用批量运行器。"""
    def __init__(self):
        super().__init__(); self.title(PROGRAM_NAME); self.geometry("980x860"); self.minsize(880,720)
        self.project_template=tk.StringVar(value=PROJECT_TEMPLATES[0]); self.queue=Queue(); self.running=False
        keys=("cq_project","cq_summary","cq_detail","cq_disease","cq_output","gd_project","gd_marking_dir","gd_guardrail_dir","gd_manual","gd_route","gd_output","gd_marking","gd_height","gd_bolt")
        self.vars={k:tk.StringVar() for k in keys}
        self.vars["gd_marking"].set("7"); self.vars["gd_height"].set("5"); self.vars["gd_bolt"].set("5")
        self._build(); self.after(100,self._poll)

    def _pick(self,key,file=False):
        path=filedialog.askopenfilename(filetypes=[("Excel","*.xlsx")]) if file else filedialog.askdirectory()
        if path:
            self.vars[key].set(path)
            if key == "gd_project":
                self._auto_detect_gd_folders(path)

    def _auto_detect_gd_folders(self, project_dir):
        """自动识别标线和护栏数据文件夹。"""
        self._append("正在自动识别标线和护栏数据文件夹...")
        try:
            marking_dir, guardrail_dir = detect_guangdong_data_folders(project_dir)
            if marking_dir: self.vars["gd_marking_dir"].set(marking_dir)
            if guardrail_dir: self.vars["gd_guardrail_dir"].set(guardrail_dir)
            self._append(f"识别完成：标线={marking_dir or '未找到'}，护栏={guardrail_dir or '未找到'}")
        except Exception as exc:
            self._append(f"自动识别失败：{exc}")

    def _row(self,parent,row,label,key,file=False):
        ttk.Label(parent,text=label,width=23).grid(row=row,column=0,sticky="w",pady=4)
        is_threshold = key in ("gd_marking","gd_height","gd_bolt")
        is_editable_path = key in ("gd_marking_dir","gd_guardrail_dir")
        state = "normal" if (is_threshold or is_editable_path) else "readonly"
        ttk.Entry(parent,textvariable=self.vars[key],state=state).grid(row=row,column=1,sticky="ew",padx=8)
        if not is_threshold: ttk.Button(parent,text="选择",command=lambda:self._pick(key,file),width=10).grid(row=row,column=2)

    def _build(self):
        root=ttk.Frame(self,padding=16); root.pack(fill="both",expand=True)
        ttk.Label(root,text=PROGRAM_NAME,font=("Microsoft YaHei UI",16,"bold")).pack(anchor="w")
        select=ttk.LabelFrame(root,text="项目模板",padding=8); select.pack(fill="x",pady=10)
        for name in PROJECT_TEMPLATES: ttk.Radiobutton(select,text=name,value=name,variable=self.project_template,command=self._switch).pack(side="left",padx=12)
        self.forms=ttk.Frame(root); self.forms.pack(fill="x")
        self.cq=ttk.LabelFrame(self.forms,text="重庆项目输入",padding=10); self.gd=ttk.LabelFrame(self.forms,text="广东项目输入",padding=10)
        for i,(label,key,file) in enumerate((("项目资料文件夹","cq_project",False),("分段汇总表","cq_summary",True),("检测明细文件夹","cq_detail",False),("病害清单文件夹","cq_disease",False),("输出文件夹","cq_output",False))): self._row(self.cq,i,label,key,file)
        for i,(label,key,file) in enumerate((("项目资料文件夹","gd_project",False),("标线数据文件夹","gd_marking_dir",False),("护栏数据文件夹","gd_guardrail_dir",False),("人工自动化对比表","gd_manual",True),("路线分类表","gd_route",True),("输出文件夹","gd_output",False),("标线一致性阈值（%）","gd_marking",False),("护栏高度一致性阈值（mm）","gd_height",False),("螺栓缺失一致性阈值（%）","gd_bolt",False))): self._row(self.gd,i,label,key,file)
        self.cq.columnconfigure(1,weight=1); self.gd.columnconfigure(1,weight=1); self._switch()
        actions=ttk.Frame(root); actions.pack(fill="x",pady=10); self.run_button=ttk.Button(actions,text="开始运行",command=self.start); self.run_button.pack(side="left"); ttk.Button(actions,text="打开输出文件夹",command=self.open_output).pack(side="left",padx=8)
        self.progress=ttk.Progressbar(root,mode="indeterminate"); self.progress.pack(fill="x")
        self.log=scrolledtext.ScrolledText(root,state="disabled",height=16); self.log.pack(fill="both",expand=True,pady=8)

    def _switch(self):
        self.cq.pack_forget(); self.gd.pack_forget(); (self.gd if self.project_template.get()==PROJECT_TEMPLATES[1] else self.cq).pack(fill="x")

    def start(self):
        if self.running:return
        try:
            if self.project_template.get()==PROJECT_TEMPLATES[1]:
                cfg=GuangdongConfig(Path(self.vars['gd_project'].get()),Path(self.vars['gd_manual'].get()),Path(self.vars['gd_route'].get()),Path(self.vars['gd_output'].get()),self.vars['gd_marking'].get(),self.vars['gd_height'].get(),self.vars['gd_bolt'].get(),marking_dir=Path(self.vars['gd_marking_dir'].get()) if self.vars['gd_marking_dir'].get() else None,guardrail_dir=Path(self.vars['gd_guardrail_dir'].get()) if self.vars['gd_guardrail_dir'].get() else None)
                worker=lambda:run_guangdong_project(cfg,lambda x:self.queue.put(("log",x)))
            else:
                base=Path(self.vars['cq_project'].get() or "."); cfg=Config(base,Path(self.vars['cq_summary'].get()),Path(self.vars['cq_detail'].get()),BUILTIN_REPORT_TEMPLATES[next(iter(BUILTIN_REPORT_TEMPLATES))],Path(self.vars['cq_output'].get() or base),Path(self.vars['cq_disease'].get()) if self.vars['cq_disease'].get() else None)
                worker=lambda:generate_statistics_and_report(cfg,lambda x:self.queue.put(("log",x)),process_height=True,process_bolts=True)
        except Exception as exc:messagebox.showerror("参数错误",str(exc));return
        self.running=True;self.run_button.configure(state="disabled");self.progress.start(10)
        def run():
            try:self.queue.put(("done",worker()))
            except Exception as exc:self.queue.put(("error",f"{exc}\n{traceback.format_exc()}"))
        Thread(target=run,daemon=True).start()

    def _poll(self):
        try:
            while True:
                kind,value=self.queue.get_nowait()
                if kind=="log":self._append(value)
                else:
                    self.running=False;self.progress.stop();self.run_button.configure(state="normal")
                    if kind=="done":self._append(f"完成：{value}");messagebox.showinfo("完成",str(value))
                    else:self._append(f"错误：{value}");messagebox.showerror("运行失败",str(value))
        except Empty:pass
        self.after(100,self._poll)

    def _append(self,text):
        self.log.configure(state="normal");self.log.insert("end",str(text)+"\n");self.log.see("end");self.log.configure(state="disabled")
    def open_output(self):
        key="gd_output" if self.project_template.get()==PROJECT_TEMPLATES[1] else "cq_output"; path=Path(self.vars[key].get() or ".")
        if path.is_dir():os.startfile(path)
        else:messagebox.showwarning("提示","输出文件夹不存在")


def launch_app():
    ReportGeneratorApp().mainloop()


if __name__ == "__main__":
    launch_app()
