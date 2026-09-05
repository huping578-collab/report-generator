from __future__ import annotations

import json
import os
import openpyxl
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn

from bridge import DesktopBridge
from backend import markdown_skeleton, minimal_docx, report_engine as engine


class FakeWindow:
    def __init__(self) -> None:
        self.scripts: list[str] = []

    def evaluate_js(self, script: str) -> None:
        self.scripts.append(script)


class MarkdownSkeletonTests(unittest.TestCase):
    def test_template_reads_toml_front_matter_and_body_blocks(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "template.md"
            path.write_text(
                "+++\n"
                "[page]\n"
                "orientation = \"landscape\"\n"
                "[body]\n"
                "size_pt = 11\n"
                "+++\n\n"
                "# 标题\n\n"
                "<!-- toc -->\n\n"
                "正文。\n",
                encoding="utf-8",
            )

            template = markdown_skeleton.read_template(path)
            self.assertEqual(markdown_skeleton.read_blocks(path), template.blocks)

        self.assertEqual(template.config["page"]["orientation"], "landscape")
        self.assertEqual(template.config["body"]["size_pt"], 11)
        self.assertEqual([block.kind for block in template.blocks], ["heading", "toc", "paragraph"])
        self.assertEqual(template.blocks[-1].text, "正文。")

    def test_template_without_front_matter_uses_legacy_defaults(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "legacy.md"
            path.write_text("# 旧模板\n\n正文。\n", encoding="utf-8")

            template = markdown_skeleton.read_template(path)
            self.assertEqual(
                markdown_skeleton.read_blocks(path),
                [
                    markdown_skeleton.Block("heading", level=1, text="旧模板"),
                    markdown_skeleton.Block("paragraph", text="正文。"),
                ],
            )

        self.assertEqual(template.config["page"]["paper"], "A4")
        self.assertEqual(template.config["page"]["orientation"], "portrait")
        self.assertEqual(template.config["body"]["east_asia"], "仿宋_GB2312")
        self.assertEqual(template.config["body"]["latin"], "Times New Roman")
        self.assertEqual(template.config["body"]["size_pt"], 10.5)
        self.assertEqual(template.config["body"]["first_line_chars"], 2)
        self.assertFalse(template.config["toc"]["enabled"])

    def test_default_config_isolated_between_template_reads(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "legacy.md"
            path.write_text("# 旧模板\n", encoding="utf-8")

            first = markdown_skeleton.read_template(path)
            first.config["body"]["size_pt"] = 99
            second = markdown_skeleton.read_template(path)

        self.assertEqual(second.config["body"]["size_pt"], 10.5)
        self.assertEqual(markdown_skeleton.DEFAULT_CONFIG["body"]["size_pt"], 10.5)

    def test_toml_multiline_string_can_contain_separator_line(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "multiline.md"
            path.write_text(
                '''+++
[toc]
enabled = false
title = """目录
+++
说明"""
+++
# 标题
''',
                encoding="utf-8",
            )

            template = markdown_skeleton.read_template(path)

        self.assertEqual(template.config["toc"]["title"], "目录\n+++\n说明")
        self.assertEqual([block.text for block in template.blocks], ["标题"])

    def test_repository_templates_declare_format_and_toc_policy(self) -> None:
        root = Path(__file__).resolve().parents[1]
        chongqing = markdown_skeleton.read_template(root / "templates" / "重庆项目报告模板.md")
        guangdong = markdown_skeleton.read_template(root / "templates" / "广东项目第五章模板.md")

        self.assertTrue(chongqing.config["toc"]["enabled"])
        self.assertTrue(any(block.kind == "toc" for block in chongqing.blocks))
        self.assertFalse(guangdong.config["toc"]["enabled"])
        self.assertFalse(any(block.kind == "toc" for block in guangdong.blocks))
        for config in (chongqing.config, guangdong.config):
            self.assertEqual(set(config), {"page", "body", "heading", "table", "caption", "toc"})

    def test_invalid_template_config_reports_path_and_field(self) -> None:
        cases = (
            ("[body]\nsize_pt = 73\n", "size_pt"),
            ("[page]\nleft_margin_cm = 11\n", "left_margin_cm"),
            ("[page]\norientation = \"diagonal\"\n", "orientation"),
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "invalid.md"
            for front_matter, field in cases:
                path.write_text(f"+++\n{front_matter}+++\n# 标题\n", encoding="utf-8")
                with self.subTest(field=field):
                    with self.assertRaises(ValueError) as raised:
                        markdown_skeleton.read_template(path)
                    self.assertIn(str(path), str(raised.exception))
                    self.assertIn(field, str(raised.exception))

    def test_toc_enabled_requires_toc_marker_but_disabled_allows_legacy_body(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "重庆模板.md"
            path.write_text("+++\n[toc]\nenabled = true\n+++\n# 标题\n", encoding="utf-8")
            with self.assertRaises(ValueError) as raised:
                markdown_skeleton.read_template(path)
            self.assertIn(str(path), str(raised.exception))
            self.assertIn("toc", str(raised.exception))

            path.write_text("+++\n[toc]\nenabled = false\n+++\n# 广东模板\n", encoding="utf-8")
            template = markdown_skeleton.read_template(path)

        self.assertEqual([block.kind for block in template.blocks], ["heading"])
        self.assertFalse(template.config["toc"]["enabled"])


class MarkdownDocxFormattingTests(unittest.TestCase):
    def _write_report(self, template_text: str) -> Path:
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        root = Path(temp_dir.name)
        template = root / "template.md"
        template.write_text(template_text, encoding="utf-8")
        output = root / "output"
        output.mkdir()
        report_config = engine.Config(
            project_dir=root,
            summary_xlsx=root / "summary.xlsx",
            detail_dir=root,
            template_docx=template,
            output_dir=output,
        )
        minimal_docx.make_report(
            report_config, [], None, [], None, [], None, [], {}, root,
            skeleton_md=template,
        )
        return output / engine.OUT_DOCX_NAME

    def test_markdown_config_controls_docx_runs_and_outline(self) -> None:
        template_text = (
            "+++\n"
            "[body]\n"
            "east_asia = \"楷体\"\n"
            "latin = \"Arial\"\n"
            "size_pt = 11\n"
            "first_line_chars = 0\n"
            "[heading.1]\n"
            "east_asia = \"黑体\"\n"
            "latin = \"Arial\"\n"
            "size_pt = 18\n"
            "outline_level = 0\n"
            "[table]\n"
            "east_asia = \"宋体\"\n"
            "latin = \"Arial\"\n"
            "size_pt = 9\n"
            "[caption]\n"
            "east_asia = \"黑体\"\n"
            "latin = \"Arial\"\n"
            "size_pt = 12\n"
            "[toc]\n"
            "enabled = true\n"
            "update_on_open = true\n"
            "+++\n"
            "# 自定义标题\n\n"
            "<!-- toc -->\n\n"
            "<!-- inject:overview -->\n"
        )
        report_path = self._write_report(template_text)
        with zipfile.ZipFile(report_path) as archive:
            document = archive.read("word/document.xml").decode("utf-8")
            settings = archive.read("word/settings.xml").decode("utf-8")
        self.assertIn('w:eastAsia="黑体"', document)
        self.assertIn('w:eastAsia="楷体"', document)
        self.assertIn('w:eastAsia="宋体"', document)
        self.assertIn('w:ascii="Arial"', document)
        self.assertIn('w:hAnsi="Arial"', document)
        self.assertIn('w:val="0"', document)
        self.assertIn('w:instrText xml:space="preserve"> TOC \\o "1-5" \\h \\z \\u </w:instrText>', document)
        self.assertIn('w:updateFields w:val="true"', settings)

    def test_disabled_toc_does_not_write_toc_field(self) -> None:
        report_path = self._write_report("+++\n[toc]\nenabled = false\n+++\n# 无目录\n")
        with zipfile.ZipFile(report_path) as archive:
            document = archive.read("word/document.xml").decode("utf-8")
        self.assertNotIn(" TOC ", document)


class GuangdongTemplateConfigTests(unittest.TestCase):
    @staticmethod
    def _bundle() -> dict:
        return {
            "marking": [],
            "height": [],
            "bolt": [],
            "notes": [],
            "comparison_detail": [],
            "weak_segments": [],
        }

    @staticmethod
    def _write_report(root: Path, template: Path, output_dir: Path | None = None) -> Path:
        return engine.GuangdongChapterWriter.write(
            "佛山市",
            GuangdongTemplateConfigTests._bundle(),
            output_dir or root / "output",
            template,
            {"marking": 7, "height": 5, "bolt": 5},
        )

    @staticmethod
    def _xml_paragraphs(document_xml: bytes):
        from xml.etree import ElementTree as ET

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = ET.fromstring(document_xml)
        return root.findall(".//w:body/w:p", ns), ns

    def test_guangdong_writer_disables_toc_and_update_fields_from_template(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template = root / "disabled-toc.md"
            template.write_text(
                "+++\n"
                "[toc]\n"
                "enabled = false\n"
                "+++\n"
                "# 五、交通安全设施技术状况检测评价情况\n\n"
                "<!-- toc -->\n",
                encoding="utf-8",
            )
            output = self._write_report(root, template)
            with zipfile.ZipFile(output) as archive:
                document = archive.read("word/document.xml").decode("utf-8")
                settings = archive.read("word/settings.xml").decode("utf-8")

        self.assertNotIn(" TOC ", document)
        self.assertNotIn("w:updateFields", settings)

    def test_guangdong_writer_uses_custom_toc_range_title_and_update_policy(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template = root / "custom-toc.md"
            template.write_text(
                "+++\n"
                "[heading.1]\n"
                "east_asia = \"目录标题字体\"\n"
                "size_pt = 18\n"
                "[toc]\n"
                "enabled = true\n"
                "title = \"自定义目录\"\n"
                "min_level = 2\n"
                "max_level = 4\n"
                "update_on_open = false\n"
                "+++\n"
                "# 五、交通安全设施技术状况检测评价情况\n\n"
                "<!-- toc -->\n",
                encoding="utf-8",
            )
            output = self._write_report(root, template)
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml")
                document = document_xml.decode("utf-8")
                settings = archive.read("word/settings.xml").decode("utf-8")

        paragraphs, ns = self._xml_paragraphs(document_xml)
        toc_titles = [
            paragraph for paragraph in paragraphs if "自定义目录" in "".join(
                node.text or "" for node in paragraph.findall(".//w:t", ns)
            )
        ]
        self.assertEqual(len(toc_titles), 1)
        if len(toc_titles) != 1:
            return
        toc_title = toc_titles[0]
        title_fonts = toc_title.find(".//w:rPr/w:rFonts", ns)
        title_size = toc_title.find(".//w:rPr/w:sz", ns)
        self.assertEqual(title_fonts.get(qn("w:eastAsia")), "目录标题字体")
        self.assertEqual(title_size.get(qn("w:val")), "36")
        self.assertIn('w:instrText xml:space="preserve"> TOC \\o "2-4" \\h \\z \\u </w:instrText>', document)
        self.assertIn(">自定义目录</w:t>", document)
        self.assertNotIn("w:updateFields", settings)

    def test_guangdong_writer_does_not_keep_template_config_after_exception(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template = root / "custom-body.md"
            template.write_text(
                "+++\n"
                "[body]\n"
                "size_pt = 17\n"
                "+++\n"
                "# 五、交通安全设施技术状况检测评价情况\n",
                encoding="utf-8",
            )
            blocked_output = root / "blocked-output"
            blocked_output.write_text("not a directory", encoding="utf-8")
            with self.assertRaises(OSError):
                self._write_report(root, template, blocked_output)

        self.assertNotIn("_template_format_config", vars(engine.GuangdongChapterWriter))
        self.assertEqual(
            engine.GuangdongChapterWriter._format_config()["body"]["size_pt"],
            markdown_skeleton.DEFAULT_CONFIG["body"]["size_pt"],
        )

    def test_guangdong_two_level_header_cells_have_center_v_align(self) -> None:
        document = Document()
        detail = [{
            "indicator": "bolt",
            "route": "G1",
            "gtype": "二波",
            "position": "左侧",
            "direction": "上行",
            "segment": "K1+000～K2+000",
            "msplice": 1,
            "mconn": 2,
            "asplice": 1,
            "aconn": 2,
            "remark": "",
        }]
        with minimal_docx.format_context(markdown_skeleton.DEFAULT_CONFIG):
            engine.GuangdongChapterWriter._comparison_table(
                document,
                detail,
                {"marking": 7, "height": 5, "bolt": 5},
            )

        table = document.tables[1]
        xml_rows = table._tbl.findall("./w:tr", table._tbl.nsmap)
        header_cells = [cell for row in xml_rows[:2] for cell in row.findall("./w:tc", table._tbl.nsmap)]
        self.assertEqual([len(row.findall("./w:tc", table._tbl.nsmap)) for row in xml_rows[:2]], [8, 10])
        for cell in header_cells:
            vertical_alignment = cell.find("./w:tcPr/w:vAlign", cell.nsmap)
            self.assertIsNotNone(vertical_alignment)
            self.assertEqual(vertical_alignment.get(qn("w:val")), "center")
        self.assertEqual(
            sum(cell.find("./w:tcPr/w:gridSpan", cell.nsmap) is not None for cell in xml_rows[0].findall("./w:tc", table._tbl.nsmap)),
            2,
        )
        self.assertEqual(
            sum(cell.find("./w:tcPr/w:vMerge", cell.nsmap) is not None for cell in header_cells),
            12,
        )

    def test_guangdong_toc_field_paragraph_has_no_body_first_line_indent(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template = root / "toc.md"
            template.write_text(
                "+++\n"
                "[toc]\n"
                "enabled = true\n"
                "title = \"目录标题\"\n"
                "+++\n"
                "# 五、交通安全设施技术状况检测评价情况\n\n"
                "<!-- toc -->\n",
                encoding="utf-8",
            )
            output = self._write_report(root, template)
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml")

        paragraphs, ns = self._xml_paragraphs(document_xml)
        field_paragraph = next(
            paragraph for paragraph in paragraphs if paragraph.find(".//w:instrText", ns) is not None
        )
        style = field_paragraph.find("./w:pPr/w:pStyle", ns)
        indent = field_paragraph.find("./w:pPr/w:ind", ns)
        self.assertIsNotNone(style)
        self.assertTrue(style.get(qn("w:val")).startswith("TOC"))
        self.assertNotEqual(None if indent is None else indent.get(qn("w:firstLine")), "420")


class SharedRulesTests(unittest.TestCase):
    def test_guardrail_height_and_bolt_rules_are_shared(self) -> None:
        self.assertEqual(engine.guardrail_type("两波护栏"), "二波")
        self.assertEqual(engine.guardrail_type("三波护栏"), "三波")
        self.assertTrue(engine.height_deviation_over_10cm("二波", 701))
        self.assertFalse(engine.height_deviation_over_10cm("三波", 797))
        self.assertAlmostEqual(engine.bolt_missing_ratio(10, 20, 3), 3 / 33)
        self.assertIsNone(engine.bolt_missing_ratio(0, 0, 0))

    def test_progress_message_round_trips_with_optional_item(self) -> None:
        message = engine.format_progress("扫描资料", 2, 5, "a.xlsx")
        self.assertEqual(message, "[progress] 扫描资料 2/5 a.xlsx")
        self.assertEqual(
            engine.parse_progress(message),
            {"stage": "扫描资料", "current": 2, "total": 5, "item": "a.xlsx"},
        )
        self.assertIsNone(engine.parse_progress("普通日志"))

    def test_make_excel_keeps_zero_denominator_bolt_rate_blank(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            config = engine.Config(root, root / "summary.xlsx", root, root / "template.md", root / "output")
            segment = {
                "county": "测试区",
                "route": "G1",
                "start": 1000.0,
                "end": 2000.0,
                "mileage": 1.0,
            }
            engine.make_excel(
                config,
                [segment],
                bolt_stats=[{
                    "segment": segment,
                    "splice": 0,
                    "connection": 0,
                    "missing": 0,
                    "rate": None,
                }],
            )
            workbook = openpyxl.load_workbook(config.out_xlsx, data_only=True)
            self.assertIsNone(workbook["螺栓缺失统计"]["J2"].value)
            workbook.close()

    def test_docx_bolt_sentence_keeps_zero_denominator_rate_blank(self) -> None:
        segment = {
            "county": "测试区",
            "route": "G1",
            "start": 1000.0,
            "end": 2000.0,
            "mileage": 1.0,
        }
        stat = {
            "segment": segment,
            "splice": 0,
            "connection": 0,
            "missing": 0,
            "rate": None,
            "points": 1,
        }
        document = Document()
        with tempfile.TemporaryDirectory() as temp_dir:
            minimal_docx._section_bolt(document, [segment], [stat], [], {}, temp_dir)
        self.assertIn("缺失率为%。", "\n".join(p.text for p in document.paragraphs))


    def test_iter_height_rows_resolves_shared_string_headers(self) -> None:
        parts = {
            "[Content_Types].xml": """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
<Override PartName="/xl/sharedStrings.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sharedStrings+xml"/>
</Types>""",
            "_rels/.rels": """<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>""",
            "xl/workbook.xml": """<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<sheets><sheet name="Sheet1" sheetId="1" r:id="rId1"/></sheets>
</workbook>""",
            "xl/_rels/workbook.xml.rels": """<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/sharedStrings" Target="sharedStrings.xml"/>
</Relationships>""",
            "xl/sharedStrings.xml": """<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" count="4" uniqueCount="4">
<si><t>护栏类型</t></si><si><t>梁板中心高度(mm)</t></si><si><t>原始桩号</t></si><si><t>三波护栏</t></si>
</sst>""",
            "xl/worksheets/sheet1.xml": """<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
<sheetData>
<row r="1"><c r="A1" t="s"><v>0</v></c><c r="B1" t="s"><v>1</v></c><c r="C1" t="s"><v>2</v></c></row>
<row r="2"><c r="A2" t="s"><v>3</v></c><c r="B2"><v>697</v></c><c r="C2"><v>2101.167</v></c></row>
</sheetData></worksheet>""",
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "shared-strings.xlsx"
            with zipfile.ZipFile(path, "w") as archive:
                for name, content in parts.items():
                    archive.writestr(name, content)
            rows = list(engine.iter_height_rows(path))
        self.assertEqual(rows[0]["护栏类型"], "三波护栏")
        self.assertEqual(rows[0]["梁板中心高度(mm)"], 697.0)

    def test_segment_lookup_requires_source_route_when_routes_overlap(self) -> None:
        segments = [
            {"route": "G319", "start": 2100000.0, "end": 2200000.0},
            {"route": "G210", "start": 2100000.0, "end": 2200000.0},
        ]
        self.assertEqual(engine._segment_index(segments, 2150000.0, "G210"), 1)
        self.assertIsNone(engine._segment_index(segments, 2150000.0, "G351"))

    def test_tci_rows_use_data_route_column_not_header_label(self) -> None:
        segments = [
            {"route": "G319", "start": 2100000.0, "end": 2200000.0},
            {"route": "G210", "start": 2100000.0, "end": 2200000.0},
        ]
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "tci.xlsx"
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.append(["区域", "路线编号", "原始桩号", "电子修正桩号", "防护设施缺损", "标志缺损", "标线缺损"])
            sheet.append(["", "", "", "", "轻", "", ""])
            sheet.append(["测试区", "G210", "K2150+000", "", 1, 0, 0])
            workbook.save(path)
            records = engine.collect_tci_records(segments, path)

        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["segment"], 1)


class DesktopBridgeTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.project = self.root / "project"
        self.project.mkdir()
        self.output = self.root / "output"
        self.summary = self.root / "summary.xlsx"
        self.manual = self.root / "manual.xlsx"
        self.route = self.root / "route.xlsx"
        for path in (self.summary, self.manual, self.route):
            path.write_bytes(b"test")
        self.detail = self.root / "detail"
        self.disease = self.root / "disease"
        self.tci = self.root / "tci"
        self.detail.mkdir()
        self.disease.mkdir()
        self.tci.mkdir()
        self.cq_template = self.root / "重庆项目报告模板.md"
        self.gd_template = self.root / "广东项目第五章模板.md"
        self.cq_template.write_text("# 重庆模板\n", encoding="utf-8")
        self.gd_template.write_text("# 广东模板\n", encoding="utf-8")
        self.bridge = DesktopBridge()

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def templates(self) -> dict[str, Path]:
        return {
            "重庆项目报告模板": self.cq_template,
            "广东项目第五章模板": self.gd_template,
        }

    def cq_payload(self) -> dict:
        return {
            "template": "cq",
            "values": {
                "projectPath": str(self.project),
                "summaryPath": str(self.summary),
                "detailPath": str(self.detail),
                "diseasePath": str(self.disease),
                "tciPath": str(self.tci),
                "outputPath": str(self.output),
            },
        }

    def test_native_window_reference_is_private_from_js_api(self) -> None:
        window = FakeWindow()
        self.bridge.attach_window(window)
        self.assertNotIn("window", vars(self.bridge))
        self.assertIs(self.bridge._window, window)

    def test_validate_chongqing_payload_creates_output(self) -> None:
        with patch.object(self.bridge, "_template_paths", return_value=self.templates()):
            self.bridge._validate_payload(self.cq_payload())
        self.assertTrue(self.output.is_dir())

    def test_start_run_validation_error_releases_lock(self) -> None:
        payload = self.cq_payload()
        payload["values"]["summaryPath"] = ""
        with patch.object(self.bridge, "_template_paths", return_value=self.templates()):
            result = self.bridge.start_run(payload)
        self.assertFalse(result["ok"])
        self.assertIn("分段汇总表", result["error"])
        self.assertFalse(self.bridge._running)

    def test_chongqing_payload_requires_tci_folder(self) -> None:
        payload = self.cq_payload()
        payload["values"]["tciPath"] = ""
        with patch.object(self.bridge, "_template_paths", return_value=self.templates()):
            result = self.bridge.start_run(payload)
        self.assertFalse(result["ok"])
        self.assertIn("TCI数据文件夹", result["error"])
        self.assertFalse(self.bridge._running)

    def test_chongqing_run_wires_tci_path_and_process_tci(self) -> None:
        configs = []
        calls = []

        def fake_run(config, **kwargs):
            configs.append(config)
            calls.append(kwargs)

        with patch.object(self.bridge, "_template_paths", return_value=self.templates()), \
                patch.object(engine, "generate_statistics_and_report", side_effect=fake_run):
            self.bridge._run_chongqing(self.cq_payload()["values"])
        self.assertEqual(str(configs[0].tci_path), str(self.tci))
        self.assertTrue(calls[0]["process_tci"])

    def test_discover_paths_detects_tci_folder(self) -> None:
        tci_dir = self.project / "TCI数据"
        tci_dir.mkdir()
        (tci_dir / "万州区-G348.xlsx").write_bytes(b"x")
        summary, detail, disease, tci = engine.discover_paths(self.project)
        self.assertEqual(tci, tci_dir)

    def test_guangdong_specialized_folders_are_optional(self) -> None:
        payload = {
            "template": "gd",
            "values": {
                "projectPath": str(self.project),
                "markingPath": "",
                "guardrailPath": "",
                "manualPath": str(self.manual),
                "routePath": str(self.route),
                "outputPath": str(self.output),
                "markingThreshold": "7",
                "heightThreshold": "5",
                "boltThreshold": "5",
            },
        }
        with patch.object(self.bridge, "_template_paths", return_value=self.templates()):
            self.bridge._validate_payload(payload)
        self.assertTrue(self.output.is_dir())

    def test_missing_specialized_template_fails_guangdong_validation(self) -> None:
        templates = self.templates()
        templates["广东项目第五章模板"] = self.root / "missing-gd.md"
        scalar = {
            "template": "gd",
            "values": {
                "projectPath": str(self.project),
                "markingPath": "",
                "guardrailPath": "",
                "manualPath": str(self.manual),
                "routePath": str(self.route),
                "outputPath": str(self.output),
                "markingThreshold": "7",
                "heightThreshold": "5",
                "boltThreshold": "5",
            },
        }
        with patch.object(self.bridge, "_template_paths", return_value=templates):
            with self.assertRaisesRegex(FileNotFoundError, "templates 文件夹"):
                self.bridge._validate_payload(scalar)

    def test_worker_emits_complete_and_releases_running_state(self) -> None:
        window = FakeWindow()
        self.bridge.attach_window(window)
        self.bridge._running = True
        with patch.object(self.bridge, "_run_chongqing"):
            self.bridge._run_worker(self.cq_payload())
        self.assertFalse(self.bridge._running)
        events = [json.loads(script.split("desktopEvents(", 1)[1][:-1]) for script in window.scripts]
        self.assertEqual(events[0]["status"], "running")
        self.assertEqual(events[-1]["status"], "complete")
        self.assertEqual(events[-1]["progress"], 100)

    def test_progress_mapping(self) -> None:
        self.assertEqual(self.bridge._progress_from_log("正在扫描资料"), (34, 1))
        self.assertEqual(self.bridge._progress_from_log("图表已生成"), (72, 2))
        self.assertEqual(self.bridge._progress_from_log("已保存文件"), (92, 3))

    def test_structured_progress_mapping_is_incremental(self) -> None:
        first = self.bridge._progress_from_log(engine.format_progress("扫描资料", 1, 4))
        last = self.bridge._progress_from_log(engine.format_progress("扫描资料", 4, 4))
        self.assertEqual(first[1], 1)
        self.assertEqual(last[1], 1)
        self.assertLess(first[0], last[0])

    def test_resource_template_path_uses_application_root(self) -> None:
        self.assertEqual(
            engine.resource_template_path("x.md"),
            engine.application_root() / "templates" / "x.md",
        )

    def test_guangdong_run_uses_resource_template_for_bundles(self) -> None:
        values = {
            "projectPath": str(self.project),
            "markingPath": "",
            "guardrailPath": "",
            "manualPath": str(self.manual),
            "routePath": str(self.route),
            "outputPath": str(self.output),
            "markingThreshold": "7",
            "heightThreshold": "5",
            "boltThreshold": "5",
        }
        scanned = {
            "marking": [{"city": "佛山市", "route": "G1"}],
            "height": [],
            "bolt": [],
            "notes": [],
            "issues": [],
        }
        expected_template = self.root / "resource-template.md"
        expected_template.write_text("# 模板", encoding="utf-8")
        route_index = type("RouteIndex", (), {"mapping": {}})()
        with (
            patch.object(engine.RouteCategoryIndex, "from_file", return_value=route_index),
            patch.object(engine.GuangdongInputScanner, "scan", return_value=scanned),
            patch.object(engine.ManualAutoComparator, "read_file", return_value=([], [])),
            patch.object(engine.GuangdongBatchRunner, "build_bundles", return_value={}),
            patch.object(engine.GuangdongBatchRunner, "run_bundles", return_value={}) as run_bundles,
            patch.object(engine, "resource_template_path", return_value=expected_template) as resource_path,
        ):
            self.bridge._run_guangdong(values)

        resource_path.assert_called_once_with("广东项目第五章模板.md")
        self.assertEqual(run_bundles.call_args.args[2], expected_template)


class UpstreamCountyTests(unittest.TestCase):
    def test_read_segments_reads_county_column(self) -> None:
        import openpyxl

        from backend.report_engine import read_segments

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.append(["序号", "区县", "路线编号", "路线名", "公路等级", "起点桩号", "止点桩号", "里程", "总里程"])
            ws.append([1, "两江新区", "G210", "满都拉－防城港", "一级公路", 2157.392, 2159.964, 2.572, 37.476])
            ws.append([1, "两江新区", "G210", "满都拉－防城港", "一级公路", 2159.964, 2184.977, 25.013, 37.476])
            wb.save(tmp_path / "summary.xlsx")
            segs = read_segments(tmp_path / "summary.xlsx")
            self.assertEqual(segs[0]["county"], "两江新区")
            self.assertEqual(segs[0]["route"], "G210")
            self.assertEqual(segs[0]["route_name"], "满都拉－防城港")


def build_demo_segments():
    return [
        {
            "grade": "一级",
            "manager": "重庆交通行政执法总队",
            "start": 2264000.0,
            "end": 2265000.0,
            "mileage": 1.0,
        },
        {
            "grade": "一级",
            "manager": "重庆交通行政执法总队",
            "start": 2265000.0,
            "end": 2266000.0,
            "mileage": 1.0,
        },
    ]


def build_demo_height_records():
    records = []
    for index, (start, end) in enumerate([(2264000.0, 2265000.0), (2265000.0, 2266000.0)]):
        for offset in (10, 20, 30):
            records.append({
                "file": "交安设施现场检测明细.xlsx",
                "direction": "上行",
                "station": start + offset * 10,
                "raw_station": start + offset * 10,
                "electronic_station": start + offset * 10,
                "basis": "电子修正桩号",
                "kind": "二波",
                "height": 580.0 + (offset * 2 % 40),
                "segment": index,
            })
        for offset in (40, 50):
            records.append({
                "file": "交安设施现场检测明细.xlsx",
                "direction": "上行",
                "station": start + offset * 10,
                "raw_station": start + offset * 10,
                "electronic_station": start + offset * 10,
                "basis": "电子修正桩号",
                "kind": "三波",
                "height": 677.0 + (offset % 20),
                "segment": index,
            })
    return records


def build_demo_bolt_records():
    return [
        {
            "file": "交安设施现场检测明细.xlsx",
            "direction": "下行",
            "station": 2264010.0,
            "raw_station": 2264010.0,
            "electronic_station": 2264010.0,
            "basis": "原始桩号",
            "segment": 0,
            "splice": 80,
            "splice_missing": 3,
            "connection": 120,
            "connection_missing": 1,
        }
    ]


def build_demo_stats(segments, records):
    stats = []
    for index in range(len(segments)):
        kinds = {"二波": [], "三波": []}
        bins = {"二波": [0] * 5, "三波": [0] * 5}
        for record in records:
            if record["segment"] != index:
                continue
            kinds[record["kind"]].append(record["height"])
            bin_index = report_engine_bin(record)
            bins[record["kind"]][bin_index] += 1
        types = {}
        for kind in ("二波", "三波"):
            heights = kinds[kind]
            pcts = [round(v * 100 / len(heights), 2) if heights else 0 for v in bins[kind]]
            types[kind] = {
                "count": len(heights),
                "bins": bins[kind],
                "pcts": pcts,
                "pass": pcts[2] if heights else 0,
            }
        stats.append({"segment": segments[index], "types": types})
    return stats


def report_engine_bin(record):
    kind = record["kind"]
    if kind == "二波":
        limits = (560, 580, 620, 640)
    else:
        limits = (657, 677, 717, 737)
    height = record["height"]
    if height < limits[0]:
        return 0
    if height < limits[1]:
        return 1
    if height <= limits[2]:
        return 2
    if height <= limits[3]:
        return 3
    return 4


class UpdaterTests(unittest.TestCase):
    def test_updater_core_validation(self) -> None:
        from updater import is_newer, parse_version, select_installer_asset, validate_download

        self.assertEqual(parse_version("v1.2.3"), (1, 2, 3))
        self.assertTrue(is_newer((1, 2, 4), (1, 2, 3)))
        self.assertFalse(is_newer((1, 2, 3), (1, 2, 3)))
        asset = select_installer_asset({
            "assets": [{
                "name": "report-generator-Setup.exe",
                "browser_download_url": "https://github.com/a/b/releases/download/v1/报告生成工具-Setup.exe",
                "size": 2,
            }],
        })
        self.assertEqual(asset["name"], "report-generator-Setup.exe")
        with self.assertRaises(ValueError):
            validate_download(b"not-an-exe", 11)


    def test_updater_rejects_untrusted_assets_and_checks_release(self) -> None:
        from updater import check_for_update, select_installer_asset

        release = {
            "tag_name": "v0.2.0",
            "assets": [{
                "name": "report-generator-Setup.exe",
                "browser_download_url": "https://objects.githubusercontent.com/a/b.exe",
                "size": 2,
            }],
        }
        result = check_for_update("0.1.0", lambda: release)
        self.assertTrue(result["update_available"])
        self.assertEqual(result["latest_version"], "0.2.0")

        malicious = dict(release)
        malicious["assets"] = [{
            "name": "report-generator-Setup.exe",
            "browser_download_url": "http://example.com/update.exe",
            "size": 2,
        }]
        with self.assertRaises(ValueError):
            select_installer_asset(malicious)

    def test_updater_launcher_does_not_use_shell(self) -> None:
        from updater import launch_installer

        with patch("updater.subprocess.Popen") as popen:
            launch_installer(Path("C:/Temp/update.exe"))
        popen.assert_called_once_with([str(Path("C:/Temp/update.exe"))], close_fds=True)


class MinimalDocxTests(unittest.TestCase):
    def test_generates_report_without_template(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir)
            skeleton = Path(temp_dir) / "重庆项目报告模板.md"
            skeleton.write_text(
                "# 2026年普通公路国省道交通安全设施自动化检测报告\n\n"
                "## 1.1 项目概况\n\n"
                "## 3.1 G210线整体情况\n\n"
                "## 5.1 G210线整体情况\n\n"
                "## 6 结论与建议\n",
                encoding="utf-8",
            )
            config = engine.Config(
                project_dir=Path(temp_dir),
                summary_xlsx=Path(temp_dir) / "summary.xlsx",
                detail_dir=Path(temp_dir),
                template_docx=skeleton,
                output_dir=output,
                disease_dir=None,
            )
            segments = build_demo_segments()
            records = build_demo_height_records()
            bolt_records = build_demo_bolt_records()
            height_stats = build_demo_stats(segments, records)
            bolt_stats = engine.make_bolt_stats(segments, bolt_records)
            messages = []
            config.out_docx.unlink(missing_ok=True)
            result = minimal_docx.run(
                config,
                segments,
                height_stats,
                records,
                bolt_stats,
                bolt_records,
                None,
                log=messages.append,
                skeleton_md=skeleton,
            )
            self.assertTrue(result.is_file())
            self.assertGreater(result.stat().st_size, 12000)
            self.assertTrue(any("Markdown" in message or "程序化" in message for message in messages))

    def test_engine_routes_to_minimal_without_template(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            output = root / "output"
            skeleton = root / "重庆项目报告模板.md"
            skeleton.write_text(
                "# 2026年普通公路国省道交通安全设施自动化检测报告\n\n"
                "## 1.1 项目概况\n\n"
                "## 3.1 G210线整体情况\n\n"
                "## 5.1 G210线整体情况\n\n"
                "## 6 结论与建议\n",
                encoding="utf-8",
            )
            config = engine.Config(
                project_dir=root,
                summary_xlsx=root / "summary.xlsx",
                detail_dir=root,
                template_docx=skeleton,
                output_dir=output,
                disease_dir=None,
            )
            segments = build_demo_segments()
            records = build_demo_height_records()
            bolt_records = build_demo_bolt_records()
            height_stats = build_demo_stats(segments, records)
            bolt_stats = engine.make_bolt_stats(segments, bolt_records)
            messages = []
            engine.make_docx(
                config,
                segments,
                height_stats=height_stats,
                height_records=records,
                bolt_stats=bolt_stats,
                bolt_records=bolt_records,
                disease_image_index=None,
                log=messages.append,
                require_template=False,
            )
            self.assertTrue(config.out_docx.is_file())
            self.assertGreater(config.out_docx.stat().st_size, 12000)
            self.assertTrue(any("Markdown" in message or "程序化" in message for message in messages))

    def test_generates_report_from_markdown_skeleton(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            skeleton = root / "重庆项目报告模板.md"
            skeleton.write_text(
                "# 2026年普通公路国省道交通安全设施自动化检测报告\n\n"
                "## 1.1 项目概况\n\n"
                "本报告依据委托单位提供的数据生成。\n\n"
                "## 3.1 G210线整体情况\n\n"
                "### 3.2 G210线K2264+000~K2265+000段（样例）\n\n"
                "## 5 结论与建议\n\n"
                "## 6 建议\n\n"
                "标志牌安装情况检查合格。\n",
                encoding="utf-8",
            )
            config = engine.Config(
                project_dir=root,
                summary_xlsx=root / "summary.xlsx",
                detail_dir=root,
                template_docx=skeleton,
                output_dir=root / "output",
                disease_dir=None,
            )
            segments = build_demo_segments()
            records = build_demo_height_records()
            bolt_records = build_demo_bolt_records()
            height_stats = build_demo_stats(segments, records)
            bolt_stats = engine.make_bolt_stats(segments, bolt_records)
            result = minimal_docx.run(
                config,
                segments,
                height_stats,
                records,
                bolt_stats,
                bolt_records,
                None,
                skeleton_md=skeleton,
            )
            self.assertTrue(result.is_file())
            document = Document(result)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("自动化检测报告", text)
            self.assertIn("本报告依据委托单位提供的数据生成。", text)
            self.assertIn("波形梁护栏横梁中心高度检测结果", text)
            self.assertIn("结论与建议", text)
            self.assertNotIn("（样例）", text)

    def test_overview_table_has_county_column(self) -> None:
        segs = [{"county": "两江新区", "route": "G210", "route_name": "满都拉－防城港", "grade": "一级公路", "start": 2157392, "end": 2159964, "mileage": 2.572, "total_mileage": 37.476, "manager": ""}]
        doc = Document()
        minimal_docx._section_overview(doc, None, segs)
        table = doc.tables[0]
        headers = [c.text for c in table.rows[0].cells]
        self.assertEqual(headers[1], "区县")
        self.assertEqual(headers[3], "路线名")
        self.assertEqual(len(headers), 9)
        self.assertEqual(headers[0], "序号")
        row = [c.text for c in table.rows[1].cells]
        self.assertEqual(row[1], "两江新区")
        self.assertEqual(row[3], "满都拉－防城港")

    def test_height_grouped_by_county(self) -> None:
        segments = [
            {"county": "两江新区", "route": "G210", "route_name": "", "grade": "一级", "manager": "", "start": 2264000.0, "end": 2265000.0, "mileage": 1.0, "total_mileage": 1.0},
            {"county": "北碚区", "route": "G210", "route_name": "", "grade": "一级", "manager": "", "start": 2265000.0, "end": 2266000.0, "mileage": 1.0, "total_mileage": 1.0},
        ]
        records = build_demo_height_records()
        height_stats = build_demo_stats(segments, records)
        bolt_records = [
            {"file": "交安设施现场检测明细.xlsx", "direction": "下行", "station": 2264010.0, "raw_station": 2264010.0, "electronic_station": 2264010.0, "basis": "原始桩号", "segment": 0, "splice": 80, "splice_missing": 3, "connection": 120, "connection_missing": 1},
            {"file": "交安设施现场检测明细.xlsx", "direction": "下行", "station": 2265010.0, "raw_station": 2265010.0, "electronic_station": 2265010.0, "basis": "原始桩号", "segment": 1, "splice": 80, "splice_missing": 3, "connection": 120, "connection_missing": 1},
        ]
        bolt_stats = engine.make_bolt_stats(segments, bolt_records)
        with tempfile.TemporaryDirectory() as tmp:
            import matplotlib.pyplot as plt

            images = {}
            for idx in range(len(segments)):
                for kind in ("二波", "三波"):
                    if height_stats[idx]["types"][kind]["count"] == 0:
                        continue
                    fig, ax = plt.subplots(figsize=(1, 1))
                    ax.plot([0, 1], [0, 1])
                    line_path = Path(tmp) / f"line_{idx}_{kind}.png"
                    fig.savefig(line_path)
                    plt.close(fig)
                    fig2, ax2 = plt.subplots(figsize=(1, 1))
                    ax2.pie([1, 1])
                    pie_path = Path(tmp) / f"pie_{idx}_{kind}.png"
                    fig2.savefig(pie_path)
                    plt.close(fig2)
                    images[(idx, kind)] = {"line": line_path, "pie": pie_path}
            doc = Document()
            minimal_docx._section_height(doc, segments, height_stats, records, images, tmp)
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
            self.assertTrue(any("两江新区整体情况" in h for h in headings), f"missing 两江新区整体情况 in {headings}")
            self.assertTrue(any("北碚区整体情况" in h for h in headings), f"missing 北碚区整体情况 in {headings}")
            doc2 = Document()
            minimal_docx._section_bolt(doc2, segments, bolt_stats, bolt_records, None, tmp)
            headings2 = [p.text for p in doc2.paragraphs if p.style.name.startswith("Heading")]
            self.assertTrue(any("两江新区整体情况" in h for h in headings2), f"bolt missing 两江新区整体情况 in {headings2}")
            self.assertTrue(any("北碚区整体情况" in h for h in headings2), f"bolt missing 北碚区整体情况 in {headings2}")
            tables = doc.tables
            has_county_col = any(any("区县" in c.text for c in t.rows[0].cells) for t in tables) if tables else False
            tables2 = doc2.tables
            has_county_col2 = any(any("区县" in c.text for c in t.rows[0].cells) for t in tables2) if tables2 else False
            self.assertTrue(has_county_col, f"height tables missing 区县 column, headers: {[[c.text for c in t.rows[0].cells] for t in tables]}")
            self.assertTrue(has_county_col2, f"bolt tables missing 区县 column, headers: {[[c.text for c in t.rows[0].cells] for t in tables2]}")


class GuangdongBusinessRegressionTests(unittest.TestCase):
    def test_scanner_preserves_bridge_guardrail_note(self) -> None:
        scanner = engine.GuangdongInputScanner(".")
        note_row = scanner._convert(
            "height",
            {
                "地市": "佛山市",
                "路线编号": "G1",
                "方向": "上行",
                "电子修正桩号": "K1+000",
                "检测区段": "K1+000～K2+000",
                "护栏类型": "二波",
                "梁板中心高度(mm)": "",
                "备注标记": "桥梁地段",
            },
            Path("guardrail.xlsx"),
            "Sheet1",
        )
        self.assertIsNotNone(note_row)
        self.assertEqual(note_row["guardrail_note"], "桥梁地段")

    def test_writer_handles_single_wave_and_bridge_only_segments(self) -> None:
        def height(segment: str, kind: str, value: float) -> dict:
            return {
                "category": "高速公路",
                "route": "G1",
                "direction": "上行",
                "segment": segment,
                "guardrail_type": kind,
                "height": value,
            }

        def note(segment: str) -> dict:
            return {
                "category": "高速公路",
                "route": "G1",
                "direction": "上行",
                "segment": segment,
                "guardrail_note": "桥梁地段",
            }

        bundle = {
            "marking": [],
            "height": [
                height("K1+000～K2+000", "二波", 600),
                height("K2+000～K3+000", "三波", 700),
            ],
            "bolt": [],
            "notes": [note("K3+000～K4+000")],
            "comparison_detail": [],
            "weak_segments": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template = root / "template.md"
            template.write_text(
                "# 五、交通安全设施技术状况检测评价情况\n\n"
                "本章为{{地市}}交通安全设施技术状况检测评价内容。\n",
                encoding="utf-8",
            )
            output = engine.GuangdongChapterWriter.write(
                "佛山市",
                bundle,
                root / "output",
                template,
                {"marking": 7, "height": 5, "bolt": 5},
            )
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("其中二波护栏", text)
        self.assertIn("其中三波护栏", text)
        self.assertEqual(text.count("当前区段为桥梁路段，无有效检测点位。"), 2)
        self.assertNotIn("区段内无护栏", text)
        self.assertNotIn("共检测0个有效点，其中。", text)

    def test_writer_deduplicates_height_and_note_segments_with_suffix(self) -> None:
        bundle = {
            "marking": [],
            "height": [{
                "category": "高速公路",
                "route": "G1",
                "direction": "上行",
                "segment": "K1+000～K2+000",
                "guardrail_type": "二波",
                "height": 600,
            }],
            "bolt": [],
            "notes": [{
                "category": "高速公路",
                "route": "G1",
                "direction": "上行",
                "segment": "K1+000～K2+000段",
                "guardrail_note": "桥梁地段",
            }],
            "comparison_detail": [],
            "weak_segments": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template = root / "template.md"
            template.write_text(
                "# 五、交通安全设施技术状况检测评价情况\n\n"
                "本章为{{地市}}交通安全设施技术状况检测评价内容。\n",
                encoding="utf-8",
            )
            output = engine.GuangdongChapterWriter.write(
                "佛山市",
                bundle,
                root / "output",
                template,
                {"marking": 7, "height": 5, "bolt": 5},
            )
            document = Document(output)
            headings = []
            height_text = []
            in_height_section = False
            for paragraph in document.paragraphs:
                if paragraph.text == "（2）护栏中心高度":
                    in_height_section = True
                elif paragraph.text == "（3）螺栓安装情况":
                    in_height_section = False
                elif in_height_section:
                    height_text.append(paragraph.text)
                    if (
                        paragraph.style.name == "Heading 5"
                        and "K1+000～K2+000段" in paragraph.text
                    ):
                        headings.append(paragraph.text)

        self.assertEqual(headings, ["a. K1+000～K2+000段"])
        self.assertNotIn("当前区段为桥梁路段，无有效检测点位。", "\n".join(height_text))

    def test_writer_accepts_markdown_template_with_city_placeholder(self) -> None:
        bundle = {
            "marking": [],
            "height": [{
                "category": "高速公路",
                "route": "G1",
                "direction": "上行",
                "segment": "K1+000～K2+000",
                "guardrail_type": "二波",
                "height": 600,
            }],
            "bolt": [],
            "notes": [],
            "comparison_detail": [],
            "weak_segments": [],
        }
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template = root / "template.md"
            template.write_text(
                "# 五、交通安全设施技术状况检测评价情况\n\n"
                "本章为{{地市}}交通安全设施技术状况检测评价内容。\n\n"
                "## （一）高速公路交安设施技术状况\n\n"
                "### 1.沿线设施技术状况TCI\n",
                encoding="utf-8",
            )
            output = engine.GuangdongChapterWriter.write(
                "佛山市",
                bundle,
                root / "output",
                template,
                {"marking": 7, "height": 5, "bolt": 5},
            )
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("本章为佛山市交通安全设施技术状况检测评价内容。", text)
        self.assertIn("五、交通安全设施技术状况检测评价情况", text)
        self.assertIn("（2）护栏中心高度", text)
        self.assertIn("其中二波护栏", text)
        self.assertNotIn("{{地市}}", text)

    def test_comparison_table_emits_expected_two_level_bolt_header_xml(self) -> None:
        document = Document()
        engine.GuangdongChapterWriter._comparison_table(
            document,
            [{
                "indicator": "bolt",
                "route": "G1",
                "gtype": "二波",
                "position": "左侧",
                "direction": "上行",
                "segment": "K1+000～K2+000",
                "msplice": 1,
                "mconn": 2,
                "asplice": 1,
                "aconn": 2,
                "remark": "",
            }],
            {"marking": 7, "height": 5, "bolt": 5},
        )
        table = document.tables[1]
        xml_rows = table._tbl.findall("./w:tr", table._tbl.nsmap)
        cell_counts = [len(row.findall("./w:tc", table._tbl.nsmap)) for row in xml_rows]
        top_cells = xml_rows[0].findall("./w:tc", table._tbl.nsmap)
        spans = []
        for cell in top_cells:
            span = cell.find("./w:tcPr/w:gridSpan", cell.nsmap)
            spans.append(int(span.get(qn("w:val"))) if span is not None else 1)
        grid_columns = table._tbl.findall("./w:tblGrid/w:gridCol", table._tbl.nsmap)

        self.assertEqual(cell_counts, [8, 10, 10])
        self.assertEqual(len(grid_columns), 10)
        self.assertEqual(spans, [1, 1, 1, 1, 1, 2, 2, 1])

    def test_comparison_analysis_uses_average_without_min_max_range(self) -> None:
        document = Document()
        detail = [
            {
                "indicator": "height",
                "absolute_difference": 1.0,
                "relative_deviation": None,
                "within_threshold": True,
            },
            {
                "indicator": "height",
                "absolute_difference": 3.0,
                "relative_deviation": None,
                "within_threshold": False,
            },
        ]
        engine.GuangdongChapterWriter._comparison_table(
            document,
            detail,
            {"marking": 7, "height": 5, "bolt": 5},
        )
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("平均偏差", text)
        self.assertIn("一致性占比", text)
        self.assertNotIn("1.00～3.00", text)


REAL_E2E_ENV = "REPORT_E2E_REAL"
_REAL_E2E_PATHS = {
    "route_xlsx": ("REPORT_E2E_ROUTE_XLSX", "file"),
    "manual_xlsx": ("REPORT_E2E_MANUAL_XLSX", "file"),
    "template": ("REPORT_E2E_TEMPLATE_DOCX", "file"),
    "foshan_marking": ("REPORT_E2E_FOSHAN_MARKING_DIR", "dir"),
    "foshan_guardrail": ("REPORT_E2E_FOSHAN_GUARDRAIL_DIR", "dir"),
    "zhuhai_marking": ("REPORT_E2E_ZHUHAI_MARKING_DIR", "dir"),
    "zhuhai_guardrail": ("REPORT_E2E_ZHUHAI_GUARDRAIL_DIR", "dir"),
}


def _required_real_path(key: str) -> Path:
    env_name, kind = _REAL_E2E_PATHS[key]
    value = os.environ.get(env_name, "").strip()
    if not value:
        raise AssertionError(f"真实数据测试必须设置环境变量：{env_name}")
    path = Path(value)
    valid = path.is_file() if kind == "file" else path.is_dir()
    if not valid:
        raise AssertionError(f"{env_name} 路径不存在或类型错误：{path}")
    return path


def _real_e2e_inputs() -> dict[str, Path]:
    return {key: _required_real_path(key) for key in _REAL_E2E_PATHS}


class RealDataConfigurationTests(unittest.TestCase):
    def test_real_e2e_requires_input_environment_variables(self) -> None:
        env = {REAL_E2E_ENV: "1"}
        env.update({env_name: "" for env_name, _ in _REAL_E2E_PATHS.values()})
        with patch.dict(os.environ, env, clear=False):
            with self.assertRaisesRegex(AssertionError, "REPORT_E2E_ROUTE_XLSX"):
                _real_e2e_inputs()


def _scan_city(city: dict, route_index) -> dict:
    """真实数据扫描：与 run_guangdong_project 中标线/护栏分别扫描的同款逻辑。"""
    scanned = {"height": [], "bolt": [], "marking": [], "notes": [], "issues": []}
    for key, kinds in (("marking", ("marking",)), ("guardrail", ("height", "bolt", "notes"))):
        root = city.get(key)
        if not root or not root.is_dir():
            continue
        partial = engine.GuangdongInputScanner(root, route_index).scan()
        for kind in kinds:
            scanned[kind].extend(partial.get(kind, []))
        scanned["issues"].extend(partial.get("issues", []))
    return scanned


def _run_real_pipeline(
    cities: list[dict],
    artifact_root: Path,
    route_xlsx: Path,
    manual_xlsx: Path,
    template: Path,
) -> dict:
    """使用运行环境提供的真实输入：直接拼装 scan+build_bundles+run_bundles。"""
    artifact_root.mkdir(parents=True, exist_ok=True)
    route_index = engine.RouteCategoryIndex.from_file(route_xlsx)
    manual_records, manual_issues = engine.ManualAutoComparator.read_file(manual_xlsx)
    thresholds = {"marking": 7, "height": 5, "bolt": 5}

    result = {"success": [], "failed": {}, "warnings": []}
    for city in cities:
        # 每个城市独立 scan/build/run，避免跨城市路线记录进入同一个 bundle。
        scanned = _scan_city(city, route_index)
        scanned["issues"] = list(manual_issues) + scanned["issues"]
        bundles = engine.GuangdongBatchRunner.build_bundles(scanned, route_index, manual_records, thresholds)
        partial = engine.GuangdongBatchRunner.run_bundles(
            bundles, artifact_root, template, thresholds
        )
        result["success"].extend(partial["success"])
        result["failed"].update(partial["failed"])
        result["warnings"].extend(partial["warnings"])
    return result


def _collect_docx_text(docx_path: Path) -> str:
    document = Document(str(docx_path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _collect_segment_order(text: str, marker: str) -> list[str]:
    """按指标汇总表分段提取 a. 区段标题中的桩号范围。"""
    import re

    summary_markers = (
        "标线逆反射区段汇总表",
        "护栏中心高度区段汇总表",
        "护栏螺栓缺失区段汇总表",
    )
    aliases = {
        "标线逆反射区段": "标线逆反射区段汇总表",
        "护栏中心高度": "护栏中心高度区段汇总表",
        "螺栓": "护栏螺栓缺失区段汇总表",
    }
    wanted = aliases.get(marker, marker)
    if wanted not in summary_markers:
        raise ValueError(f"未知指标汇总表 marker：{marker}")

    title = re.compile(r"^\s*[a-z]\.\s+")
    pattern = re.compile(r"K\+?\d+(?:\.\d+)?[~～]K\+?\d+(?:\.\d+)?")
    result: list[str] = []
    active = False
    for line in text.splitlines():
        found = next((item for item in summary_markers if item in line), None)
        if found is not None:
            active = found == wanted
            continue
        if not active or not title.match(line):
            continue
        match = pattern.search(line)
        if match:
            result.append(match.group(0))
    return result


@unittest.skipUnless(os.environ.get(REAL_E2E_ENV) == "1", f"set {REAL_E2E_ENV}=1 to run")
class RealDataE2ETests(unittest.TestCase):
    """真实佛山/珠海端到端回归；产物仅落在 tests/artifacts/real-data/，不进 git。"""

    @classmethod
    def setUpClass(cls) -> None:
        paths = _real_e2e_inputs()
        cls.route_xlsx = paths["route_xlsx"]
        cls.manual_xlsx = paths["manual_xlsx"]
        cls.template = paths["template"]
        cls.foshan = {
            "name": "佛山",
            "marking": paths["foshan_marking"],
            "guardrail": paths["foshan_guardrail"],
        }
        cls.zhuhai = {
            "name": "珠海",
            "marking": paths["zhuhai_marking"],
            "guardrail": paths["zhuhai_guardrail"],
        }

    def _run(self, cities: list[dict], sub: str) -> dict:
        artifact_root = Path(__file__).resolve().parent / "artifacts" / "real-data" / sub
        if artifact_root.exists():
            # 清空旧产物，避免上轮 run 的 docx 干扰本次断言。
            import shutil

            shutil.rmtree(artifact_root, ignore_errors=True)
        return _run_real_pipeline(cities, artifact_root, self.route_xlsx, self.manual_xlsx, self.template)

    def test_foshan_minimal_integration_runs_end_to_end(self) -> None:
        result = self._run([self.foshan], "foshan")
        self.assertIn("佛山市", result["success"])
        foshan_dir = Path(__file__).resolve().parent / "artifacts" / "real-data" / "foshan" / "佛山市"
        docx_path = next(foshan_dir.glob("*第五部分.docx"), None)
        self.assertIsNotNone(docx_path, f"未生成佛山 docx：{foshan_dir}")
        self._assert_docx_complete(docx_path, "佛山")
        text = _collect_docx_text(docx_path)
        self.assertEqual(text.count("区段内无护栏"), 0)
        self.assertEqual(text.count("共检测0个有效点，其中。"), 0)

    def test_foshan_and_zhuhai_full_e2e_artifact_checks(self) -> None:
        result = self._run([self.foshan, self.zhuhai], "foshan-zhuhai")
        self.assertIn("佛山市", result["success"])
        self.assertIn("珠海市", result["success"], msg=f"珠海失败：{result['failed']}")
        self.assertFalse(result["failed"], f"运行失败：{result['failed']}")

        artifact_root = Path(__file__).resolve().parent / "artifacts" / "real-data" / "foshan-zhuhai"
        foshan_docx = next((artifact_root / "佛山市").glob("*第五部分.docx"), None)
        zhuhai_docx = next((artifact_root / "珠海市").glob("*第五部分.docx"), None)
        self.assertIsNotNone(foshan_docx, "佛山 docx 缺失")
        self.assertIsNotNone(zhuhai_docx, "珠海 docx 缺失")
        self._assert_docx_complete(foshan_docx, "佛山")
        self._assert_docx_complete(zhuhai_docx, "珠海")

        # 仅抽取检查中用得到的城市子串，避免佛山段落污染珠海或反之。
        foshan_text = _collect_docx_text(foshan_docx)
        zhuhai_text = _collect_docx_text(zhuhai_docx)

        for label, text in (("佛山", foshan_text), ("珠海", zhuhai_text)):
            self.assertEqual(text.count("区段内无护栏"), 0, f"{label}出现旧句：区段内无护栏")
            self.assertEqual(text.count("共检测0个有效点，其中。"), 0, f"{label}出现残句：共检测0个有效点，其中。")

        # 标线/高度/螺栓三章节共同区段顺序一致：抽取各章节首次出现的区段链，比较相等。
        for label, text in (("佛山", foshan_text), ("珠海", zhuhai_text)):
            for marker in ("标线逆反射区段", "护栏中心高度", "螺栓"):
                segments = _collect_segment_order(text, marker)
                self.assertGreater(len(segments), 1, f"{label}/{marker} 区段数不足，无法比序")
            marking = _collect_segment_order(text, "标线逆反射区段")
            height = _collect_segment_order(text, "护栏中心高度")
            bolt = _collect_segment_order(text, "螺栓")
            common = [s for s in marking if s in set(height) & set(bolt)]
            self.assertGreater(len(common), 1, f"{label} 共同区段数={len(common)}，无法比序")
            height_in_common = [s for s in height if s in set(common)]
            bolt_in_common = [s for s in bolt if s in set(common)]
            self.assertEqual(common, height_in_common, f"{label} 标线/高度共同区段相对顺序不一致")
            self.assertEqual(common, bolt_in_common, f"{label} 标线/螺栓共同区段相对顺序不一致")

        # 抽查珠海 docx 的一个螺栓对比表：物理 w:tc 8/10/10、两个父表头 gridSpan=2。
        bolt_table_index, bolt_table = self._find_bolt_comparison_table(zhuhai_docx)
        with zipfile.ZipFile(zhuhai_docx) as archive:
            xml = archive.read(f"word/document.xml")
        self._assert_two_level_bolt_header(
            xml, bolt_table_index, label=f"珠海 螺栓对比表 #{bolt_table_index}"
        )

    def _find_bolt_comparison_table(self, docx_path: Path):
        document = Document(str(docx_path))
        for index, table in enumerate(document.tables):
            text = "\n".join(cell.text or "" for row in table.rows for cell in row.cells)
            rows = table._tbl.findall("./w:tr", table._tbl.nsmap)
            cell_counts = [len(row.findall("./w:tc", table._tbl.nsmap)) for row in rows]
            if "拼接螺栓" in text and cell_counts[:3] == [8, 10, 10]:
                return index, table
        self.fail(f"未在 {docx_path} 中找到物理列为8/10/10的螺栓对比表")

    def _assert_docx_complete(self, docx_path: Path, label: str) -> None:
        with zipfile.ZipFile(docx_path) as archive:
            self.assertIsNone(archive.testzip(), f"{label} docx ZIP 存在损坏成员")

    def _assert_two_level_bolt_header(self, document_xml: bytes, table_index: int, label: str) -> None:
        from xml.etree import ElementTree as ET

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = ET.fromstring(document_xml)
        tables = root.findall(f".//{{{ns['w']}}}tbl")
        if not 0 <= table_index < len(tables):
            self.fail(f"{label} XML 表格索引越界：{table_index}/{len(tables)}")
        table = tables[table_index]
        text = "".join(t.text or "" for t in table.iter(f"{{{ns['w']}}}t"))
        self.assertIn("拼接螺栓", text, f"{label} XML 表格内容不符")
        rows = table.findall(f"{{{ns['w']}}}tr")
        cell_counts = [len(r.findall(f"{{{ns['w']}}}tc")) for r in rows]
        self.assertGreaterEqual(len(cell_counts), 3, f"{label} 行数不足：{cell_counts}")
        self.assertEqual(cell_counts[:3], [8, 10, 10], f"{label} 物理列不符：{cell_counts}")
        self.assertTrue(all(count == 10 for count in cell_counts[2:]), f"{label} 数据行物理列不符：{cell_counts}")
        grid = table.find(f"{{{ns['w']}}}tblGrid")
        grid_columns = [] if grid is None else grid.findall(f"{{{ns['w']}}}gridCol")
        self.assertEqual(len(grid_columns), 10, f"{label} 物理网格列不符：{len(grid_columns)}")
        top = rows[0].findall(f"{{{ns['w']}}}tc")
        spans = []
        for cell in top:
            span = cell.find(f"{{{ns['w']}}}tcPr/{{{ns['w']}}}gridSpan")
            spans.append(int(span.get(f"{{{ns['w']}}}val")) if span is not None else 1)
        self.assertEqual(spans.count(2), 2, f"{label} 父表头 gridSpan=2 应有2个：{spans}")
        self.assertEqual(len(spans), 8, f"{label} 父表头数量不符：{spans}")


class ChongqingCountyNamingTests(unittest.TestCase):
    def _config(self, root: Path, county=None) -> object:
        return engine.Config(
            root, root / "summary.xlsx", root, root / "template.md", root / "output",
            county=county,
        )

    def test_default_names_keep_legacy(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            config = self._config(Path(temp_dir))
            self.assertEqual(config.out_docx.name, engine.OUT_DOCX_NAME)
            self.assertEqual(config.out_xlsx.name, engine.OUT_XLSX_NAME)

    def test_county_names_follow_chongqing_pattern(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            config = self._config(Path(temp_dir), county="万州区")
            self.assertEqual(config.out_docx.name, "重庆市万州区交安设施检测报告.docx")
            self.assertEqual(config.out_xlsx.name, "重庆市万州区交安设施检测报告.xlsx")

    def test_county_stem_does_not_duplicate_prefix(self) -> None:
        self.assertEqual(engine.county_report_stem("重庆市万州区"), "重庆市万州区交安设施检测报告")
        self.assertEqual(engine.county_report_stem("城口县"), "重庆市城口县交安设施检测报告")


class CountyRoutingTests(unittest.TestCase):
    def _segments(self) -> list:
        return [
            {"county": "万州区", "route": "G210", "start": 2264000.0, "end": 2265000.0},
            {"county": "渝北区", "route": "G210", "start": 2265000.0, "end": 2266000.0},
        ]

    def test_extract_full_short_and_multi(self) -> None:
        known = ["万州区", "渝北区", "城口县"]
        self.assertEqual(
            engine.extract_counties_from_filenames(["重庆市-万州区-交安设施现场检测-明细.xlsx"], known),
            ["万州区"],
        )
        self.assertEqual(engine.extract_counties_from_filenames(["渝北-G210-明细.xlsx"], known), ["渝北区"])
        self.assertEqual(
            engine.extract_counties_from_filenames(["万州区-明细.xlsx", "城口-病害清单.xlsx"], known),
            ["万州区", "城口县"],
        )
        self.assertEqual(engine.extract_counties_from_filenames(["G210-明细.xlsx"], known), [])

    def test_resolve_without_files_or_dimension_returns_all_or_empty(self) -> None:
        self.assertEqual(
            engine.resolve_report_counties(self._segments(), []),
            ["万州区", "渝北区"],
        )
        self.assertEqual(engine.resolve_report_counties([{"route": "G210"}], ["万州区-明细.xlsx"]), [])

    def test_resolve_without_match_raises_with_candidates(self) -> None:
        with self.assertRaises(ValueError) as raised:
            engine.resolve_report_counties(self._segments(), ["G210-明细.xlsx"])
        self.assertIn("手动选择", str(raised.exception))
        self.assertIn("万州区", str(raised.exception))

    def test_resolve_override_short_ok_and_missing_raises(self) -> None:
        self.assertEqual(engine.resolve_report_counties(self._segments(), [], override="渝北"), ["渝北区"])
        with self.assertRaises(ValueError) as raised:
            engine.resolve_report_counties(self._segments(), [], override="江北区")
        self.assertIn("无对应路线分段", str(raised.exception))


class ChongqingHeadingNumberingTests(unittest.TestCase):
    def _document_xml(self, template_text: str) -> tuple:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template = root / "template.md"
            template.write_text(template_text, encoding="utf-8")
            output = root / "output"
            output.mkdir()
            config = engine.Config(root, root / "summary.xlsx", root, template, output)
            minimal_docx.make_report(
                config, [], None, [], None, [], None, [], {}, root,
                skeleton_md=template,
            )
            with zipfile.ZipFile(output / engine.OUT_DOCX_NAME) as archive:
                return (
                    archive.read("word/document.xml").decode("utf-8"),
                    archive.read("word/numbering.xml").decode("utf-8"),
                )

    def _paragraphs_by_style(self, document_xml: str) -> dict:
        from xml.etree import ElementTree as ET

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = ET.fromstring(document_xml)
        result: dict = {}
        for paragraph in root.findall(".//w:body/w:p", ns):
            style = paragraph.find("w:pPr/w:pStyle", ns)
            if style is None:
                continue
            name = style.get(f"{{{ns['w']}}}val")
            text = "".join(t.text or "" for t in paragraph.findall("w:r/w:t", ns))
            numbered = paragraph.find("w:pPr/w:numPr", ns) is not None
            result.setdefault(name, []).append((text, numbered))
        return result

    def test_multilevel_numbering_and_stripped_text(self) -> None:
        document, numbering = self._document_xml(
            "# 1 概况\n\n## 1.1 项目概况\n\n### 2.3.1 评价\n\n#### 四级标题\n\n##### 五级标题\n"
        )
        self.assertIn('w:val="%1."', numbering)
        self.assertIn('w:val="%1.%2"', numbering)
        self.assertIn('w:val="%1.%2.%3"', numbering)
        by_style = self._paragraphs_by_style(document)
        self.assertEqual(by_style["Heading1"][0][0], "概况")
        self.assertTrue(all(numbered for _, numbered in by_style["Heading1"]))
        self.assertTrue(all(numbered for _, numbered in by_style["Heading2"]))
        self.assertTrue(all(numbered for _, numbered in by_style["Heading3"]))
        self.assertEqual(by_style["Heading3"][0][0], "评价")
        self.assertTrue(all(not numbered for _, numbered in by_style.get("Heading4", [])))
        self.assertTrue(all(not numbered for _, numbered in by_style.get("Heading5", [])))

    def test_year_prefix_not_stripped(self) -> None:
        document, _ = self._document_xml("# 2026年普通公路检测报告\n")
        by_style = self._paragraphs_by_style(document)
        self.assertEqual(by_style["Heading1"][0][0], "2026年普通公路检测报告")


class ChongqingTableHeaderTests(unittest.TestCase):
    def test_template_fonts_bold_and_gray_shading(self) -> None:
        template = markdown_skeleton.read_template(engine.resource_template_path("重庆项目报告模板.md"))
        self.assertEqual(template.config["body"]["east_asia"], "宋体")
        self.assertEqual(template.config["table"]["east_asia"], "宋体")
        self.assertTrue(template.config["table"]["header_bold"])
        self.assertEqual(template.config["table"]["header_shading"], "D9D9D9")
        for level in ("1", "2", "3"):
            self.assertEqual(template.config["heading"][level]["east_asia"], "黑体")
            self.assertTrue(template.config["heading"][level]["bold"])
        for level in ("4", "5"):
            self.assertEqual(template.config["heading"][level]["east_asia"], "黑体")
            self.assertFalse(template.config["heading"][level]["bold"])

    def test_docx_header_gray_bold(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            output = root / "output"
            output.mkdir()
            template = engine.resource_template_path("重庆项目报告模板.md")
            config = engine.Config(root, root / "summary.xlsx", root, template, output)
            segment = {
                "county": "万州区", "route": "G210", "route_name": "", "grade": "一级",
                "start": 2264000.0, "end": 2265000.0, "mileage": 1.0, "total_mileage": 1.0,
            }
            minimal_docx.make_report(
                config, [segment], None, [], None, [], None, [], {}, root,
                skeleton_md=template,
            )
            with zipfile.ZipFile(output / engine.OUT_DOCX_NAME) as archive:
                document = archive.read("word/document.xml").decode("utf-8")
        self.assertIn('w:fill="D9D9D9"', document)
        self.assertIn("w:eastAsia=\"宋体\"", document)

    def test_excel_header_gray_bold(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            config = engine.Config(root, root / "summary.xlsx", root, root / "template.md", root / "output")
            segment = {
                "county": "测试区", "route": "G1", "start": 1000.0, "end": 2000.0, "mileage": 1.0,
            }
            engine.make_excel(
                config,
                [segment],
                bolt_stats=[{
                    "segment": segment, "splice": 0, "connection": 0,
                    "missing": 0, "rate": None,
                }],
            )
            workbook = openpyxl.load_workbook(config.out_xlsx, data_only=True)
            header = workbook["螺栓缺失统计"]["A1"]
            self.assertTrue(header.font.bold)
            self.assertTrue(str(header.fill.start_color.rgb).upper().endswith("D9D9D9"))
            self.assertTrue(str(header.font.color.rgb).upper().endswith("000000"))
            workbook.close()


class ChongqingCountyEndToEndTests(unittest.TestCase):
    def _write_summary(self, path: Path) -> None:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "各区县项目概况"
        sheet.append(["序号", "区县", "路线编号", "路线名", "公路等级", "起点桩号", "止点桩号", "里程", "总里程"])
        sheet.append([1, "万州区", "G210", "", "一级", 2264.0, 2265.0, 1.0, 1.0])
        sheet.append([2, "渝北区", "G210", "", "一级", 2265.0, 2266.0, 1.0, 1.0])
        workbook.save(path)

    def _write_detail(self, path: Path, station: str) -> None:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.append(["护栏类型", "梁板中心高度(mm)", "原始桩号", "电子修正桩号", "方向", "路线编号", "异常标记"])
        sheet.append(["两波护栏", 600, station, station, "上行", "G210", ""])
        workbook.save(path)

    def _skeleton(self, path: Path) -> None:
        path.write_text("# 报告\n\n## 1.1 项目概况\n", encoding="utf-8")

    def _run(self, root: Path, detail_names: list) -> object:
        summary = root / "summary.xlsx"
        self._write_summary(summary)
        detail = root / "detail"
        detail.mkdir()
        stations = {"wanzhou": "K2264+100", "yubei": "K2265+100"}
        for name in detail_names:
            key = "yubei" if "渝北" in name or "yubei" in name else "wanzhou"
            self._write_detail(detail / name, stations[key])
        skeleton = root / "template.md"
        self._skeleton(skeleton)
        output = root / "output"
        config = engine.Config(root, summary, detail, skeleton, output)
        engine.generate_statistics_and_report(
            config, log=lambda _: None, process_height=True,
            process_bolts=False, process_tci=False, require_template=False,
        )
        return config

    def test_multi_county_files_generate_each_county_report(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            self._run(root, ["重庆市-万州区-交安设施现场检测-明细.xlsx", "渝北-G210-明细.xlsx"])
            for county in ("万州区", "渝北区"):
                stem = engine.county_report_stem(county)
                self.assertTrue((root / "output" / county / f"{stem}.docx").is_file(), county)
                self.assertTrue((root / "output" / county / f"{stem}.xlsx").is_file(), county)

    def test_single_county_file_filters_segments_and_names_top_files(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            config = self._run(root, ["重庆市-万州区-交安设施现场检测-明细.xlsx"])
            self.assertEqual(config.out_docx.name, "重庆市万州区交安设施检测报告.docx")
            self.assertEqual(config.out_xlsx.name, "重庆市万州区交安设施检测报告.xlsx")
            self.assertTrue(config.out_docx.is_file())
            workbook = openpyxl.load_workbook(config.out_xlsx, data_only=True)
            detail = workbook["检测明细"]
            counties = {row[7] for row in detail.iter_rows(min_row=2, values_only=True)}
            workbook.close()
            self.assertEqual(counties, {"万州区"})


if __name__ == "__main__":
    unittest.main()
